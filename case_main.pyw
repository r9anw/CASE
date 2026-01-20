# © Copyright Sussex Police
#
# This software is licensed 'as-is'.  You bear the risk of using it.  In
# consideration for use of the software, you agree that you have not relied upon
# any, and we have made no, warranties, whether oral, written, or implied, to
# you in relation to the software.  To the extent permitted at law, we disclaim
# any and all warranties, whether express, implied, or statutory, including, but
# without limitation, implied warranties of non-infringement of third party
# rights, merchantability and fitness for purpose.
# 
# In no event will we be held liable to you for any loss or damage (including
# without limitation loss of profits or any indirect or consequential losses)
# arising from the use of this software.
#
# Permission is granted to LAW ENFORCEMENT ONLY to use this software free of charge
# for any purpose and to alter it and redistribute it freely, subject to the
# following restrictions:
#
# 1. The origin of this software must not be misrepresented; you must not
# claim that you wrote the original software. If you use this software
# in a product, an acknowledgment in the form of
# "Copyright Sussex Police" in the product
# documentation would be appreciated but is not required.
#
# 2. Altered versions of the source code must be plainly marked as such, and
# must not be misrepresented as being the original software.
#
# 3. This copyright notice and disclaimer may not be removed from or varied in
# any copy of the software (whether in its original form or any altered version)
#
# DESCRIPTION:
# CASE
#
# Author:
# Ryan Ward - Sussex Police
#
# Contact: ryan.ward@sussex.pnn.police.uk
#
# Version: 3.0.0 - LargeScreen
############################################################################################
import sqlite3
import os
import hashlib
import re
from sys import argv
import string
import secrets
import shutil
import ctypes
import glob
import time
import datetime
from datetime import date
import tempfile
from time import sleep
import threading
from random import choice, shuffle
from string import printable
from tkinter import * 
from tkinter import ttk, messagebox
from tkinter import filedialog
import tkinter as tk
from tkinter import font
import subprocess
import docx
from docx import *
from docx import Document
from docx.shared import Pt
import xlsxwriter
from PIL import ImageTk, Image, ImageGrab
import random
from random import choice
#from wonderwords import RandomWord
from wordlist import *
#from individual_cases import *

############################################################################################
timestamp = '{:%Y-%m-%d %H:%M:%S}'.format(datetime.datetime.now())
datestamp = '{:%Y-%m-%d}'.format(datetime.datetime.now())
############################################################################################
############################################################################################
def createprofile():
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    try:
        c.execute("""create table assignedcases (id_no, year, full_dft_ref, crime_ref, no_of_exhibit, oic, datestamp, notes, analyst, notified, timestamp)""")
        c.execute("""create table analysts (name, warrant)""")
        c.execute("""create table version (swversion, build_date, install_timestamp)""")
        c.execute("""create table profile (casecount, case_dir, template_dir, contemp_file, contemp_copy, contemp_populate, sfr_file, sfr_copy, sfr_populate, disclosure_file, disclosure_copy, disclosure_populate, examiner, pw_length, pw_incdft, startup)""")
        c.execute("""create table cases (id_no, year, dft_ref_only, crime_ref, exhibit, bag_seal, operation, oic, suspect, g83, analyst, full_dft_ref, case_password, timestamp, exhib_notes, status)""")
        c.execute("""create table triage (casecount, year, dftonly, crime_ref, exhibit, bag_seal, operation, oic, suspect, g83, analyst, dft_ref, case_password, timestamp, exhib_notes, status)""")
        c.execute("""create table folders (id_no, f1, pp1, f2, pp2, f3, pp3, f4, pp4)""")
        c.execute("""create table officers (warrant, officer)""")
        c.execute("""create table notepad (notes, last_written)""")
        c.execute("""create table pindecryptlog (id_no, full_dft_ref, crime_ref, exhibit, oic, datestamp, days_running, analyst, pin_notes)""")
        c.execute("""create table opencases (id_no, full_dft_ref, crime_ref, no_of_exhib, oic, datestamp, note, analyst, suspect, type)""")
        c.execute("""create table teamleaders (name, warrant, key, salt)""")
        c.execute("""create table notifications (notific, timestamp)""")
        c.execute("""insert into notifications values ('Welcome!', ?);""", (timestamp,))
        c.execute("""insert into notepad values ('Welcome!', ?);""", (timestamp,))
        c.execute("""insert into analysts values ('R WARD', '32533')""")
        c.execute("""insert into version values ('200','181221', ?);""", (timestamp,))
        c.execute("""insert into profile values ('0','C:\\','C:\\', 'C:\\', '0', '0', 'C:\\', '0', '0', 'C:\\', '0', '0', 'R WARD 32533', '12', '1', 'Case');""")
        c.execute("""insert into folders values ('0','XRY', '0','UFED', '0','Graykey', '0','FTK', '0','Manual', '0','-', '0');""")
    except:
        pass
    conn.commit()
    conn.close()
############################################################################################
assignedcasevariables=[]
overviewvariables=[]
opencasevariables=[]
closedcasevariables=[]
pindecryptvariables=[]
individualcasevariables=[]
individual_pd_casevariables=[]
new_notification_variables=[]
startupmode=['Case Creator', 'Case']
status=['On Hold', 'Open - Imaging', 'Open - Analysis', 'With OIC', 'Pending Viewing']
property_of=['Victim', 'Suspect']
location=['R WARD 32533', 'Mobile LAB', 'Exhibit store']
action=['Admin', 'Exhibit movement', 'Strategy', 'RQC', 'Camera Calibration', 'Pre-imaging', 'PIN decryption', 'Imaging', 'Analysis', 'Grading', 'Reseal']

if not os.path.exists("C:\\Case Creator\\"):
    os.makedirs("C:\\Case Creator\\")
    print('made folder')
sqlprolocal = ("C:\\Case Creator\\casecreator.sqlite3")
if not os.path.isfile(sqlprolocal):
    createprofile()
    print('made sql pro')
else:
    pass
############################################################################################
def database():
        global sqllocal
        systempath = 'C:\\DFT Tools\\DFT Password Generator\\'
        if not os.path.exists(systempath):      
            os.makedirs(systempath)
            
        source = cwd+'\\Password.sqlite3'
        destination = 'C:\\DFT Tools\\DFT Password Generator\\Password.sqlite3'
        
        sqllocal = ("C:\\DFT_Password_Generator\\Password.sqlite3")
        # Check if the destination file does not exist
        if not os.path.exists(destination):
            shutil.copy(source, destination)
        else:
            pass

        sqllocal = ("C:\\DFT Tools\\DFT Password Generator\\Password.sqlite3")
############################################################################################
def gen_pw():
    global res, res_len
    os.chdir("C:\\Case Creator\\")
    words=[]
    for each in range(3):
        rw = random.choice(word_list)
        rw = str.replace(rw, "-'",'')
        words.append(str(rw))
    pw_words = []
    for each in words:
        t_word = ''.join(map(str, each))
        word = str.replace(t_word, "('",'')
        word = str.replace(word, "('",'')
        word = word.capitalize()
        pw_words.append(word)
    res = (str(pw_words[0]+pw_words[1]+pw_words[2]))
    res_len = (len(res))
    pw_test()
############################################################################################
def pw_test():
    if res_len <= 14:
        gen_pw()
    else:
        pass
############################################################################################
def callback_pw(sv):
    pass
############################################################################################
def remove_duplicates():
        global unique_list
        # Removing duplicates using dictionary keys
        unique_list = list(dict.fromkeys(word_list))
        write_to_db()
############################################################################################
def write_to_db():
        count=1
        for each in unique_list:
                conn = sqlite3.connect(sqllocal, isolation_level=None)
                conn.execute('pragma journal_mode=wal')
                c=conn.cursor()
                c.execute(""" insert into words values (?, ?)""", (count, each))
                count+=1
        conn.commit()
        conn.close()
############################################################################################
def exportpwtodesktop():
    global desktop
    case_password = pw_customa.get()
    # Prints the current user's directory: C:\Users\Ryan
    print(os.environ['USERPROFILE'])
    # Or print(os.environ.get("USERPROFILE"))
    desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    os.chdir(desktop)
    output_file = open(case_password + ".txt", "w")
    output_file.write(case_password)
    output_file.close()
############################################################################################
def exportpwtolocation():
    global desktop
    current_setting = startup_e2a.get()
    print('current_setting', current_setting)
    d = argv[1] if len(argv)>1 else filedialog.askdirectory(initialdir=desktop)
    D = os.path.realpath(d)
    if d == '':
        print('no data', current_setting)
        startup_e2b.delete(0,END)
        startup_e2b.insert(10,current_setting)
    else:
        print('qwerty')
        startup_e2b.delete(0,END)
        startup_e2b.insert(10,D)
############################################################################################
def copytoclip():
    clippw=pw_customa.get()
    print("clip Pw: "+clippw)
    cmd='echo '+clippw.strip()+'|clip'
    return subprocess.check_call(cmd, shell=True)
############################################################################################
def about_gui():
    global popup
    gui = ("popup")
    popup = Toplevel()
    popup.title("CASE - About")
    popup.resizable(width=False, height=False)

    frames=["popup"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(popup, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    Label(popup, text=" Version: 1.0.0").grid(row=0, column=1, sticky=W)
    Label(popup, text=" Build date: 11/09/2025").grid(row=1, column=1, sticky=W)
    Label(popup, text=" Author: Ryan Ward 32533").grid(row=2, column=1, sticky=W)
    Label(popup, text=" Contact: ryan.ward@sussex.police.uk").grid(row=3, column=1, sticky=W)
    Label(popup, text=" ").grid(row=5, column=0, columnspan=3, sticky=W)
    popup.attributes('-topmost',True)
    popup.mainloop()
############################################################################################
def zoom_in(event):
    pw_customb.config(font=zoomed_font)
    pw_customb.config(width=30)
############################################################################################
def zoom_out(event):
    pw_customb.config(font=default_font)
    pw_customb.config(width=45)
############################################################################################
def readDB():
    global sqlitedbs, sqlprolocal
    sqlprolocal = ("C:\\Case Creator\\casecreator.sqlite3")
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    sqlitedbs=[]
    c.execute(""" select * from databases""")
    for each in c:
        sqlitedbs.append(each)
        print(each[0])
    conn.close()
    sqlprolocal = (str(each[0]))
    sqlprolocal = (str(each[1]))
    print(sqlitedbs)

#sqlpro = ("\\\\sxiso1p.sdft.local\\Cases\\CaseCreator\\casecreator.sqlite3") # NEW SVR sdft.local
sqlpro = ("C:\\Case Creator\\casecreator.sqlite3")
sqlprolocal = ("C:\\Case Creator\\casecreator.sqlite3")
sqlitedbs = (sqlprolocal) #sqlpro
############################################################################################
cwd = os.getcwd()
file = ("\\ccicon.ico")
iconfile = (cwd + file)
if not os.path.isfile(iconfile):
    ICON = (b'\x00\x00\x01\x00\x01\x00\x10\x10\x00\x00\x01\x00\x08\x00h\x05\x00\x00'
            b'\x16\x00\x00\x00(\x00\x00\x00\x10\x00\x00\x00 \x00\x00\x00\x01\x00'
            b'\x08\x00\x00\x00\x00\x00@\x05\x00\x00\x00\x00\x00\x00\x00\x00\x00\x00'
            b'\x00\x01\x00\x00\x00\x01') + b'\x00'*1282 + b'\xff'*64

    _, ICON_PATH = tempfile.mkstemp()
    with open(ICON_PATH, 'wb') as iconfile:
        iconfile.write(ICON)
        iconfile.close()
    iconfile=(ICON_PATH)

updateversion = 146
updatebuild = 1020920
############################################################################################
def os_scandir():
    years=('16', '17', '18', '19', '20', '21')
    for YY in years:
        if not os.path.exists(case_dir + "\\DFT\\" + YY):
            pass
        else:
            print(case_dir)
            fu=os.listdir(case_dir + "\\DFT\\" + YY)
            print('fu:',fu)
############################################################################################
def popupclose():
    popup.destroy()
############################################################################################
def addclose():
    add.destroy()
############################################################################################
def search_pwclose():
    search_pw.destroy()
############################################################################################
def confirmclose():
    confirm.destroy()
############################################################################################
def close():
    global confirm, yes, yes_b, no, no_b
    gui = ("confirm")
    confirm = Toplevel()
#    confirm.geometry('315x200')
    confirm.title("Confirm?")
    confirm.resizable(width=False, height=False)
    confirm.lift(aboveThis=root)

    frames=["confirm"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(confirm, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1

    Label(confirm, text="Are you sure you wish to close this case?").grid(row=1, column=1, columnspan=3, sticky=EW)
    yes= StringVar()
    yes_b=Button(confirm, text="Yes", width=12, command=deleteentry2)
    yes_b.grid(row=2, column=1, columnspan=1, sticky=EW)

    no= StringVar()
    no_b=Button(confirm, text="No", width=12, command=confirmclose)
    no_b.grid(row=2, column=3, columnspan=1, sticky=EW)

    Label(confirm, text=" ").grid(row=3, column=1, columnspan=1, sticky=W)
    
    confirm.lift()
    confirm.attributes('-topmost',True)
    confirm.after_idle(root.attributes,'-topmost',False)
    confirm.mainloop()
############################################################################################
def dbcheck():
    for sql in sqlitedbs:
        # establish connection to database
        conn = sqlite3.connect(sql)
        c=conn.cursor()    


        # Request table names from database
        c.execute("SELECT name FROM sqlite_master WHERE type='table';")
        table_names = c.fetchall()

        # loop out table names from database
        for table in table_names:
            tablename = (table)
            
        # Request and loop out column names from database
            sql = "select * from %s" % tablename
            c.execute(sql)
            result = c.description
            col_names = [i[0] for i in c.description]

        conn.close()
############################################################################################
def writetoopencases():
    if any(dft_ref in s for s in opencases):
        pass
    else:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute(""" insert into opencases values(?,?,?,?,?,?,?,?,?,?);""", ('0', dft_ref, crime_ref, no_of_exhib, oic, datestamp, '', analyst, suspect, '-'))
        conn.commit()
        conn.close()
############################################################################################
def writecasehistory():
    global case_password
    case_password = res
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    print('Write case historyDB NAME:', sqlprolocal)
    c.execute(""" insert into cases values(?,?,?,?,?,?,?,?,?,?,?,?,?,?,?);""", (casecount, year, dftonly, crime_ref, exhib_ref, bag_seal, operation, oic, suspect, g83, analyst, dft_ref, case_password, timestamp, '-'))#, '-', '-'))
    conn.commit()
    conn.close()
    #writetoopencases()
############################################################################################
def writeofficers():
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    #c.execute(""" insert into officers values(?,?);""", (warrant, officer))
    print('DB NAME:', sqlprolocal)
    c.execute("SELECT * FROM officers WHERE warrant like '%'||?||'%'", (warrant,))
    officers=[]
    for each in c:
        officers.append(each)
    print(officers)
    if (len(officers)) == 0:
        print("no warrant match")
        c.execute(""" insert into officers values(?,?);""", (warrant, officer))
    else:
        print("warrant match")
        c.execute(""" update officers SET officer=? WHERE warrant=?;""", (warrant, officer))
    conn.commit()
    conn.close()
############################################################################################
def readanalysts():
    global analysts
    try:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        print(sqlprolocal)
        c=conn.cursor()
        analysts=[]
        c.execute(""" select * from analysts""")
        for each in c:
            add=(str(each[0] + ' ' + (str(each[1]))))
            analysts.append(add)
        conn.close()
    except:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        print(sqlprolocal)
        c=conn.cursor()
        analysts=[]
        c.execute(""" select * from analysts""")
        for each in c:
            add=(str(each[0] + ' ' + each[1]))
            analysts.append(add)
        conn.close()
    if (len(analysts))==0:
        analysts.append('R WARD 32533')
    else:
        pass
    print('analysts', analysts)
############################################################################################
def readversion():
    global version, build_date
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()

    swversion=[]
    c.execute(""" select * from version""")

    for each in c:
        swversion.append(each)
    conn.close()
    version = (int(each[0]))
    build_date = (int(each[1]))
    print("Version:", version)
    build_date = each[1]
############################################################################################
def readopencases():
    global opencases, casecount, case_dir, template_dir, contemp_file, examiner
    try:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        opencases=[]
        c.execute(""" select * from opencases WHERE analyst like '%'||?||'%'""", (analyst,))
        for each in c:
            opencases.append(each)
        conn.close()
        
    except:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        opencases=[]
        c.execute(""" select * from opencases WHERE analyst like '%'||?||'%'""", (analyst,))
        for each in c:
            opencases.append(each)
            opencases_easy_view.append(each[11]+" - "+each[3])
        conn.close()
    casecount = (len(cases))
############################################################################################
def readcases():
    global case_password, opencases_easy_view, cases, casecount, case_dir, template_dir, contemp_file, examiner
    print('in readcases 370', sqlprolocal)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    cases=[]
    opencases_easy_view=[]
    case_password=[]
    c.execute(""" select * from cases""")
    for each in c:
        cases.append(each)
        opencases_easy_view.append(each[11]+" - "+each[3])
        case_password.append(each[11]+" - "+each[3])
    conn.close()
    casecount = (len(cases))
    opencases_easy_view = list(set(opencases_easy_view))
    case_password = list(set(case_password))
###########################################################################################
def readofficers():
    global cases, casecount, case_dir, template_dir, contemp_file, examiner
    try:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        officers=[]
        c.execute(""" select * from officers""")
        for each in c:
            officers.append(each)
        conn.close()
    except:
        pass
############################################################################################    
def readprofile():
    global startup_method, prodata, cases, case_dir, template_dir, contemp_file, contemp_copy, contemp_populate, sfr_file, sfr_copy, sfr_populate, disclosure_file, disclosure_copy, disclosure_populate, examiner, analyst, pw_length, pw_incdft
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    prodata=[]

    c.execute(""" select * from profile""")
    for each in c:
        print('in read profile: ', each)
        case_dir = each[1]
        template_dir = each[2]
        contemp_file = each[3]
        contemp_copy = each[4]
        contemp_populate = each[5]
        sfr_file = each[6]
        sfr_copy =each[7]
        sfr_populate = each[8]
        disclosure_file = each[9]
        disclosure_copy = each[10]
        disclosure_populate = each[11]
        examiner = each[12]
        analyst = each[12]
        pw_length = each[13]
        pw_incdft = each[14]
        startup_method = each[15]
    
    prodata.append(case_dir)
    prodata.append(template_dir)
    prodata.append(contemp_file)
    prodata.append(contemp_copy)
    prodata.append(contemp_populate)
    prodata.append(sfr_file)
    prodata.append(sfr_copy)
    prodata.append(sfr_populate)
    prodata.append(disclosure_file)
    prodata.append(disclosure_copy)
    prodata.append(disclosure_populate)
    prodata.append(examiner)
    prodata.append(pw_length)
    prodata.append(pw_incdft)
############################################################################################
def readfolders():
    global fld1, fld2, fld3, fld4, pp1, pp2, pp3, pp4
    folders=[]
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" select * from folders""")
    for each in c:
        folders.append(each)
        print(each)
        fld1 = each[1]
        pp1 = each[2]
        fld2 = each[3]
        pp2 = each[4]
        fld3 = each[5]
        pp3 = each[6]
        fld4 = each[7]
        pp4 = each[8]
    conn.close()
    # Catch error if c == 0
    if len(folders) == 0:
        fld1 = '-'
        pp1 = 0
        fld2 = '-'
        pp2 = 0
        fld3 = '-'
        pp3 = 0
        fld4 = '-'
        pp4 = 0
############################################################################################    
def updatefolders():
    global fld1, fld2, fld3, fld4
    fld1=e13.get()
    fld2=e14.get()
    fld3=e15.get()
    fld4=e16.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" DELETE FROM folders""")
    c.execute(""" insert into folders values ('1', ?, ?, ?, ?, ?, ?, ?, ?);""", (fld1, pp1, fld2, pp2, fld3, pp3, fld4, pp4))
    conn.commit()
    conn.close()
    settings_tab.update
    settingslock()
############################################################################################
############################################################################################
def searchcases():
    print('db name:', sqlpro)
    try:
        print('try ')
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute(""" select * from cases""")
        c.execute("SELECT * FROM cases WHERE full_dft_ref like '%'||?||'%'", (e1a.get().upper(),))
        for each in c:
            if each[11] == (e1a.get().upper()):
                enableload()
            else:
                disableload()
        conn.commit()
        conn.close()
    except:
        pass
############################################################################################
def searchcases_crime():
    print('db name:', sqlpro)
    try:
        print('try ')
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute(""" select * from cases""")
        c.execute("SELECT * FROM cases WHERE crime_ref like '%'||?||'%'", (e2a.get().upper(),))
        for each in c:
            if each[3] == (e2a.get().upper()):
                enableload_crime()
            else:
                disableload_crime()
        conn.commit()
        conn.close()
    except:
        pass
############################################################################################
def duplicate_entry():
    check=0
    x=e1a.get().upper()
    y=e3a.get().upper()
    z=e4a.get().upper()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" select * from cases""")
    c.execute("SELECT * FROM cases WHERE full_dft_ref like '%'||?||'%' and exhibit = ? and bag_seal = ?", (x, y, z,))
    print('get data: ', x, y, z)
    
    for each in c:
        print(each)
        print('data from loop:', each[11], each[4], each[5])
        if each[11] == x:
            check+=1
        elif each[4] == y:
            check+=1
        elif each[5] == z:
            check+=1
        else:
            pass
    print('check value: ', check)
    if check >= 3:
        print('duplicate case')
        check=0
        duplicate_next_step()
    else:
        check=0
        getData()
    conn.commit()
    conn.close()
############################################################################################
def duplicate_next_step():
    global popup
    x=e1a.get().upper() # DFT
    y=e2a.get().upper() # CRIME
    y=e3a.get().upper()
    z=e4a.get().upper()
    oic = (e5a.get().upper())
    suspect = (e7a.get().upper())
    gui = ("popup")
    popup = Toplevel()
    popup.title("Duplicate entry...")
    popup.resizable(width=False, height=False)
    popup.lift(aboveThis=root)

    frames=["popup"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(popup, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    Label(popup, text=" DFT Ref: "+x+" - Exhibit Ref: "+y+" - Bag Seal: "+z+"...").grid(row=1, column=1, columnspan=5, sticky=W)
    Label(popup, text=" What do you want to do?").grid(row=2, column=1, columnspan=5, sticky=W)
    Label(popup, text=" ").grid(row=3, column=0, columnspan=3, sticky=W)


    Button(popup, text="Re-enter details", width=17, command=popupclose).grid(row=4, column=1, sticky=W)
    Button(popup, text="Add to CASE only", width=17, command=change_case_status).grid(row=4, column=3, sticky=W)
    Button(popup, text="Add Folder & CASE",  width=17, command=getData).grid(row=4, column=5, sticky=W)
    Label(popup, text=" ").grid(row=5, column=0, columnspan=3, sticky=W)
    popup.attributes('-topmost',True)
    popup.mainloop() 
############################################################################################
def change_case_status():
    popupclose()
    x=e1a.get().upper()
    y=e3a.get().upper()
    z=e4a.get().upper()
    oic = (e5a.get().upper())
    suspect = (e7a.get().upper())
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("SELECT * FROM opencases WHERE full_dft_ref like '%'||?||'%'", (x,))
    print('get data: ', x, y, z)
    entries=[]
    for each in c:
        print(each)
        entries.append(each)
    if (len(entries)) == 0:
        c.execute(""" insert into opencases values(?,?,?,?,?,?,?,?,?);""", ('0', x, y, 'Update', oic, datestamp, '', analyst, suspect))
        c.execute("UPDATE cases SET status=? WHERE full_dft_ref like '%'||?||'%'", ('open', x,))
    else:
        pass
    conn.commit()
    conn.close()
    new_case.destroy()
    #overview()
############################################################################################   
def loadcases_dft():
    global casetype, case_password_loaded
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.update()
        else:
            pass
    except:
        pass
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("SELECT * FROM cases WHERE full_dft_ref like '%'||?||'%'", (e1a.get().upper(),))
    '''if (len(c))==0:
        c.execute("SELECT * FROM triage WHERE full_dft_ref like '%'||?||'%'", (e1a.get().upper(),))
    else:
        c=['-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-']'''
    for each in c:
        # Delete exhibit ref
        e3b.delete(0, END)
        # Delete bag seal ref
        e4b.delete(0, END)
        # DFT ref
        # e1a.set(each[11])
        # Crime ref
        e2a.set(each[3])
        # OIC
        e5a.set(each[7])
        # Operation
        e6a.set(each[6])
        # Suspect
        e7a.set(each[8])
        # Case Password
        try:
            if 'normal' == root.state():
                case_password_loaded = each[12]
                #startup_e1a.set(case_password)
                #startup_e1b.config(state='disable')
                Refresh.config(state='disable')
                root.update()
            elif 'normal' == new_triage.state():
                new_triage.update()
            else:
                pass
        except:
            pass
        
    conn.commit()
    conn.close()
    LoadButtonb.config(state='disable')
    LoadButtonb1.config(state='disable')
    casetype = 'Loaded'
############################################################################################   
def loadcases_crime():
    global casetype, case_password_loaded
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.update()
        else:
            pass
    except:
        pass
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("SELECT * FROM cases WHERE crime_ref like '%'||?||'%'", (e2a.get().upper(),))
    '''if (len(c))==0:
        c.execute("SELECT * FROM triage WHERE full_dft_ref like '%'||?||'%'", (e1a.get().upper(),))
    else:
        c=['-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-', '-']'''
    for each in c:
        # Delete exhibit ref
        e3b.delete(0, END)
        # Delete bag seal ref
        e4b.delete(0, END)
        # DFT ref
        e1a.set(each[11])
        # Crime ref
        e2a.set(each[3])
        # OIC
        e5a.set(each[7])
        # Operation
        e6a.set(each[6])
        # Suspect
        e7a.set(each[8])
        # Case Password
        try:
            if 'normal' == new_case.state():
                case_password_loaded = each[12]
                #startup_e1a.set(case_password)
                #startup_e1b.config(state='disable')
                Refresh.config(state='disable')
                root.update()
            elif 'normal' == new_triage.state():
                new_triage.update()
            else:
                pass
        except:
            pass
        
    conn.commit()
    conn.close()
    LoadButtonb.config(state='disable')
    LoadButtonb1.config(state='disable')
    casetype = 'Loaded'
############################################################################################
def callback(sv):
    print(sv.get())
############################################################################################
##def callback_pw(sv):
##    new_pw_guigenPw()
############################################################################################    
def callback_pw_search(sv):
    global no_of_results
    pw_search_textbox.delete(1.0, END)
    search_pw_result=[]
    get=Newentry1a.get()
    print('sql pro: ', sqlpro)
    conn = sqlite3.connect(sqlpro, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("SELECT * FROM cases WHERE full_dft_ref like '%'||?||'%'", (get,))
    for each in c:
        add=(str(each[11] + ' - ' + each[12]))
        search_pw_result.append(add)
        print(add)
        pw_search_textbox.insert(INSERT, add+'\n')
    conn.commit()
    conn.close()
    no_of_results = (len(search_pw_result))
############################################################################################
def callback_new_triage(sv):
    dft_ref=(str(e1a.get().upper()))
    dft_ref=(dft_ref[:7])
    if len(dft_ref) == 3:
        if dft_ref[-1:3] != '-':
            e1a.set(dft_ref[:2])
            dft_ref=(dft_ref[:2])
        else:
            e1a.set(dft_ref[:7])
    elif len(dft_ref) >= 7:
        e1a.set(dft_ref[:7])
        dftrefcheck()
    else:
        e1a.set(dft_ref[:7])
        e1b.config({"background": "white"})
    new_triage.update()
    searchcases() 
############################################################################################
def callback1(sv):
    dft_ref=(str(e1a.get().upper()))
    dft_ref=(dft_ref[:7])
    if len(dft_ref) == 3:
        if dft_ref[-1:3] != '-':
            e1a.set(dft_ref[:2])
            dft_ref=(dft_ref[:2])
        else:
            e1a.set(dft_ref[:7])
    elif len(dft_ref) >= 7:
        e1a.set(dft_ref[:7])
        dftrefcheck()
    else:
        e1a.set(dft_ref[:7])
        e1b.config({"background": "white"})
    new.update()
    searchcases() 
############################################################################################
def callback2(sv):
    crime_ref=(str(e2a.get().upper()))
    print(sv.get())
    
    if len(crime_ref) >= 11:
        crimerefcheck()
    else:
        e2b.config({"background": "white"})
    new.update()
    searchcases_crime()
    updatePw()  
############################################################################################
def callback3(sv):
    exhib_ref=(str(e3a.get().upper()))
    print(sv.get())
    if len(exhib_ref) >= 2:
        exhibrefcheck()
    else:
        e3b.config({"background": "white"})
############################################################################################
def callbacke26(sv):
    global pp1
    pp1 = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE folders SET pp1 = ? WHERE _rowid_ = 1', (pp1,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke27(sv):
    global pp2
    pp2 = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE folders SET pp2 = ? WHERE _rowid_ = 1', (pp2,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke28(sv):
    global pp3
    pp3 = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE folders SET pp3 = ? WHERE _rowid_ = 1', (pp3,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke29(sv):
    global pp4
    pp4 = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE folders SET pp4 = ? WHERE _rowid_ = 1', (pp4,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke30(sv):
    global pp5
    pp5 = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE folders SET pp5 = ? WHERE _rowid_ = 1', (pp5,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke31(sv):
    global pp6
    pp6 = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE folders SET pp6 = ? WHERE _rowid_ = 1', (pp6,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke33(sv):
    global contemp_copy
    contemp_copy = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET contemp_copy = ? WHERE _rowid_ = 1', (contemp_copy,))
    if contemp_copy == 0:
        contemp_populate = int(0)
        c.execute('UPDATE profile SET contemp_populate = ? WHERE _rowid_ = 1', (contemp_populate,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke34(sv):
    global contemp_populate
    contemp_populate = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET contemp_populate = ? WHERE _rowid_ = 1', (contemp_populate,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke35(sv):
    global sfr_copy
    sfr_copy = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET sfr_copy = ? WHERE _rowid_ = 1', (sfr_copy,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke36(sv):
    global sfr_populate
    sfr_populate = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET sfr_populate = ? WHERE _rowid_ = 1', (sfr_populate,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke37(sv):
    global disclosure_copy
    print(disclosure_copy)
    disclosure_copy = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET disclosure_copy = ? WHERE _rowid_ = 1', (disclosure_copy,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke38(sv):
    global disclosure_populate
    disclosure_populate = sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET disclosure_populate = ? WHERE _rowid_ = 1', (disclosure_populate,))
    conn.commit()
    conn.close()
############################################################################################
def e33checkbox():
    if contemp_copy ==0:
        e34a.set(0)
    new.update	
############################################################################################
def e35checkbox():
    if sfr_copy ==0:
        e36a.set(0)
    new.update	
############################################################################################
def e37checkbox():
    if disclosure_copy ==0:
        e38a.set(0)
    new.update	
############################################################################################
def callbackstartup_e1a():
    if pw_incdft == 0:
        startup_e1a.set('ryan')
        new.update()
############################################################################################ 
def genPw():
    global pw, pwdft
    if casetype != 'Loaded':
        dft_ref=(str(e1a.get().upper()))
        pwlength = (int(e21a.get()))
        if (type(pwlength)) == int:
            length=pwlength
            e21a.set(length)
        else:
            length = 12
            e21a.set(length)

        if len(dft_ref)==0:
            '''P = list(printable[:62]) + ['@', '$', '%', '!', '+']
            shuffle(P)
            dft_ref=(str(e1a.get().upper()))
            pw=''.join(P[:length])
            pwdft=dft_ref+''.join(pw)
            print(pw)
            incdft=e22a.get()
            print(incdft)'''
            incdft=e22a.get()
            if incdft == 1:
                print("on")
                P = list(printable[:62]) + ['@', '$', '%', '!', '+']
                shuffle(P)
                dft_ref=(str(e1a.get().upper()))
                pw=''.join(P[:length])
                pwdft=dft_ref+''.join(pw)
                startup_e1a.set(pwdft)
                root.update()
            else:
                print("off")
                P = list(printable[:62]) + ['@', '$', '%', '!', '+']
                shuffle(P)
                pw=''.join(P[:length])
                print(pw)
                startup_e1a.set(pw)
                root.update()
        else:
            P = list(printable[:62]) + ['@', '$', '%', '!', '+']
            shuffle(P)
            dft_ref=(str(e1a.get().upper()))
            pw=''.join(P[:length])
            pwdft=dft_ref+''.join(pw)
            print(pwdft)
            startup_e1a.set(pwdft)
            root.update()
    else:
        startup_e1a.set(case_password)
    root.update()
############################################################################################ 
def zznew_pw_guigenPw():
    global pw, pwdft
    casetype = ''
    if casetype != 'Loaded':
        #dft_ref=(str(e1a.get().upper()))
        try:
            pwlength = (int(e21a.get()))
        except:
            conn = sqlite3.connect(sqlpro, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("SELECT pw_length FROM profile")
            for each in c:
                print(each)
                pwlength = each[0]
            conn.close()
            casecount = (len(cases))
            
        if (type(pwlength)) == int:
            length=pwlength
            e21a.set(length)
        else:
            length = 12
            e21a.set(length)


    P = list(printable[:62]) + ['@', '$', '%', '!', '+']
    shuffle(P)
    customise=(str(pw_customa.get().upper()))
    pw=''.join(P[:length])
    pwdft=customise+''.join(pw)
    print(pwdft)
    startup_e1a.set(pwdft)
    pw_generator.update()
############################################################################################
def new_pw_guiupdatePw():
    global pw, pwdft
    select_random_entry()
    pw_customa.set(res)
    pw_generator.update()
############################################################################################
def updatePw():
    length=12
    P = list(printable[:62]) + ['@', '$', '%', '!', '+']
    shuffle(P)
    dft_ref=(str(e1a.get().upper()))
    pw=''.join(P[:length])
    pwdft=dft_ref+''.join(pw)
    new.update()
############################################################################################
def dftrefcheck():
    e1b.config({"background": "pale green"})
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.update()
        else:
            pass
    except:
        pass
############################################################################################
def crimerefcheck():
    e2b.config({"background": "pale green"})
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.update()
        else:
            pass
    except:
        pass
############################################################################################
def exhibrefcheck():
    e3b.config({"background": "pale green"})
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.update()
        else:
            pass
    except:
        pass
############################################################################################
def enablesubmit():
    livedft=callback1()
    print("livedft: ", livedft)
    root.update()
############################################################################################
def enableload():
    LoadButtonb.config(state='normal')
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.destroy()
        else:
            pass
    except:
        pass
    
############################################################################################
def disableload():
    LoadButtonb.config(state='disable')
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.update()
        else:
            pass
    except:
        pass
############################################################################################
def enableload_crime():
    LoadButtonb1.config(state='normal')
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.destroy()
        else:
            pass
    except:
        pass
    
############################################################################################
def disableload_crime():
    LoadButtonb1.config(state='disable')
    try:
        if 'normal' == new_case.state():
            root.update()
        elif 'normal' == new_triage.state():
            new_triage.update()
        else:
            pass
    except:
        pass    
############################################################################################
def getData():
    global dft_ref, crime_ref, exhib_ref, bag_seal, oic, operation, suspect, g83, analyst, analystinitials, warrant, officer, sus_op_crime, no_of_exhib, casetype
    casetype = 'new'
    oicmatch=['0']
    dft_ref = (e1a.get().upper())
    crime_ref = (e2a.get().upper())
    exhib_ref = (e3a.get().upper())
    bag_seal = (e4a.get().upper())
    if (len(bag_seal))==0:
            bag_seal =("-")
    oic = (e5a.get().upper())
    if (len(oic))==0:
            oic =("-")
            officer = oic
            warrant =("-")
    else:
        cut = int(len(oic))-5
        print("Cut"), print(cut)
        officer = oic[:-5]
        warrant = oic[cut:]
        print("officer"), print(officer), print("warrant"), print(warrant)
    operation = (e6a.get().upper())
    if (len(operation))==0:
            operation =("-")
    suspect = (e7a.get().upper())
    if (len(suspect))==0:
            suspect =("-")
    
    g83 = (e8.get().upper())
    if (len(g83))==0:
            g83 =("-")
    #analyst = (e9a.get())
    initials = (analyst.split())

    init1=(str(initials[0]))
    
    if (len(init1))==1:
        len1=1
        init1=init1[:len1]
    else:
        len1=(int(len(init1)-1))
        init1=init1[:-len1]
             
    init2=(str(initials[1]))
    len2=int(len(init2)-1)
    init2=init2[:-len2]

    analystinitials = init1+''.join(init2)

    slashes = ('/', ':')
    for item in slashes:
        dft_ref = str.replace(dft_ref, item,'-')
        crime_ref = str.replace(crime_ref, item,'-')
        exhib_ref = str.replace(exhib_ref, item,'-')
        officer = str.replace(officer, item,'-')
        warrant = str.replace(warrant, item,'-')
        suspect = str.replace(suspect, item,'-')
        operation = str.replace(operation, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        dft_ref = str.replace(dft_ref, item,'')
        crime_ref = str.replace(crime_ref, item,'')
        exhib_ref = str.replace(exhib_ref, item,'')
        officer = str.replace(officer, item,'')
        warrant = str.replace(warrant, item,'')
        suspect = str.replace(suspect, item,'')
        operation = str.replace(operation, item,'')
    casespecifics=(dft_ref, crime_ref, exhib_ref, bag_seal, oic, operation, suspect, g83, analyst)
    sus_op_crime = suspect
    writeofficers()
    for item in casespecifics:
        print(item)
    no_of_exhib = '-'
    processEntry()
############################################################################################
def processEntry():
    global year, dftonly, sus_op_crime
    print('in process entry 1254')
    if len(dft_ref)!=7: 
        e1b.config({"background": "orange red"})
        root.update()
        dftok=0
    else:
        dftok=1
    
    if len(crime_ref)<=10: 
        e2b.config({"background": "orange red"})
        root.update()
        crimeok=0
    else:
        crimeok=1
        
    if len(exhib_ref)<=1: 
        e3b.config({"background": "orange red"})
        root.update()
        exhibok=0
    else:
        exhibok=1

    
        
    if dftok ==1:
        if crimeok ==1:
            if exhibok==1:
                year=dft_ref[:-5]
                dftonly=dft_ref[3:]
                print("year: "+(str(year)))
                print("dft only: "+dftonly)
                print('ROOOT_STATE', new.state())
                new_state = new.state()
                
                if new_state == 'normal':
                    print('processEntry(): new case', timestamp)
                    readcases()
                    readopencases()
                    lockCase()
                    createCase()
                else:
                    print('ROOOT_STATE 2', new.state())
                    print('passed 1', timestamp)
###########################################################################################
def casefolderStruct():
    global popup
    gui = ("popup")
    popup = Toplevel()
    popup.title("!")
    popup.resizable(width=False, height=False)
    popup.lift(aboveThis=root)

    frames=["popup"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(popup, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    Label(popup, text=" The server creates case folders using the Suspect's name...").grid(row=1, column=1, columnspan=5, sticky=W)
    Label(popup, text=" Suspect field is empty. What do you want to do?").grid(row=2, column=1, columnspan=5, sticky=W)
    Label(popup, text=" ").grid(row=3, column=0, columnspan=3, sticky=W)


    Button(popup, text="Enter Suspect name", width=17, command=entersuspectname).grid(row=4, column=1, sticky=W)
    Button(popup, text="Use Operation name", width=17, command=useoperationname).grid(row=4, column=3, sticky=W)
    Button(popup, text="Use Crime reference",  width=17, command=usecrimereference).grid(row=4, column=5, sticky=W)
    Label(popup, text=" ").grid(row=5, column=0, columnspan=3, sticky=W)
    
    popup.mainloop() 
############################################################################################
def entersuspectname():
    popup.destroy()
############################################################################################    
def useoperationname():
    global sus_op_crime, year, dftonly, startup, submitb
    if (len(operation))<=3:
        e1b.config(state='disabled')
        e2b.config(state='disabled')
        e3b.config(state='disabled')
        e4.config(state='disabled')
        e5b.config(state='disabled')
        e6b.config(state='normal')
        e7b.config(state='disabled')
        e8.config(state='disabled')
        e24b.config(state='disabled')
        submitb.config(command=processopentry)
        submitb.config(state='normal')
        startup.update
    else:
        e7a.set('-')
        e7b.config({"background": "white"})
        sus_op_crime = operation
        year=dft_ref[:-5]
        dftonly=dft_ref[3:]
        readcases()
        lockCase()
        createCase()
    popup.destroy()
############################################################################################    
def usecrimereference():
    global sus_op_crime, year, dftonly, startup, submitb
    if (len(crime_ref))<=3:
        e1b.config(state='disabled')
        e2b.config(state='disabled')
        e3b.config(state='normal')
        e4.config(state='disabled')
        e5b.config(state='disabled')
        e6b.config(state='disabled')
        e7b.config(state='disabled')
        e8.config(state='disabled')
        e24b.config(state='disabled')
        submitb.config(command=processcrimeentry)
        submitb.config(state='normal')
        startup.update
    else:
        e7a.set('-')
        e7b.config({"background": "white"})
        sus_op_crime = crime_ref
        year=dft_ref[:-5]
        dftonly=dft_ref[3:]
        readcases()
        lockCase()
        createCase()
    popup.destroy()
############################################################################################
def processopentry():
    global year, dftonly, sus_op_crime
    operation = (e6a.get().upper())
    slashes = ('/', ':')
    for item in slashes:
        operation = str.replace(operation, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        operation = str.replace(operation, item,'')
    casespecifics=(dft_ref, crime_ref, exhib_ref, bag_seal, oic, operation, suspect, g83, analyst)
    sus_op_crime = operation
    year=dft_ref[:-5]
    dftonly=dft_ref[3:]
    readcases()
    lockCase()
    createCase()
############################################################################################
def processcrimeentry():
    global year, dftonly, crime_ref, sus_op_crime
    crime_ref = (e2a.get().upper())
    slashes = ('/', ':')
    for item in slashes:
        crime_ref = str.replace(crime_ref, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        crime_ref = str.replace(crime_ref, item,'')
    casespecifics=(dft_ref, crime_ref, exhib_ref, bag_seal, oic, operation, suspect, g83, analyst)
    sus_op_crime = crime_ref
    year=dft_ref[:-5]
    dftonly=dft_ref[3:]
    readcases()
    lockCase()
    createCase()
############################################################################################
def createCase():
    global DftFolder, cf_0, cf_1, cf_2, casecount, year, dftonly, crime_ref, exhib_ref, bag_seal, operation, oic, g83, analyst, dft_ref, case_password, case_password_loaded, timestamp
    #readprofile()
    print('in createCase() 1430')
##    cf_0 = case_dir+'\\'+crime_ref+'\\'
##    if not os.path.exists(cf_0):      
##        os.makedirs(cf_0)

    NicheFolder = (case_dir+"\\"+crime_ref+"\\")
    if not os.path.exists(NicheFolder):
        os.makedirs(NicheFolder)

    os.chdir(NicheFolder)
        
##    cf_1 = cf_0+dft_ref+'\\'
##    if not os.path.exists(cf_1):      
##        os.makedirs(cf_1)

    DftFolder = (NicheFolder+dft_ref+"\\")
    if not os.path.exists(DftFolder):
        os.makedirs(DftFolder)

    os.chdir(DftFolder)

    case_data = 'Case Data'
    case_dataFolder = (DftFolder+case_data+"\\")
    if not os.path.exists(case_dataFolder):
        os.makedirs(case_dataFolder)

    os.chdir(case_dataFolder)

    ExhibitFolder_S = (case_dataFolder+"\\"+exhib_ref+"\\")
    if not os.path.exists(ExhibitFolder_S):
        os.makedirs(ExhibitFolder_S)

    os.chdir(DftFolder)

    generated = 'Generated Material'
    GeneratedFolder = (DftFolder+generated+"\\")
    if not os.path.exists(GeneratedFolder):
        os.makedirs(GeneratedFolder)

    os.chdir(GeneratedFolder)

    ExhibitFolder_G = (GeneratedFolder+"\\"+exhib_ref+"\\")
    if not os.path.exists(ExhibitFolder_G):
        os.makedirs(ExhibitFolder_G)

    os.chdir(ExhibitFolder_G)
    
    # Add Photosgraphs folder to Crime/Niche>DFT ref>Generated Material>Exhibit Ref folder
    folders=("Photographs", )
    for folder in folders:
        if not os.path.exists(folder):
            os.makedirs(folder)

    os.chdir(DftFolder)

    forensic = 'Forensic Images'
    forensicFolder = (DftFolder+forensic+"\\")
    if not os.path.exists(forensicFolder):
        os.makedirs(forensicFolder)

    os.chdir(forensicFolder)

    ExhibitFolder_G = (forensicFolder+"\\"+exhib_ref+"\\")
    if not os.path.exists(ExhibitFolder_G):
        os.makedirs(ExhibitFolder_G)
    
    os.chdir(DftFolder)    
##    
##    cf_2_1 = cf_1+'Source Material'+'\\'
##    if not os.path.exists(cf_2_1):      
##        os.makedirs(cf_2_1)
##
##    cf_2_2 = cf_1+'Generated Material'+'\\'
##    if not os.path.exists(cf_2_2):      
##        os.makedirs(cf_2_2)
##        
##    cf_3 = cf_2_1+exhib_ref+'\\'
##    if not os.path.exists(cf_3):      
##        os.makedirs(cf_3)
##
##    cf_3 = cf_2_2+exhib_ref+'\\'
##    if not os.path.exists(cf_3):      
##        os.makedirs(cf_3)
          

    file_name = "("+dft_ref+") Case Details - "+exhib_ref+".txt"
    text_file = open(file_name, "w")
    text_file.write('DFT Ref: '+dft_ref+'\n')
    text_file.write('Exhibit Ref: '+exhib_ref+'\n')
    text_file.write('Bag Seal: '+bag_seal+'\n')
    text_file.write('Property Ref: '+g83+'\n')
    text_file.write('OIC: '+oic+'\n')
    text_file.write('Operation Name: '+operation+'\n')
    text_file.write('Suspect: '+sus_op_crime)

    text_file.close()
    
##    # Add Analyst name to Analyst folder
##    add_analyst = analyst
##    if not os.path.exists(add_analyst):      
##        os.makedirs(add_analyst)
##    # Add Exhibit Ref folder to Analyst>Analyst name folder
##    add_exhibit = analyst+'\\'+exhib_ref
##    if not os.path.exists(add_exhibit):      
##        os.makedirs(add_exhibit)
##
##    # Add Exhibit Ref folder to Exhibit>Exhibit Ref folder
##    add_exhibit = exhib_ref
##    if not os.path.exists(add_exhibit):      
##        os.makedirs(add_exhibit)
##    os.chdir(exhib_ref+'\\')
##    # Add Photosgraphs & Extractions folder to Exhibit>Exhibit Ref folder
##    folders=("Photographs", )
##    for folder in folders:
##        if not os.path.exists(folder):
##            os.makedirs(folder)
    
    #os.chdir(case_dir)
    #print(case_dir)
    print('case type: ', casetype)
    
##    if casetype == 'Loaded':
##        case_password_loaded
##        case_password=case_password_loaded
##        print('createCase(): case_password: ',case_password)
##    else:
##        length=12
##        P = list(printable[:62]) + ['@', '$', '%', '!', '+']
##        shuffle(P)
##        dft_ref=(str(e1a.get().upper()))
##        pw=''.join(P[:length])
##        case_password=dft_ref+''.join(pw)
##
##    
##    try:_password_loaded)
##    except:
##        print('CP: ',case_password)

    gen_pw()
    #Launch case folder
    os.startfile(DftFolder)
    new.destroy()
    t1 = threading.Thread(target=writecasehistory)
    t2 = threading.Thread(target=folderautobuild)#checktemplates)
    t3 = threading.Thread(target=exportpw)
    t1.start()
    t2.start()
    t3.start()
    
##    for each in assignedcases:
##        print('each assigned case: ',each)
##        if each[1] == dft_ref:
##            print('Remove assigned case: ', each[1])
##            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
##            conn.execute('pragma journal_mode=wal')
##            c=conn.cursor()
##            c.execute("DELETE FROM assignedcases WHERE full_dft_ref like '%'||?||'%' ", (each[1],))
##            conn.commit()
##            conn.close()
##            readassignedcases()
##        else:
##            pass

############################################################################################
def createfolder1():
    global fld1, fld2, fld3, fld4, fld5, fld6
    '''if len(e1a.get())!=7:
        e1b.config({"background": "orange red"})
        if len(e2a.get())<=1:
            e2b.config({"background": "orange red"})
            if len(e3a.get())<=1:
                e3b.config({"background": "orange red"})
    else:'''
    slashes = ('/', ':')
    for item in slashes:
        fld1 = str.replace(fld1, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        fld1 = str.replace(fld1, item,'')
    path = case_dir+'\\'+"DFT"+'\\'+year+'\\'+dftonly+' - '+sus_op_crime+'\\'+exhib_ref+'\\'+fld1

    if not os.path.exists(path):      
        os.makedirs(path)
        '''folder1a = StringVar()
        folder1b = Button(startup, text=fld1,  width= 10, state=DISABLED, command=0).grid(row=3, column=6, sticky=EW)
        startup.update'''
###########################################################################################
def createfolder2():
    global fld1, fld2, fld3, fld4, fld5, fld6
    '''if len(e1a.get())!=7:
        e1b.config({"background": "orange red"})
        if len(e2a.get())<=1:
            e2b.config({"background": "orange red"})
            if len(e3a.get())<=1:
                e3b.config({"background": "orange red"})
    else:'''
    slashes = ('/', ':')
    for item in slashes:
        fld2 = str.replace(fld2, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        fld2 = str.replace(fld2, item,'')
    path = case_dir+'\\'+"DFT"+'\\'+year+'\\'+dftonly+' - '+sus_op_crime+'\\'+exhib_ref+'\\'+fld2
    if not os.path.exists(path):      
        os.makedirs(path)
        '''folder2a = StringVar()
        folder2b = Button(startup, text=fld2,  width= 10, state=DISABLED, command=0).grid(row=4, column=6, sticky=EW)
        startup.update'''
###########################################################################################
def createfolder3():
    global fld1, fld2, fld3, fld4, fld5, fld6
    '''if len(e1a.get())!=7:
        e1b.config({"background": "orange red"})
        if len(e2a.get())<=1:
            e2b.config({"background": "orange red"})
            if len(e3a.get())<=1:
                e3b.config({"background": "orange red"})
    else:'''
    slashes = ('/', ':')
    for item in slashes:
        fld3 = str.replace(fld3, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        fld3 = str.replace(fld3, item,'')
    path = case_dir+'\\'+"DFT"+'\\'+year+'\\'+dftonly+' - '+sus_op_crime+'\\'+exhib_ref+'\\'+fld3
    if not os.path.exists(path):      
        os.makedirs(path)
        '''folder3a = StringVar()
        folder3b = Button(startup, text=fld3,  width= 10, state=DISABLED, command=0).grid(row=5, column=6, sticky=EW)
        startup.update'''
###########################################################################################
def createfolder4():
    global fld1, fld2, fld3, fld4, fld5, fld6
    '''if len(e1a.get())!=7:
        e1b.config({"background": "orange red"})
        if len(e2a.get())<=1:
            e2b.config({"background": "orange red"})
            if len(e3a.get())<=1:
                e3b.config({"background": "orange red"})
    else:'''
    slashes = ('/', ':')
    for item in slashes:
        fld4 = str.replace(fld4, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        fld4 = str.replace(fld4, item,'')
    path = case_dir+'\\'+"DFT"+'\\'+year+'\\'+dftonly+' - '+sus_op_crime+'\\'+exhib_ref+'\\'+fld4
    if not os.path.exists(path):      
        os.makedirs(path)
        '''folder4a = StringVar()
        folder4b = Button(startup, text=fld4,  width= 10, state=DISABLED, command=0).grid(row=6, column=6, sticky=EW)
        startup.update'''
###########################################################################################
def createfolder5():
    global fld1, fld2, fld3, fld4, fld5, fld6
    '''if len(e1a.get())!=7:
        e1b.config({"background": "orange red"})
        if len(e2a.get())<=1:
            e2b.config({"background": "orange red"})
            if len(e3a.get())<=1:
                e3b.config({"background": "orange red"})
    else:'''
    slashes = ('/', ':')
    for item in slashes:
        fld5 = str.replace(fld5, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        fld5 = str.replace(fld5, item,'')
    path = case_dir+'\\'+"DFT"+'\\'+year+'\\'+dftonly+' - '+sus_op_crime+'\\'+exhib_ref+'\\'+fld5
    if not os.path.exists(path):      
        os.makedirs(path)
        '''folder5a = StringVar()
        folder5b = Button(startup, text=fld5,  width= 10, state=DISABLED, command=0).grid(row=7, column=6, sticky=EW)
        startup.update'''
############################################################################################
def createfolder6():
    global fld1, fld2, fld3, fld4, fld5, fld6
    '''if len(e1a.get())!=7:
        e1b.config({"background": "orange red"})
        if len(e2a.get())<=1:
            e2b.config({"background": "orange red"})
            if len(e3a.get())<=1:
                e3b.config({"background": "orange red"})
    else:'''
    slashes = ('/', ':')
    for item in slashes:
        fld6 = str.replace(fld6, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        fld6 = str.replace(fld6, item,'')
    path = case_dir+'\\'+"DFT"+'\\'+year+'\\'+dftonly+' - '+sus_op_crime+'\\'+exhib_ref+'\\'+fld6
    if not os.path.exists(path):      
        os.makedirs(path)
        '''folder6a = StringVar()
        folder6b = Button(startup, text=fld6,  width= 10, state=DISABLED, command=0).grid(row=8, column=6, sticky=EW)
        startup.update'''
############################################################################################
def changeexaminer():
    e9b.config(state='normal')
    Button(startup, text="  Update  ",  width= 10, command=updateexaminer).grid(row=12, column=4, sticky=W)
    startup.update()
############################################################################################
def updateexaminer(e9b):
    global analyst, examiner
#    e9b.config(state='disabled')
#    Button(startup, text="  Change  ",  width= 10, command=changeexaminer).grid(row=12, column=4, sticky=W)
    examiner=(e9a.get().upper())
    slashes = ('/', ':')
    for item in slashes:
        examiner = str.replace(examiner, item,'-')
    cleanup = ('\\', '|', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        examiner = str.replace(examiner, item,'')
    if not os.path.isfile(sqlprolocal):
        print("Examiner Profile not present, create?")
        createprofile()
    print(examiner)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET examiner = ? WHERE _rowid_ = 1', (examiner,))
    conn.commit()
    conn.close()
    analyst=examiner
    readprofile()
    readnotepad()
    if startup_method == 'Case':
        print('In updateexaminer()')
        overview()
    else:
        pass
    
    settings.update()
############################################################################################
def menu():
    menubar = Menu(root)

    casesmenu = Menu(root)
    casesmenu.add_command(label="Open Cases", command=0)
    casesmenu.add_command(label="Closed Cases", command=0)
    casesmenu.add_command(label="Exit", command=root.quit)
    menubar.add_cascade(label="Cases", menu=casesmenu)

    profilemenu = Menu(root)
    profilemenu.add_command(label="PIN Decryption log", command=pindecryptlog)
    menubar.add_cascade(label="Processes", menu=profilemenu)
    
    root.config(menu=menubar)
############################################################################################
def lockCase():
    global unlock
    e1b.config(state='disabled')
    e2b.config(state='disabled')
    e3b.config(state='disabled')
    e4b.config(state='disabled')
    e5b.config(state='disabled')
    e6b.config(state='disabled')
    e7b.config(state='disabled')
    e8.config(state='disabled')
    e24b.config(state='disabled')
    LoadButtonb.config(command=unlockCase)
    LoadButtonb.config(text="Unlock")
    LoadButtonb.config(state='normal')
    submitb.config(state='disabled')
    print('caselocked')
    new.update
############################################################################################
def unlockCase():
    
    e1b.config(state='normal')
    e2b.config(state='normal')
    e3b.config(state='normal')
    e4b.config(state='normal')
    e5b.config(state='normal')
    e6b.config(state='normal')
    e7b.config(state='normal')
    e8.config(state='normal')
    e24b.config(state='normal')
    LoadButtonb.config(command=loadcases_dft)
    LoadButtonb.config(text="Load")
    
    submitb.config(command=getData)
    submitb.config(state='normal')
##    if startup_method == 'Case':
##        startup.update
##    else:
##        new_case.update
############################################################################################
def disableloadcase():
    print("loadcasedisabled")
    submitb.config(state='disabled')
    startup.update
############################################################################################
def cleanup():
    cleanup = ('\\', '|', '/', ':', ';', '*', '?', '"', '<', '>')
    for item in cleanup:
        examiner = str.replace(examiner, item,'')
        dft_ref = str.replace(dft_ref, item,'')
        crime_ref = str.replace(crime_ref, item,'')
        exhib_ref = str.replace(exhib_ref, item,'')
############################################################################################
def restart_program():
    python = sys.executable
    os.execl(python, python, * sys.argv)
############################################################################################
def closesettings():
    updatefolders()
    settings_tab.destroy()
    sys.exit()
    overview()
############################################################################################
def closeprogram():
    root.destroy()
    sys.exit()
############################################################################################
def onclosingstartup():
    if messagebox.askokcancel("Quit", "Do you want to quit?"):
        root.destroy()
        sys.exit()
############################################################################################
##def copytoclip():
##    '''dft_ref=(str(e1a.get().upper()))
##    clippw=(str(dft_ref+''.join(pw)))
##    print("clip Pw: "+clippw)'''
##    clippw=startup_e1a.get()
##    print("clip Pw: "+clippw)
##    cmd='echo '+clippw.strip()+'|clip'
##    return subprocess.check_call(cmd, shell=True) 
############################################################################################
def copytoclip_individual_case():
    clippw=Indiv_Case8a.get()
    print("clip Pw: "+clippw)
    cmd='echo '+clippw.strip()+'|clip'
    return subprocess.check_call(cmd, shell=True) 
############################################################################################
def selectExamDir():
    d = argv[1] if len(argv)>1 else filedialog.askdirectory(initialdir="C:\\")
    D = os.path.realpath(d)
    e10.config(state='normal')
    e10.delete(0,END)
    e10.insert(10,D)
    e10.config(state='disabled')
    new.update()
    
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET case_dir = ? WHERE _rowid_ = 1', (D,))
    conn.commit()
    conn.close()
    
############################################################################################
def selectTempDir():
    settings.lower()
    d = argv[1] if len(argv)>1 else filedialog.askdirectory(initialdir="C:\\")
    D = os.path.realpath(d)
    e11.delete(0,END)
    e11.insert(10,D) # Log file name
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET template_dir = ? WHERE _rowid_ = 1', (D,))
    conn.commit()
    conn.close()
    settings.lift()
############################################################################################
def selectNetworkDB():
    settings.lower()
    d = argv[1] if len(argv)>1 else filedialog.askopenfilename(initialdir="C:\\")
    D = os.path.realpath(d)
    db10.delete(0,END)
    db10.insert(10,D) # Network db
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE databases SET network = ? WHERE _rowid_ = 1', (D,))
    conn.commit()
    conn.close()
    settings.lift()
############################################################################################
def selectLocalDB():
    settings.lower()
    d = argv[1] if len(argv)>1 else filedialog.askopenfilename(initialdir="C:\\")
    D = os.path.realpath(d)
    db11.delete(0,END)
    db11.insert(10,D) # Local db
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE databases SET local = ? WHERE _rowid_ = 1', (D,))
    conn.commit()
    conn.close()
    settings.lift()
############################################################################################
def selectContemp():
    global log_file
    settings.lower()
    templatedir = e11.get()
    if os.path.exists(templatedir):
        templatedir = e11.get()
    else:
        templatedir = "C:\\"
    currentselection = e12.get()
    f = argv[1] if len(argv)>1 else filedialog.askopenfilename(initialdir=templatedir)
    F = os.path.basename(f)
    if len(f)==0:
        f == currentselection
    else:
        from docx import Document
        try:
            Document(f)
            e12.delete(0,END)
            e12.insert(10,F) # Log file name
            log_file = F
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute('UPDATE profile SET contemp_file = ? WHERE _rowid_ = 1', (F,))
            conn.commit()
            conn.close()
        except ValueError:
            e12.delete(0,END)
            e12.insert(10,"Not a valid document type")
    settings.lift()
############################################################################################
def selectSfr():
    global sfr_file
    settings.lower()
    templatedir = e11.get()
    if os.path.exists(templatedir):
        templatedir = e11.get()
    else:
        templatedir = "C:\\"
    currentselection = e25.get()
    f = argv[1] if len(argv)>1 else filedialog.askopenfilename(initialdir=templatedir)
    F = os.path.basename(f)
    if len(f)==0:
        f == currentselection
    else:
        from docx import Document
        try:
            Document(f)
            e25.delete(0,END)
            e25.insert(10,F) # Log file name
            sfr_file = F
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute('UPDATE profile SET sfr_file = ? WHERE _rowid_ = 1', (F,))
            conn.commit()
            conn.close()
        except ValueError:
            e25.delete(0,END)
            e25.insert(10,"Not a valid document type")
    settings.lift()
############################################################################################
def selectDisclosure():
    global disclosure_file
    settings.lower()
    templatedir = e11.get()
    if os.path.exists(templatedir):
        templatedir = e11.get()
    else:
        templatedir = "C:\\"
    currentselection = e32.get()
    f = argv[1] if len(argv)>1 else filedialog.askopenfilename(initialdir=templatedir)
    F = os.path.basename(f)
    if len(f)==0:
        f == currentselection
    else:
        from docx import Document
        try:
            Document(f)
            e32.delete(0,END)
            e32.insert(10,F) # Disclosure file name
            sfr_file = F
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute('UPDATE profile SET disclosure_file = ? WHERE _rowid_ = 1', (F,))
            conn.commit()
            conn.close()
        except ValueError:
            e32.delete(0,END)
            e32.insert(10,"Not a valid document type")
    settings.lift()
############################################################################################
def updatepwlength():
    updated_pw_length = e21b.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET pw_length = ? WHERE _rowid_ = 1', (pw_length,))
    conn.commit()
    conn.close()
    if updated_pw_length == pw_length:
        print("no change")
    elif casetype == 'Loaded':
        print("Cant change pw, as this is a loaded case")
    else:
        genPw()
############################################################################################
def folderslock():

    folders_lock = Button(folders, text="  Unlock  ", width= 10, command=foldersunlock).grid(row=17, column=6, sticky=E)
    startup.update
############################################################################################
def foldersunlock():

    folders_lock = Button(folders, text="  Lock  ", width= 10, command=folderslock).grid(row=21, column=6, sticky=E)
    startup.update
############################################################################################
def settingslock():
    e13.config(state='disabled')
    e14.config(state='disabled')
    e15.config(state='disabled')
    e16.config(state='disabled')
#    e17.config(state='disabled')
#    e18.config(state='disabled')
    e26b.config(state='disabled')
    e27b.config(state='disabled')
    e28b.config(state='disabled')
    e29b.config(state='disabled')
#    e30b.config(state='disabled')
#    e31b.config(state='disabled')
    e10.config(state='disabled')
    e11.config(state='disabled')
    e12.config(state='disabled')
#    e21b.config(state='disabled')
#    e22b.config(state='disabled')
    e25.config(state='disabled')
    e32.config(state='disabled')
    e33b.config(state='disabled')
    e34b.config(state='disabled')
    e35b.config(state='disabled')
    e36b.config(state='disabled')
    e37b.config(state='disabled')
    e38b.config(state='disabled')

    casedir.config(state='disabled')
    tempdir.config(state='disabled')
    contempdir.config(state='disabled')
    sfrdir.config(state='disabled')
    disclosuredir.config(state='disabled')

    settings_lock.config(text="  Unlock  ")
    settings_lock.config(command=settingsunlock)
    e9b.config(state='disabled')
    estartupb.config(state='disabled')
#    change_analyst_button.config(state='disabled')
    db10.config(state='disabled')
    db11.config(state='disabled')

    networkdb.config(state='disabled')
    localdb.config(state='disabled')
    settings_tab.update
############################################################################################
def settingsunlock():
    e13.config(state='normal')
    e14.config(state='normal')
    e15.config(state='normal')
    e16.config(state='normal')
#    e17.config(state='normal')
#    e18.config(state='normal')
    e26b.config(state='normal')
    e27b.config(state='normal')
    e28b.config(state='normal')
    e29b.config(state='normal')
#    e30b.config(state='normal')
#    e31b.config(state='normal')
    e10.config(state='normal')
    e11.config(state='normal')
    e12.config(state='normal')
#    e21b.config(state='normal')
#    e22b.config(state='normal')
    e25.config(state='normal')
    e32.config(state='normal')
    e33b.config(state='normal')
#    e34b.config(state='normal')
    e35b.config(state='normal')
#    e36b.config(state='normal')
    e37b.config(state='normal')
    e38b.config(state='disabled')

    casedir.config(state='normal')
    tempdir.config(state='normal')
    contempdir.config(state='normal')
    sfrdir.config(state='normal')
    disclosuredir.config(state='normal')

    settings_lock.config(text="  Lock  ")
    settings_lock.config(command=settingslock)
    e9b.config(state='normal')
    estartupb.config(state='normal')
#    change_analyst_button.config(state='normal')
    db10.config(state='disabled')
    db11.config(state='disabled')

    networkdb.config(state='normal')
    localdb.config(state='normal')
    settings_tab.update
############################################################################################
def ontabchangelock():
    pass
#    e10.config(state='disabled')
#    e11.config(state='disabled')
#    e12.config(state='disabled')
#    e13.config(state='disabled')
#    e14.config(state='disabled')
#    e15.config(state='disabled')
#    e16.config(state='disabled')
#    e17.config(state='disabled')
#    e18.config(state='disabled')
#    e21b.config(state='disabled')
#    e22b.config(state='disabled')
#    e25.config(state='disabled')
#    e32.config(state='disabled')

#    e26b.config(state='disabled')
#    e27b.config(state='disabled')
#    e28b.config(state='disabled')
#    e29b.config(state='disabled')
#    e30b.config(state='disabled')
#    e31b.config(state='disabled')
#    e33b.config(state='disabled')
#    e34b.config(state='disabled')
#    e35b.config(state='disabled')
#    e36b.config(state='disabled')
#    e37b.config(state='disabled')
#    e38b.config(state='disabled')
    
#    updatefolders()

#    overview()
#    change_to_notepad()
#    settings_lock = Button(settings, text=" Unlock ", width= 10, command=settingsunlock).grid(row=21, column=6, sticky=E)

#    updatepwlength()
#    startup.update
############################################################################################
def change_to_notepad():
    get=notepad_overview.get(1.0, END)
    notepad_n59b.delete(1.0,END)
    notepad_n59b.insert(INSERT, get)
############################################################################################
def readpindecryptlog():
    global pdlog, no_of_entries, pindecryptlogs, entry
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    pindecryptlogs=[]
    c.execute(""" select * from pindecryptlog WHERE analyst like '%'||?||'%'""", (analyst,))
    count=1
    for each in c:
        pindecryptlogs.append(each)
        
        count+=1

    conn.close()

    pindecryptcount = (len(pindecryptlogs))
    print("No of logs: ", pindecryptcount)

    no_of_entries=(len(pindecryptlogs))+1
############################################################################################
def addline():
    global row, addentryb, statusbuttonb, Entry1a, Entry2a, Entry3a, Entry4a, new0, new1, new2, new3, new4, new5
    count=1
    new1 = (Newentry1a.get().upper())
    new2 = (Newentry2a.get().upper())
    new3 = (Newentry3a.get().upper())
    new4 = (Newentry4a.get().upper())
    new5 = (Newentry5a.get().upper())
    
#    populatepinlog()
#    viewTypea.set('PIN decryption log')
    pindecryptcount = (len(pindecryptlogs))

    row=pindecryptcount+1

    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" insert into pindecryptlog values (?, ?, ?, ?, ?, ?, ?, ?, '-');""", (row, new1, new2, new3, new4, datestamp,'0', analyst,))
    conn.commit()
    conn.close()
    add.destroy()
    t4 = threading.Thread(target=populatepinlog)
    t4.start()
#    overview()
    root.update()

############################################################################################
def overview():
    global left1, right1, row, notepad_overview, Entry1a, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, entry1, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, opencasevariables, notepad_n59b_overview, photoScrollv, photoScrollh, photoCanvas1, photoCanvas2
    readpindecryptlog()
    readnotepad()
    readopencases()
    for each in pindecryptvariables:
        each.destroy()
    for each in opencasevariables:
        each.destroy()
    for each in overviewvariables:
        each.destroy()
    for each in closedcasevariables:
        each.destroy()
    #viewTypea.set('Overview')

    ############################################################################################
    '''
    left_frame = Frame(open_cases, borderwidth=1,relief=RIDGE)
    left_frame.grid(row=0, column=1, columnspan=10, sticky=E) 

    photoCanvas1 = Canvas(left_frame, width=1200, height=(height)-45) #width=width, height=(height)-45)
    photoCanvas1.grid()#sticky=NSEW)

    left1 = Frame(photoCanvas1, width=width, height=10000)
    photoCanvas1.create_window(0, 0, window=left1, anchor='nw')
    ############################################################################################
    right_frame = Frame(open_cases, borderwidth=1,relief=RIDGE)
    right_frame.grid(row=0, column=11, columnspan=10, sticky=W)

    photoCanvas2 = Canvas(right_frame, width=600, height=(height)-45) #width=width, height=(height)-45)
    photoCanvas2.grid()#sticky=NSEW)

    right1 = Frame(photoCanvas2, width=width, height=10000)
    photoCanvas2.create_window(0, 0, window=right1, anchor='nw')'''

    centre_frame = Frame(open_cases, borderwidth=1,relief=RIDGE)
    centre_frame.grid(row=0, column=11, columnspan=10, sticky=W)

    photoCanvas3 = Canvas(centre_frame, width=600, height=(height)-45) #width=width, height=(height)-45)
    photoCanvas3.grid()#sticky=NSEW)

    centre1 = Frame(photoCanvas3, width=width, height=10000)
    photoCanvas3.create_window(0, 0, window=centre1, anchor='center')
    #anchor must be n, ne, e, se, s, sw, w, nw, or center
    ############################################################################################
    
    ### Notepad tab (scroll WORKING)
    notepad_overview=Text(left1, wrap=WORD, width=122, height=8) #, yscrollcommand=notepad_scrollbar.set)
    notepad_overview.grid(row=1, column=1, columnspan=15, rowspan=4, sticky=W)
    print('Len notepadentries:',(len(notepadentries)))
    if (len(notepadentries))==2:
            print('in func')
            notepadnotes1_text = (str(notepadentries[1]))
            notepadnotes1_text = (str.replace(notepadnotes1_text,"\\n","\n"))
            notepad_overview.insert(INSERT, notepadnotes1_text)
    else:
        pass
    notepad_overview.grid_propagate(True)
    notepad_overview.bind('<KeyRelease>', notepad_overview_notes1)
    notepad_overview.bind('<Double-Button-1>', lambda x: tabs.select(notepad))
    notepad_overview.config(state='normal')
    overviewvariables.append(notepad_overview)
    
    row = 1
    column = 1
    Label_pin_d = Label(right1,  text='PIN Decryption(s)')
    Label_pin_d.grid(row=row, column=column, sticky=W, padx=1)
    overviewvariables.append(Label_pin_d)
    
    row+=1
    
    count=1
    Label0 = Label(right1, text=' ', width=200,  borderwidth=3, relief="groove")
    Label0.grid(row=row, column=column, columnspan=25, sticky=W, ipady=3, ipadx=20)
    overviewvariables.append(Label0)

    v0 =  StringVar()
    Label0 = Label(right1, textvariable=v0)
    Label0.grid(row=row, column=column, sticky=W, padx=1)
    v0.set('DFT  Ref')
    overviewvariables.append(Label0)
    column+=1        
    v1 =  StringVar()
    Label1 = Label(right1, textvariable=v1)
    Label1.grid(row=row, column=column, sticky=W)
    v1.set('Crime Ref')
    overviewvariables.append(Label1)
    column+=1
    v2 =  StringVar()
    Label2 = Label(right1, textvariable=v2)
    Label2.grid(row=row, column=column, sticky=W)
    v2.set('Exhibit Ref')
    overviewvariables.append(Label2)
    column+=1
    v3 =  StringVar()
    Label3 = Label(right1, textvariable=v3)
    Label3.grid(row=row, column=column, sticky=W)
    v3.set('OIC')
    overviewvariables.append(Label3)
    column+=1
    v4 =  StringVar()
    Label4 = Label(right1, textvariable=v4)
    Label4.grid(row=row, column=column, sticky=W)
    v4.set('Date started')
    overviewvariables.append(Label4)
    column+=1
    v5 =  StringVar()
    Label5 = Label(right1, textvariable=v5)
    Label5.grid(row=row, column=column, sticky=W)
    v5.set('Day running')
    overviewvariables.append(Label5)
    column+=1
    v6 =  StringVar()
    Label6 = Label(right1, textvariable=v6)
    Label6.grid(row=row, column=column, sticky=W)
    v6.set('Notes')
    overviewvariables.append(Label6)
    column+=1
    v7 =  StringVar()
    Label7 = Label(right1, textvariable=v7)
    Label7.grid(row=row, column=column, sticky=W)
    v7.set('')
    overviewvariables.append(Label7)
    column+=1    
    v8 =  StringVar()
    Label8 = Label(right1, textvariable=v8)
    Label8.grid(row=row, column=column, sticky=W)
    v8.set('')
    overviewvariables.append(Label8)
    column+=1
    v9 =  StringVar()
    Label9 = Label(right1, textvariable=v9)
    Label9.grid(row=row, column=column, sticky=W)
    v9.set('')
    overviewvariables.append(Label9)
    column+=1    
    '''edit1b.grid(row=0, column=15, columnspan=2, sticky=W)
    edit1b.config(state='normal')'''

#    launch.grid(row=0, column=13, columnspan=1, sticky=W)
    column = 1
    row=3
    for each in pindecryptlogs:

        entry1 = (str("e"+(str(count))))
        Entry1a = StringVar()
        Entry1a.set(each[1])
        dft=each[1]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(right1, text=dft, width= 10, state='normal', textvariable=Entry1a)
        entry1.bind('<Button-1>', populate_individual_pd)
        entry1.grid(row=row, column=column, sticky=EW)
        column+=1
        
        '''entry1 = (str("e"+(str(count))))
        Entry1a = StringVar()
        Entry1a.set(each[1])
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Entry(right1, relief=SUNKEN,width=8, state='disabled', textvariable=Entry1a)
        entry1.grid(row=row, column=11, sticky=EW)'''
        pindecryptvariables.append(entry1)
        count += 1
        entry2 = (str("e"+(str(count))))
        Entry2a = StringVar()
        Entry2a.set(each[2])
        Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
        entry2 = Entry(right1, relief=SUNKEN,width=15, state='disabled', textvariable=Entry2a)
        entry2.grid(row=row, column=column, sticky=EW, ipady=3, padx=2)
        pindecryptvariables.append(entry2)
        count += 1
        column+=1
        
        entry3 = (str("e"+(str(count))))
        Entry3a = StringVar()
        Entry3a.set(each[3])
        Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
        entry3 = Entry(right1, relief=SUNKEN, width=12, state='disabled', textvariable=Entry3a)
        entry3.grid(row=row, column=column, sticky=EW, ipady=3)
        pindecryptvariables.append(entry3)
        count += 1
        column+=1
        
        entry4 = (str("e"+(str(count))))
        Entry4a = StringVar()
        Entry4a.set(each[4])
        Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
        entry4 = Entry(right1, relief=SUNKEN,width=30, state='disabled', textvariable=Entry4a)
        entry4.grid(row=row, column=column, sticky=W, ipady=3)
        pindecryptvariables.append(entry4)
        count += 1
        column+=1
        
        entry5 = (str("e"+(str(count))))
        Entry5a = StringVar()
        Entry5a.set(each[5])
        Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
        entry5 = Entry(right1, relief=SUNKEN,width=13, state='disabled', textvariable=Entry5a)
        entry5.grid(row=row, column=column, columnspan=2, sticky=W, ipady=3)
        pindecryptvariables.append(entry5)
        count += 1
        column+=1
        
        selected_month_rec = (each[5])
        start = date(int(selected_month_rec.split('-')[0]),int(selected_month_rec.split('-')[1]),int(selected_month_rec.split('-')[2]))
        today = date.today()
        res = today - start
        res.days
        entry6 = (str("e"+(str(count))))
        Entry6a = StringVar()
        Entry6a.set(res.days)
        Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
        entry6 = Entry(right1, relief=SUNKEN,width=12, state='disabled', textvariable=Entry6a)
        entry6.grid(row=row, column=column, sticky=W, ipady=3)
        pindecryptvariables.append(entry6)
        count += 1
        column+=1
        
        entry7 = (str("e"+(str(count))))
        Entry7a = StringVar()
        Entry7a.set(each[8])
        Entry7a.trace("w", lambda name, index, mode, Entry7a=Entry7a: callback(Entry7a))
        entry7 = Entry(right1, relief=SUNKEN,width=155, state='disabled', textvariable=Entry7a)
        entry7.grid(row=row, column=column, sticky=W, columnspan=50, ipady=3)
        pindecryptvariables.append(entry7)
        row += 1
        column = 1

    
#    edit1b.config(command=disable_edit2)
    readopencases()
    column=1
    count=1
    row=6

    Label21 = Label(left1, text=' ')
    Label21.grid(row=row, column=column, columnspan=1)
    overviewvariables.append(Label21)
    
    row += 1
    Label20 = Label(left1, textvariable=' ', width=140,  borderwidth=3, relief="groove")
    Label20.grid(row=row, column=1, columnspan=15, sticky=W, ipady=3)
    overviewvariables.append(Label20)
    
    v0 =  StringVar()
    Label0 = Label(left1, textvariable=v0)
    Label0.grid(row=row, column=column, sticky=W, padx=1)
    v0.set('DFT  Ref')
    overviewvariables.append(Label0)
    column+=1        
    v1 =  StringVar()
    Label1 = Label(left1, textvariable=v1)
    Label1.grid(row=row, column=column, sticky=W)
    v1.set('Crime Ref')
    overviewvariables.append(Label1)
    column+=1 
    v2 =  StringVar()
    Label2 = Label(left1, textvariable=v2)
    Label2.grid(row=row, column=column, sticky=W)
    v2.set('Suspect')
    overviewvariables.append(Label2)
    column+=1 
    v3 =  StringVar()
    Label3 = Label(left1, textvariable=v3)
    Label3.grid(row=row, column=column, sticky=W)
    v3.set('OIC')
    overviewvariables.append(Label3)
    column+=1 
    v4 =  StringVar()
    Label4 = Label(left1, textvariable=v4)
    Label4.grid(row=row, column=column, sticky=W)
    v4.set('Date started')
    overviewvariables.append(Label4)
    column+=1 
    v5 =  StringVar()
    Label5 = Label(left1, textvariable=v5)
    Label5.grid(row=row, column=column, sticky=W)
    v5.set('# of exhibits')
    overviewvariables.append(Label5)
    column+=1 
    v6 =  StringVar()
    Label6 = Label(left1, textvariable=v6)
    Label6.grid(row=row, column=column, sticky=EW)
    '''v6.set('Pre-Img')
    overviewvariables.append(Label6)
    column+=1 
    v7 =  StringVar()
    Label7 = Label(left1, textvariable=v7)
    Label7.grid(row=row, column=column, sticky=EW)
    v7.set('Image')
    overviewvariables.append(Label7)'''
    column+=1
    v8 =  StringVar()
    Label8 = Label(left1, textvariable=v8)
    Label8.grid(row=row, column=column, sticky=EW)
    v8.set('QC')
    overviewvariables.append(Label8)
    column+=1 
    v9 =  StringVar()
    Label9 = Label(left1, textvariable=v9)
    Label9.grid(row=row, column=column, sticky=EW)
    v9.set('Server')
    overviewvariables.append(Label9)


    '''edit1b.grid(row=0, column=15, columnspan=1, sticky=W)
    edit1b.grid_forget()'''
    

#    launch.grid(row=0, column=13, columnspan=1, sticky=W)


    row += 1
    column = 1
    for each in opencases:
        Entry1a = StringVar()
        Entry1a.set(each[1])
        dft=each[1]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(left1, text=dft, width= 10, textvariable=Entry1a)
        entry1.bind('<Button-1>', populate_individual_case)
        entry1.grid(row=row, column=column, sticky=EW)

        column+=1
        print('opencases: ', each[1])

        opencasevariables.append(entry1)
        count += 1
        entry2 = (str("e"+(str(count))))
        Entry2a = StringVar()
        Entry2a.set(each[2])
        Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
        entry2 = Entry(left1, relief=SUNKEN,width=15, state='disabled', textvariable=Entry2a)
        entry2.grid(row=row, column=column, sticky=EW, ipady=3, padx=2)
        opencasevariables.append(entry2)
        count += 1
        column+=1
        
        entry3 = (str("e"+(str(count))))
        Entry3a = StringVar()
        Entry3a.set(each[8])
        Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
        entry3 = Entry(left1, relief=SUNKEN,width=25, state='disabled', textvariable=Entry3a)
        entry3.grid(row=row, column=column, sticky=EW, ipady=3)
        opencasevariables.append(entry3)
        count += 1
        column+=1
        
        entry4 = (str("e"+(str(count))))
        Entry4a = StringVar()
        Entry4a.set(each[4])
        Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
        entry4 = Entry(left1, relief=SUNKEN,width=30, state='disabled', textvariable=Entry4a)
        entry4.grid(row=row, column=column, sticky=W, ipady=3)
        opencasevariables.append(entry4)
        count += 1
        column+=1
        
        entry5 = (str("e"+(str(count))))
        Entry5a = StringVar()
        Entry5a.set(each[5])
        Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
        entry5 = Entry(left1, relief=SUNKEN,width=12, state='disabled', textvariable=Entry5a)
        entry5.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        opencasevariables.append(entry5)
        count += 1
        column+=1

        entry7 = (str("e"+(str(count))))
        Entry7a = StringVar()
        Entry7a.set(each[3])
        Entry7a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry7a))
        entry7 = Entry(left1, relief=SUNKEN,width=12, state='disabled', textvariable=Entry7a)
        entry7.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        opencasevariables.append(entry7)
        count += 1
        column+=2
        
        '''e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        overviewvariables.append(e30b)
        column+=2
        overviewvariables.append(e30b)
        
        e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        column+=1
        overviewvariables.append(e30b)'''
        
        e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        column+=1
        opencasevariables.append(e30b)
        
        e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        opencasevariables.append(e30b)
        column=1
        
        # NOTES
        row+=1
        entry6 = (str("e"+(str(count))))
        Entry6a = StringVar()
        Entry6a.set(each[6])
        Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
        entry6 = Entry(left1, relief=SUNKEN, width=112,  state='disabled', textvariable=Entry6a)
        entry6.grid(row=row, column=2, columnspan=14, sticky=EW, ipady=0)
        opencasevariables.append(entry6)
        count += 1
        row += 1

        #label=Label(left1, text=" - ").grid(row=row, column=2, columnspan=1, sticky=EW)
        
        #Label0 = Label(left1, text=' ', width=10)
        #Label0.grid(row=row, column=1, columnspan=1, sticky=EW, ipady=3, ipadx=20)
        #overviewvariables.append(Label0)

        #row += 2
    photoScrollv = Scrollbar(left_frame, orient=VERTICAL)
    photoScrollv.config(command=photoCanvas1.yview)
    photoScrollh = Scrollbar(left_frame, orient=HORIZONTAL)
    photoScrollh.config(command=photoCanvas1.xview)
    photoCanvas1.config(yscrollcommand=photoScrollv.set)
    photoCanvas1.config(xscrollcommand=photoScrollh.set)
    photoScrollv.grid(row=0, column=1, sticky="ns")
    photoScrollh.grid(row=7, column=0, sticky="ew")
    left1.bind("<Configure>", update_scrollregion)
    photoScrollv.bind("<MouseWheel>", update_scrollregion)

    photoScrollv = Scrollbar(right_frame, orient=VERTICAL)
    photoScrollv.config(command=photoCanvas2.yview)
    photoScrollh = Scrollbar(right_frame, orient=HORIZONTAL)
    photoScrollh.config(command=photoCanvas2.xview)
    photoCanvas2.config(yscrollcommand=photoScrollv.set)
    photoCanvas2.config(xscrollcommand=photoScrollh.set)
    photoScrollv.grid(row=0, column=1, sticky="ns")
    photoScrollh.grid(row=7, column=0, sticky="ew")
    right1.bind("<Configure>", update_scrollregion)
    photoScrollv.bind("<MouseWheel>", update_scrollregion)


    root.update()
############################################################################################
def populatepinlog():
    global row, Entry1a, Entry2a, Entry3a, Entry4a, Entry5a, Entry6b

    readpindecryptlog()
    count=1
    
    for each in pindecryptvariables:
        each.destroy()
    '''for each in assignedcasevariables:
        each.destroy()
    for each in individualcasevariables:
        each.destroy()
    for each in overviewvariables:
        each.destroy()
    for each in closedcasevariables:
        each.destroy()'''
    column = 1
    row = 3
    for each in pindecryptlogs:

        entry1 = (str("e"+(str(count))))
        Entry1a = StringVar()
        Entry1a.set(each[1])
        dft=each[1]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(right1, text=dft, width= 15, state='normal', textvariable=Entry1a)
        entry1.bind('<Button-1>', populate_individual_pd)
        entry1.grid(row=row, column=column, sticky=EW)
        pindecryptvariables.append(entry1)

        column+=1
        count += 1
        entry2 = (str("e"+(str(count))))
        Entry2a = StringVar()
        Entry2a.set(each[2])
        Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
        entry2 = Entry(right1, relief=SUNKEN,width=12, state='disabled', textvariable=Entry2a)
        entry2.grid(row=row, column=column, sticky=EW, ipady=3, padx=2)
        pindecryptvariables.append(entry2)

        column+=1
        count += 1
        entry3 = (str("e"+(str(count))))
        Entry3a = StringVar()
        Entry3a.set(each[3])
        Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
        entry3 = Entry(right1, relief=SUNKEN, state='disabled', textvariable=Entry3a)
        entry3.grid(row=row, column=column, sticky=EW, ipady=3)
        pindecryptvariables.append(entry3)

        column+=1
        count += 1
        entry4 = (str("e"+(str(count))))
        Entry4a = StringVar()
        Entry4a.set(each[4])
        Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
        entry4 = Entry(right1, relief=SUNKEN,width=35, state='disabled', textvariable=Entry4a)
        entry4.grid(row=row, column=column, sticky=W, ipady=3)
        pindecryptvariables.append(entry4)

        column+=1
        count += 1
        entry5 = (str("e"+(str(count))))
        Entry5a = StringVar()
        Entry5a.set(each[5])
        Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
        entry5 = Entry(right1, relief=SUNKEN,width=15, state='disabled', textvariable=Entry5a)
        entry5.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        pindecryptvariables.append(entry5)

        column+=1
        count += 1
        selected_month_rec = (each[5])
        start = date(int(selected_month_rec.split('-')[0]),int(selected_month_rec.split('-')[1]),int(selected_month_rec.split('-')[2]))
        today = date.today()
        res = today - start
        res.days
        entry6 = (str("e"+(str(count))))
        Entry6a = StringVar()
        Entry6a.set(res.days)
        Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
        entry6 = Entry(right1, relief=SUNKEN, width=12, state='disabled', textvariable=Entry6a)
        entry6.grid(row=row, column=column, sticky=W, ipady=3)
        pindecryptvariables.append(entry6)

        column+=1
        count += 1
        entry7 = (str("e"+(str(count))))
        Entry7a = StringVar()
        Entry7a.set(each[8])
        Entry7a.trace("w", lambda name, index, mode, Entry7a=Entry7a: callback(Entry7a))
        entry7 = Entry(right1, relief=SUNKEN,width=120, state='disabled', textvariable=Entry7a)
        entry7.grid(row=row, column=column, sticky=W, columnspan=2, ipady=3)
        pindecryptvariables.append(entry7)
        row += 1
        column = 1
    right1.update
############################################################################################
def populateopencases():
    pass
############################################################################################
def Xpopulateopencases():
    global row, Entry1a, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, entry1, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, opencasevariables
#    edit1b.config(command=disable_edit2)
    readopencases()
    print('writing open cases: ')
    print(len(opencases))
    count=1
    
    for each in opencasevariables:
        each.destroy()
        
    row=8
    column = 1
    for each in opencases:
        Entry1a = StringVar()
        Entry1a.set(each[1])
        dft=each[1]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(left1, text=dft, width= 10, textvariable=Entry1a)
        entry1.bind('<Button-1>', populate_individual_case)
        entry1.grid(row=row, column=column, sticky=EW)

        column+=1
        print('opencases: ', each[1])

        opencasevariables.append(entry1)
        count += 1
        entry2 = (str("e"+(str(count))))
        Entry2a = StringVar()
        Entry2a.set(each[2])
        Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
        entry2 = Entry(left1, relief=SUNKEN,width=15, state='disabled', textvariable=Entry2a)
        entry2.grid(row=row, column=column, sticky=EW, ipady=3, padx=2)
        opencasevariables.append(entry2)
        count += 1
        column+=1
        
        entry3 = (str("e"+(str(count))))
        Entry3a = StringVar()
        Entry3a.set(each[8])
        Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
        entry3 = Entry(left1, relief=SUNKEN,width=25, state='disabled', textvariable=Entry3a)
        entry3.grid(row=row, column=column, sticky=EW, ipady=3)
        opencasevariables.append(entry3)
        count += 1
        column+=1
        
        entry4 = (str("e"+(str(count))))
        Entry4a = StringVar()
        Entry4a.set(each[4])
        Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
        entry4 = Entry(left1, relief=SUNKEN,width=30, state='disabled', textvariable=Entry4a)
        entry4.grid(row=row, column=column, sticky=W, ipady=3)
        opencasevariables.append(entry4)
        count += 1
        column+=1
        
        entry5 = (str("e"+(str(count))))
        Entry5a = StringVar()
        Entry5a.set(each[5])
        Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
        entry5 = Entry(left1, relief=SUNKEN,width=12, state='disabled', textvariable=Entry5a)
        entry5.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        opencasevariables.append(entry5)
        count += 1
        column+=1

        entry7 = (str("e"+(str(count))))
        Entry7a = StringVar()
        Entry7a.set(each[3])
        Entry7a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry7a))
        entry7 = Entry(left1, relief=SUNKEN,width=12, state='disabled', textvariable=Entry7a)
        entry7.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        opencasevariables.append(entry7)
        count += 1
        column+=2
        
        '''e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        overviewvariables.append(e30b)
        column+=2
        overviewvariables.append(e30b)
        
        e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        column+=1
        overviewvariables.append(e30b)'''
        
        e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        column+=1
        opencasevariables.append(e30b)
        
        e30a = IntVar()
        e30a.set(0)
        e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
        e30b = Checkbutton(left1, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
        e30b.grid(row=row, column=column, sticky=EW)
        opencasevariables.append(e30b)
        column=1
        
        # NOTES
        row+=1
        entry6 = (str("e"+(str(count))))
        Entry6a = StringVar()
        Entry6a.set(each[6])
        Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
        entry6 = Entry(left1, relief=SUNKEN,width=112, state='disabled', textvariable=Entry6a)
        entry6.grid(row=row, column=2, columnspan=14, sticky=EW, ipady=0)
        opencasevariables.append(entry6)
        count += 1
        row += 1
############################################################################################
def populateclosedcases():
    global row, Entry1a, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, entry1, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, opencasevariables
#    edit1b.config(command=disable_edit2)
    readclosedcases()
        
    for each in pindecryptvariables:
        each.destroy()
    for each in opencasevariables:
        each.destroy()
    for each in individualcasevariables:
        each.destroy()
    for each in overviewvariables:
        each.destroy()
        
    row=2
    Label20 = Label(startup, textvariable='DFT Ref', width=165,  borderwidth=3, relief="groove")
    Label20.grid(row=row, column=1, columnspan=50, sticky=EW, ipady=3)
    closedcasevariables.append(Label20)
    print('writing open cases: ')
    print(len(opencases))
    count=1
    v0 =  StringVar()
    Label0 = Label(startup, textvariable=v0)
    Label0.grid(row=row, column=2, sticky=W)
    v0.set('DFT  Ref')
    closedcasevariables.append(Label0)
        
    v1 =  StringVar()
    Label1 = Label(startup, textvariable=v1)
    Label1.grid(row=row, column=3, sticky=W)
    v1.set('Crime Ref')
    closedcasevariables.append(Label1)

    v2 =  StringVar()
    Label2 = Label(startup, textvariable=v2)
    Label2.grid(row=row, column=4, sticky=W)
    v2.set('Exhibit Ref')
    closedcasevariables.append(Label2)

    v3 =  StringVar()
    Label3 = Label(startup, textvariable=v3)
    Label3.grid(row=row, column=5, sticky=W)
    v3.set('Suspect')
    closedcasevariables.append(Label3)

    v4 =  StringVar()
    Label4 = Label(startup, textvariable=v4)
    Label4.grid(row=row, column=6, sticky=W)
    v4.set('OIC')
    closedcasevariables.append(Label4)

    v5 =  StringVar()
    Label5 = Label(startup, textvariable=v5)
    Label5.grid(row=row, column=7, sticky=W)
    v5.set('Case password')
    closedcasevariables.append(Label5)

    v6 =  StringVar()
    Label6 = Label(startup, textvariable=v6)
    Label6.grid(row=row, column=8, sticky=W)
    v6.set('Date Closed')
    closedcasevariables.append(Label6)

    v7 =  StringVar()
    Label7 = Label(startup, textvariable=v7)
    Label7.grid(row=row, column=11, sticky=W)
    v7.set('')
    closedcasevariables.append(Label7)
    
    v8 =  StringVar()
    Label8 = Label(startup, textvariable=v8)
    Label8.grid(row=row, column=12, sticky=W)
    v8.set('')
    closedcasevariables.append(Label8)

    v9 =  StringVar()
    Label9 = Label(startup, textvariable=v9)
    Label9.grid(row=row, column=13, sticky=W)
    v9.set('')
    closedcasevariables.append(Label9)

    edit1b.grid(row=0, column=15, columnspan=1, sticky=W)
    edit1b.grid_forget()

#    launch.grid(row=0, column=13, columnspan=1, sticky=W)

    row += 1
    column = 2
    for each in closedcases:
        # DFT full
        Entry1a = StringVar()
        Entry1a.set(each[11])
        dft=each[11]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(startup, text=dft, width= 15,  state='disabled', textvariable=Entry1a)
        #entry1.bind('<Button-1>', populate_individual_case)
        entry1.grid(row=row, column=column, sticky=EW)

        column+=1
        print('closed cases: ', each[11])
        # Crime ref
        closedcasevariables.append(entry1)
        count += 1
        entry2 = (str("e"+(str(count))))
        Entry2a = StringVar()
        Entry2a.set(each[3])
        Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
        entry2 = Entry(startup, relief=SUNKEN,width=15, state='disabled', textvariable=Entry2a)
        entry2.grid(row=row, column=column, sticky=EW, ipady=3, padx=2)
        closedcasevariables.append(entry2)
        count += 1
        column+=1
        # Exhibit
        entry3 = (str("e"+(str(count))))
        Entry3a = StringVar()
        Entry3a.set(each[4])
        Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
        entry3 = Entry(startup, relief=SUNKEN,width=25, state='disabled', textvariable=Entry3a)
        entry3.grid(row=row, column=column, sticky=W, ipady=3)
        closedcasevariables.append(entry3)
        count += 1
        column+=1
        # Suspect
        entry4 = (str("e"+(str(count))))
        Entry4a = StringVar()
        Entry4a.set(each[8])
        Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
        entry4 = Entry(startup, relief=SUNKEN,width=35, state='disabled', textvariable=Entry4a)
        entry4.grid(row=row, column=column, sticky=W, ipady=3)
        closedcasevariables.append(entry4)
        count += 1
        column+=1
        # OIC
        entry5 = (str("e"+(str(count))))
        Entry5a = StringVar()
        Entry5a.set(each[7])
        Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
        entry5 = Entry(startup, relief=SUNKEN,width=35, state='disabled', textvariable=Entry5a)
        entry5.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        closedcasevariables.append(entry5)
        count += 1
        column+=1
        # Case password
        entry6 = (str("e"+(str(count))))
        Entry6a = StringVar()
        Entry6a.set(each[12])
        Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
        entry6 = Entry(startup, relief=SUNKEN,width=25, state='disabled', textvariable=Entry6a)
        entry6.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        closedcasevariables.append(entry6)
        count += 1
        column+=1
        # Date closed
        entry7 = (str("e"+(str(count))))
        Entry7a = StringVar()
        Entry7a.set(each[13])
        Entry7a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry7a))
        entry7 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry7a)
        entry7.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        closedcasevariables.append(entry7)
        

        count += 1
        row += 1
        column = 2

    photoScrollv = Scrollbar(photoFrame1, orient=VERTICAL)
    photoScrollv.config(command=photoCanvas1.yview)
    photoScrollh = Scrollbar(photoFrame1, orient=HORIZONTAL)
    photoScrollh.config(command=photoCanvas1.xview)
    photoCanvas1.config(yscrollcommand=photoScrollv.set)
    photoCanvas1.config(xscrollcommand=photoScrollh.set)
    photoScrollv.grid(row=0, column=1, sticky="ns")
    photoScrollh.grid(row=row, column=0, sticky="ew")
    startup.bind("<Configure>", update_scrollregion)
    root.update
############################################################################################
def populateassignedcases():
    global row, Entry1a, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, opencasevariables, startAssigneda
    edit1b.config(command=disable_edit3)
    readassignedcases()
    count=1
    v0.set('')
    v1.set('DFT Ref')
    v2.set('Crime Ref')
    v3.set('# of exhibits')
    v4.set('OIC')
    v5.set('Date assigned')
    v6.set('Notes')
    v7.set('')

    edit1b.grid_forget()#(row=0, column=15, columnspan=1, sticky=W)
    
    launch.grid(row=0, column=13, columnspan=1, sticky=W)
        
    startAssigneda = StringVar()
    startAssigneda.set('Start case')
    startAssigneda.trace("w", lambda name, index, mode, startAssigneda=startAssigneda: callback(startAssigneda))
    startAssignedb = OptionMenu(startup, startAssigneda,*assigned_cases_options, command=start_assigned_from_dropdown)
    startAssignedb.grid(row=0, column=16, sticky=W)
    assignedcasevariables.append(startAssignedb)
    
    for each in pindecryptvariables:
        each.destroy()
    for each in opencasevariables:
        each.destroy()
    for each in individualcasevariables:
        each.destroy()


    
    row = 2
    for each in assignedcases:
        entry = (str("e"+(str(row))))
        entry1 = (str("e"+(str(row))))
        entry1a = (str("e"+(str(row))))
        entry1a = StringVar()
        entry1a.set(each[1])
        
        '''a=each[1]
        dft=a
        entry = Deleteopencase(dft)
        entry.dft   =   a
        entry1 = Entry(startup, relief=SUNKEN,width=10, state='disabled', textvariable=entry1a)
        entry1.grid(row=row, column=11, sticky=EW)'''

        Entry1a = StringVar()
        Entry1a.set(each[1])
        dft=each[1]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(startup, text=dft, width= 10, state='disabled', textvariable=entry1a)
        entry1.bind('<Button-1>', populate_individual_case)
        entry1.grid(row=row, column=11, sticky=W)
        
        assignedcasevariables.append(entry1)
        count += 1
        entry2 = (str("e"+(str(count))))
        Entry2a = StringVar()
        Entry2a.set(each[2])
        Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
        entry2 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry2a)
        entry2.grid(row=row, column=12, sticky=EW, ipady=3, padx=2)
        assignedcasevariables.append(entry2)
        count += 1
        entry3 = (str("e"+(str(count))))
        Entry3a = StringVar()
        Entry3a.set(each[3])
        Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
        entry3 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry3a)
        entry3.grid(row=row, column=13, sticky=W, ipady=3)
        assignedcasevariables.append(entry3)
        count += 1
        entry4 = (str("e"+(str(count))))
        Entry4a = StringVar()
        Entry4a.set(each[4])
        Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
        entry4 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry4a)
        entry4.grid(row=row, column=14, sticky=W, ipady=3)
        assignedcasevariables.append(entry4)
        count += 1
        entry5 = (str("e"+(str(count))))
        Entry5a = StringVar()
        Entry5a.set(each[5])
        Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
        entry5 = Entry(startup, relief=SUNKEN,width=25, state='disabled', textvariable=Entry5a)
        entry5.grid(row=row, column=15, columnspan=2, sticky=W, ipady=3)
        assignedcasevariables.append(entry5)
        count += 1
        entry6 = (str("e"+(str(count))))
        Entry6a = StringVar()
        Entry6a.set(each[6])
        Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
        entry6 = Entry(startup, relief=SUNKEN,width=140, state='disabled', textvariable=Entry6a)
        entry6.grid(row=row, column=16, columnspan=2, sticky=W, ipady=3)
        assignedcasevariables.append(entry6)
        count += 1
        row += 1

    startup.update
############################################################################################
def populatetriagecases():
    global row, Entry1a, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, entry1, Entry2a, Entry3a, Entry4a, Entry5a, Entry6a, opencasevariables
#    edit1b.config(command=disable_edit2)
    readclosedcases()
        
    for each in pindecryptvariables:
        each.destroy()
    for each in opencasevariables:
        each.destroy()
    for each in individualcasevariables:
        each.destroy()
    for each in overviewvariables:
        each.destroy()
        
    row=2
    Label20 = Label(startup, textvariable='DFT Ref', width=165,  borderwidth=3, relief="groove")
    Label20.grid(row=row, column=1, columnspan=50, sticky=EW, ipady=3)
    closedcasevariables.append(Label20)
    print('writing open cases: ')
    print(len(opencases))
    count=1
    v0 =  StringVar()
    Label0 = Label(startup, textvariable=v0)
    Label0.grid(row=row, column=2, sticky=W)
    v0.set('DFT  Ref')
    closedcasevariables.append(Label0)
        
    v1 =  StringVar()
    Label1 = Label(startup, textvariable=v1)
    Label1.grid(row=row, column=3, sticky=W)
    v1.set('Crime Ref')
    closedcasevariables.append(Label1)

    v2 =  StringVar()
    Label2 = Label(startup, textvariable=v2)
    Label2.grid(row=row, column=4, sticky=W)
    v2.set('Exhibit Ref')
    closedcasevariables.append(Label2)

    v3 =  StringVar()
    Label3 = Label(startup, textvariable=v3)
    Label3.grid(row=row, column=5, sticky=W)
    v3.set('Suspect')
    closedcasevariables.append(Label3)

    v4 =  StringVar()
    Label4 = Label(startup, textvariable=v4)
    Label4.grid(row=row, column=6, sticky=W)
    v4.set('OIC')
    closedcasevariables.append(Label4)

    v5 =  StringVar()
    Label5 = Label(startup, textvariable=v5)
    Label5.grid(row=row, column=7, sticky=W)
    v5.set('Case password')
    closedcasevariables.append(Label5)

    v6 =  StringVar()
    Label6 = Label(startup, textvariable=v6)
    Label6.grid(row=row, column=8, sticky=W)
    v6.set('Date Closed')
    closedcasevariables.append(Label6)

    v7 =  StringVar()
    Label7 = Label(startup, textvariable=v7)
    Label7.grid(row=row, column=11, sticky=W)
    v7.set('')
    closedcasevariables.append(Label7)
    
    v8 =  StringVar()
    Label8 = Label(startup, textvariable=v8)
    Label8.grid(row=row, column=12, sticky=W)
    v8.set('')
    closedcasevariables.append(Label8)

    v9 =  StringVar()
    Label9 = Label(startup, textvariable=v9)
    Label9.grid(row=row, column=13, sticky=W)
    v9.set('')
    closedcasevariables.append(Label9)

    edit1b.grid(row=0, column=15, columnspan=1, sticky=W)
    edit1b.grid_forget()

#    launch.grid(row=0, column=13, columnspan=1, sticky=W)

    row += 1
    column = 2
    for each in triage:
        # DFT full
        Entry1a = StringVar()
        Entry1a.set(each[11])
        dft=each[11]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(startup, text=dft, width= 15,  state='disabled', textvariable=Entry1a)
        #entry1.bind('<Button-1>', populate_individual_case)
        entry1.grid(row=row, column=column, sticky=EW)

        column+=1
        print('closed cases: ', each[11])
        # Crime ref
        closedcasevariables.append(entry1)
        count += 1
        entry2 = (str("e"+(str(count))))
        Entry2a = StringVar()
        Entry2a.set(each[3])
        Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
        entry2 = Entry(startup, relief=SUNKEN,width=15, state='disabled', textvariable=Entry2a)
        entry2.grid(row=row, column=column, sticky=EW, ipady=3, padx=2)
        closedcasevariables.append(entry2)
        count += 1
        column+=1
        # Exhibit
        entry3 = (str("e"+(str(count))))
        Entry3a = StringVar()
        Entry3a.set(each[4])
        Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
        entry3 = Entry(startup, relief=SUNKEN,width=25, state='disabled', textvariable=Entry3a)
        entry3.grid(row=row, column=column, sticky=W, ipady=3)
        closedcasevariables.append(entry3)
        count += 1
        column+=1
        # Suspect
        entry4 = (str("e"+(str(count))))
        Entry4a = StringVar()
        Entry4a.set(each[8])
        Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
        entry4 = Entry(startup, relief=SUNKEN,width=35, state='disabled', textvariable=Entry4a)
        entry4.grid(row=row, column=column, sticky=W, ipady=3)
        closedcasevariables.append(entry4)
        count += 1
        column+=1
        # OIC
        entry5 = (str("e"+(str(count))))
        Entry5a = StringVar()
        Entry5a.set(each[7])
        Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
        entry5 = Entry(startup, relief=SUNKEN,width=35, state='disabled', textvariable=Entry5a)
        entry5.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        closedcasevariables.append(entry5)
        count += 1
        column+=1
        # Case password
        entry6 = (str("e"+(str(count))))
        Entry6a = StringVar()
        Entry6a.set(each[12])
        Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
        entry6 = Entry(startup, relief=SUNKEN,width=25, state='disabled', textvariable=Entry6a)
        entry6.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        closedcasevariables.append(entry6)
        count += 1
        column+=1
        # Date closed
        entry7 = (str("e"+(str(count))))
        Entry7a = StringVar()
        Entry7a.set(each[13])
        Entry7a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry7a))
        entry7 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry7a)
        entry7.grid(row=row, column=column, columnspan=1, sticky=W, ipady=3)
        closedcasevariables.append(entry7)
        

        count += 1
        row += 1
        column = 2

    photoScrollv = Scrollbar(photoFrame1, orient=VERTICAL)
    photoScrollv.config(command=photoCanvas1.yview)
    photoScrollh = Scrollbar(photoFrame1, orient=HORIZONTAL)
    photoScrollh.config(command=photoCanvas1.xview)
    photoCanvas1.config(yscrollcommand=photoScrollv.set)
    photoCanvas1.config(xscrollcommand=photoScrollh.set)
    photoScrollv.grid(row=0, column=1, sticky="ns")
    photoScrollh.grid(row=row, column=0, sticky="ew")
    startup.bind("<Configure>", update_scrollregion)
    root.update
##############################################################################################
def populate_individual_case():
    global TEXT
    case_get = d1a.get()
    print('case get', case_get)
    TEXT = case_get[:7]
    get_individual_case_data()
    try:
        case_get = d1a.get()
        print('case get', case_get)
        TEXT = case_get[:7]
        get_individual_case_data()
    except:
        pass
############################################################################################
def populate_individual_case_2():
    global TEXT
    try:
        case_get = d1a.get()
        print('case get', case_get)
        TEXT = case_get[:7]
        get_individual_case_data()
    except:
        pass
############################################################################################
##def populate_individual_pd(event):
##    global TEXT
##    print(event.widget["text"])
###    event.widget["text"] = "DONE"
###    event.widget["bg"] = "green"
##    TEXT = (event.widget["text"])
##    print('button text: ', TEXT)
###    print(entry1.keys())
###    for each in pindecryptvariables:
###        each.destroy()
###    print('open case vars: ', opencasevariables)
##    get_individual_pd_data()
############################################################################################
def writeentry():
    global row, pindecryptlogs, statusbuttonb
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" insert into pindecryptlog values (?, ?, ?, ?, ?, ?, ?, ?, '-');""", (row, new1, new2, new3, new4, datestamp,'0', analyst,))
    conn.commit()
    conn.close()

    Newentry1a.set('')
    Newentry2a.set('')
    Newentry3a.set('')
    Newentry4a.set('')
    
    pindecryptcount = (len(pindecryptlogs))
    print("No of logs: ", pindecryptcount)

    readpindecryptlog()
############################################################################################
def func(value):
    pass
############################################################################################
def widthNormal():
    root.geometry('610x720')
    width.config(text='>')
    width.config(command=widthWide)
#    notepad_n59b.configure(width=74)
    root.update()
############################################################################################
def widthWide():
    # getting screen's height in pixels 
    screen_height = root.winfo_screenheight()
    # getting screen's width in pixels 
    screen_width = root.winfo_screenwidth()
    print('screen_width: ', screen_width)
    if analyst == 'R WARD 32533':
        screen_width = 1860
    screen_resolution = str(screen_width)+'x'+str(screen_height)
    root.geometry(screen_resolution)
    width.config(text='<')
    width.config(command=widthNormal)
    root.update()
############################################################################################
def addtopdlog():
    global row
    count=1
    addtopdl.config(state='disabled')
    dft_ref = (e1a.get().upper())
    crime_ref = (e2a.get().upper())
    exhib_ref = (e3a.get().upper())
    oic = (e5a.get().upper())
    
    populatepinlog()
    viewTypea.set('PIN decryption log')
    
    entry1 = (str("e"+(str(count))))
    Entry1a = StringVar()
    Entry1a.set(dft_ref)
    Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
    entry1 = Entry(startup, relief=SUNKEN,width=10, state='disabled', textvariable=Entry1a)
    entry1.grid(row=row, column=11, sticky=W)
    pindecryptvariables.append(entry1)
    count += 1
    entry2 = (str("e"+(str(count))))
    Entry2a = StringVar()
    Entry2a.set(crime_ref)
    Entry2a.trace("w", lambda name, index, mode, Entry2a=Entry2a: callback(Entry2a))
    entry2 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry2a)
    entry2.grid(row=row, column=12, sticky=EW)
    pindecryptvariables.append(entry2)
    count += 1
    entry3 = (str("e"+(str(count))))
    Entry3a = StringVar()
    Entry3a.set(exhib_ref)
    Entry3a.trace("w", lambda name, index, mode, Entry3a=Entry3a: callback(Entry3a))
    entry3 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry3a)
    entry3.grid(row=row, column=13, sticky=W)
    pindecryptvariables.append(entry3)
    count += 1
    entry4 = (str("e"+(str(count))))
    Entry4a = StringVar()
    Entry4a.set(oic)
    Entry4a.trace("w", lambda name, index, mode, Entry4a=Entry4a: callback(Entry4a))
    entry4 = Entry(startup, relief=SUNKEN,width=20, state='disabled', textvariable=Entry4a)
    entry4.grid(row=row, column=14, sticky=W)
    pindecryptvariables.append(entry4)
    count += 1
    entry5 = (str("e"+(str(count))))
    Entry5a = StringVar()
    Entry5a.set(datestamp)
    Entry5a.trace("w", lambda name, index, mode, Entry5a=Entry5a: callback(Entry5a))
    entry5 = Entry(startup, relief=SUNKEN,width=12, state='disabled', textvariable=Entry5a)
    entry5.grid(row=row, column=15, columnspan=2, sticky=W)
    pindecryptvariables.append(entry5)
    count += 1
    selected_month_rec = (datestamp)
    start = date(int(selected_month_rec.split('-')[0]),int(selected_month_rec.split('-')[1]),int(selected_month_rec.split('-')[2]))
    today = date.today()
    res = today - start
    res.days
    entry6 = (str("e"+(str(count))))
    Entry6a = StringVar()
    Entry6a.set(res.days)
    Entry6a.trace("w", lambda name, index, mode, Entry6a=Entry6a: callback(Entry6a))
    entry6 = Entry(startup, relief=SUNKEN,width=12, state='disabled', textvariable=Entry6a)
    entry6.grid(row=row, column=16, sticky=W)
    pindecryptvariables.append(entry6)

    row+=1

    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" insert into pindecryptlog values (?, ?, ?, ?, ?, ?, ?, ?, '-');""", (row, dft_ref, crime_ref, exhib_ref, oic, datestamp,'0', analyst,))
    conn.commit()
    conn.close()

    t4 = threading.Thread(target=overview)
    t4.start()
    #populatepinlog()
    root.update()
############################################################################################
def clearall():
    for each in pindecryptvariables:
        each.destroy()
    
    pindecryptcount = (len(pindecryptlogs))
    print("No of logs: ", pindecryptcount)
    populatepinlog()
############################################################################################
def clearall2():
    for each in opencasevariables:
        each.destroy()
    populateopencases()
############################################################################################
def clearall3():
    for each in assignedcasevariables:
        each.destroy()
    populateassignedcases()
############################################################################################
def deleteentry():
    global options
    confirm.destroy()
    pd_case.destroy()

    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("DELETE FROM pindecryptlog WHERE full_dft_ref = '%'||?||'%' and exhibit = ?", (d, x,))
    conn.commit()
    conn.close()

    readpindecryptlog()
    pindecryptcount = (len(pindecryptlogs))
    print("No of logs: ", pindecryptcount)

    clearall2()
    overview()
    enable_edit1()
############################################################################################
def deleteentry2():
    global options
    confirmclose()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("DELETE FROM opencases WHERE full_dft_ref like '%'||?||'%'", (d,))
    c.execute(" update cases SET status=? WHERE full_dft_ref like '%'||?||'%'", ('closed', d,))
    conn.commit()
    conn.close()

    readopencases()
    clearall()
    #t10 = threading.Thread(target=populateopencases)
    #t10.start()
    case.destroy()
    overview()

############################################################################################
def deleteentry3():
    global options

    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("DELETE FROM assignedcases WHERE full_dft_ref like '%'||?||'%'", (d,))
    conn.commit()
    conn.close()

    readopencases()
    a1a.set('')
    a2a.set('')
    a3a.set('')
    a4a.set('')
    a5a.set('')
    clearall()
    populateassignedcases
    enable_edit3()
############################################################################################
def exportpw():
    global level8, level7, level6, case_password
    if casetype =='Loaded':
        case_password == case_password_loaded
    else:
        case_password == case_password
    os.chdir(DftFolder)
    print('3732 - exportpw', cwd)
    
    output_file = open(case_password + ".txt", "w")
    output_file.write(case_password)
    output_file.close()
############################################################################################
def exportpwtolocation2():
    global desktop
    d = argv[1] if len(argv)>1 else filedialog.askdirectory(initialdir="C:\\")
    D = os.path.realpath(d)
    startup_e2b.delete(0,END)
    startup_e2b.insert(10,D)
#############################################################################################Copy MPER and Log over to exhibit folder folder
def copycontemp():
    if e33b == 1:
        shutil.copy(template_dir + '\\' + log_file, exhibfolder + '\\' + log_file)
        
    if e35b == 1:
        shutil.copy(template_dir + '\\' + mper_file, exhibfolder + '\\' + mper_file)
############################################################################################
def do_file(FN):
    hash_type = e23a.get()
    e23b.config(width=20, state='disabled')
    buf = 2**20
    size = os.path.getsize(FN)
    PB.configure(maximum=size)
    File = open(FN, 'rb')
    if hash_type == 'md5':
        MD5 = hashlib.md5()
    elif hash_type == 'sha1':
        MD5 = hashlib.sha1()
    elif hash_type == 'sha256':
        MD5 = hashlib.sha256()
    def _com():
            while True:
                    d = File.read(buf)
                    if not d: break
                    PB.step(buf)
                    MD5.update(d)
            File.close()
            PB.stop()
            PB.pack_forget()
            e19a.set(MD5.hexdigest())
            if not os.path.isfile(FN+"."+hash_type):
                    open(FN+"."+hash_type, 'w').write("{0}\t{1}".format(e19a.get(), os.path.split(FN)[1]))
            Browse.config(state='disabled')
            #Reset.config(state='normal')
            
    Thread(target=_com).start()
    e20b.config(state='normal')
############################################################################################
def select_file():
    e20b.config(width=20, state='disabled')
    f = argv[1] if len(argv)>1 else filedialog.askopenfilename(initialdir=os.getcwd())
    F = os.path.realpath(f)
    (do_file(F), e19a.set('Please wait...')) if os.path.isfile(F) else print("Error")
    root.update
############################################################################################	
def to_clip():
    tools.clipboard_clear()
    tools.clipboard_append(e19a.get())
    e19b.config(foreground='green')
    e19b.selection_range(0, END)
############################################################################################
def hash_reset():
    e23b.config(width=20, state='normal')
    e20b.config(width=20, state='disabled')
    Reset.config(state='disabled')
    Browse.config(state='normal')
    e19a.set("")
############################################################################################
def callbacke21(sv):
    length=sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET pw_length = ? WHERE _rowid_ = 1', (length,))
    conn.commit()
    conn.close()
############################################################################################
def callbacke22(sv):
    incdft=sv.get()
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute('UPDATE profile SET pw_incdft = ? WHERE _rowid_ = 1', (incdft,))
    conn.commit()
    conn.close()

    print("callback22", incdft)
    if incdft == 1:
        #startup_e1a.set(pwdft)
        print("on")
        genPw()
    else:
        #startup_e1a.set(pw)
        print("off")
        updatePw()
############################################################################################
def checktemplates():
    preimg_doc()
    if contemp_copy == 1:
        if contemp_populate == 1:
            contemp_copy_populate()
        else:
            contemp_copy_only()
    else:
        print("contemp - Do not copy")
        
    if sfr_copy == 1:
        if sfr_populate == 1:
            sfr_copy_populate()
        else:
            sfr_copy_only()
    else:
        print("sfr - Do not copy")
        
    if disclosure_copy == 1:
        if disclosure_populate == 1:
            disclosure_copy_populate()
        else:
            disclosure_copy_only()
    else:
        print("disclosure - Do not copy")
    #folderautobuild()
############################################################################################
def contemp_copy_only():
    # Open template
    #contemp_file = e12.get()
    _templates = template_dir
    _temp = os.path.join(_templates, contemp_file)
    document = docx.Document(_temp)
    level8 = case_dir+'\\'+"DFT"+'\\'+year+'\\'+crime_ref+'\\'+analyst+'\\'+exhib_ref+'\\'
    document.save(level8+"\\"+analystinitials+"-"+dft_ref+"-"+exhib_ref+"-NOTES.docx")
############################################################################################
def sfr_copy_only():
    # Open template
    #sfr_file = e25.get()
    _templates = template_dir
    _temp = os.path.join(_templates, sfr_file)
    print(_temp)
    document = docx.Document(_temp)
    document.save(level6+"\\"+analystinitials+"-"+dft_ref+"-SFR-1.docx")
############################################################################################
def disclosure_copy_only():
    if disclosure_copy == 1:
        # Open template
        #disclosure_file = e32.get()
        try:
            _templates = template_dir
            _temp = os.path.join(_templates, disclosure_file)
            document = docx.Document(_temp)
            document.save(level6+"\\"+analystinitials+"-"+dft_ref+"-Disclosure-1.docx")
        except:
            pass
############################################################################################
def preimg_doc():
    print('template_dir: ', template_dir)
    try:
        if os.path.isfile(template_dir+'\\PRE-IMG.docx'):
            print('Pre-img is file')
            pre_img = 'PRE-IMG.docx'
            _templates = template_dir
            _temp = os.path.join(_templates, pre_img)
            print(_temp)
            document = docx.Document(_temp)
            '''style = document.styles['Normal']
            font = style.font
            table1 = document.tables[0]
            table2 = document.tables[1]
            table3 = document.tables[2]
            table4 = document.tables[3]
            #Presented as (row no, column no)
                
            #table.style = 'TableGrid'
            font.name = 'Arial Narrow'
            font.size = Pt(28)
            
            cell_1 = table1.cell(0,1).text = 'DFT-'+dft_ref
            cell_2 = table1.cell(1,1).text = 'Exhibit: '+ori_exhibit
            def make_rows_bold(*rows):
                for row in rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            make_rows_bold(table1.rows[1])

            cell_1 = table2.cell(0,1).text = 'DFT-'+dft_ref
            cell_2 = table2.cell(1,1).text = 'Exhibit: '+ori_exhibit+'-SIM-1'
            def make_rows_bold(*rows):
                for row in rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            make_rows_bold(table2.rows[1])

            cell_1 = table3.cell(0,1).text = 'DFT-'+dft_ref
            cell_2 = table3.cell(1,1).text = 'Exhibit: '+ori_exhibit+'-SIM-2'
            def make_rows_bold(*rows):
                for row in rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            make_rows_bold(table3.rows[1])

            cell_1 = table4.cell(0,1).text = 'DFT-'+dft_ref
            cell_2 = table4.cell(1,1).text = 'Exhibit: '+ori_exhibit+'-MC'
            def make_rows_bold(*rows):
                for row in rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
            make_rows_bold(table4.rows[1])'''
            level8 = case_dir+'\\'+"DFT"+'\\'+year+'\\'+crime_ref+'\\'+analyst+'\\'+exhib_ref+'\\'
            document.save(level8+"\\"+analystinitials+"-"+dft_ref+"-"+exhib_ref+"-Pre-Img.docx")#(level7+"\\Pre-Img.docx")
            
        else:
            pass 
    except:
        pass
############################################################################################
def folderautobuild():
    
    if pp1==1:
        createfolder1()
    else:
        print("o")
    if pp2==1:
        createfolder2()
    else:
        print("o")
    if pp3==1:
        createfolder3()
    else:
        print("o")
    if pp4==1:
        createfolder4()
    else:
        print("o")

############################################################################################
def update_scrollregion(event):
    photoCanvas1.configure(scrollregion=photoCanvas1.bbox("all"))
############################################################################################
def showallcasedata():
    global photoFrame, photoCanvas, canvasFrame
    readcases()
    row=1
    column=1
    case = Toplevel()
    width  = root.winfo_screenwidth()
    height = root.winfo_screenheight()
    print(width)
    print(height)
    
    photoFrame = Frame(case, width=width, height=height)
    photoFrame.grid()
    photoFrame.rowconfigure(0, weight=1) 
    photoFrame.columnconfigure(0, weight=1) 

    photoCanvas = Canvas(photoFrame, width=(width)-85, height=(height)-45)
    photoCanvas.grid(sticky=NSEW)

    canvasFrame = Frame(photoCanvas, width=width, height=10000)
    photoCanvas.create_window(0, 0, window=canvasFrame, anchor='nw')

    Label1=Label(canvasFrame,text='#', width=2, anchor='w').grid(row=1,column=1, sticky=W)
    Label2=Label(canvasFrame,text='Year', width=5, anchor='w').grid(row=1,column=2, sticky=W)
    Label3=Label(canvasFrame,text='DFT Ref', width=7, anchor='w').grid(row=1,column=3, sticky=W)
    Label4=Label(canvasFrame,text='Crime Ref', width=12, anchor='w').grid(row=1,column=4)
    Label(canvasFrame,text='Exhibit', width=20, anchor='w').grid(row=1,column=5, sticky=W)
    Label(canvasFrame,text='Bag Seal', width=15, anchor='w').grid(row=1,column=6, sticky=W)
    Label(canvasFrame,text='Op Name', width=15, anchor='w').grid(row=1,column=7, sticky=W)
    Label(canvasFrame,text='OIC', width=15, anchor='w').grid(row=1,column=8, sticky=W)
    Label(canvasFrame,text='Supect', width=15, anchor='w').grid(row=1,column=9, sticky=W)
    Label(canvasFrame,text='Property Ref', width=10, anchor='w').grid(row=1,column=10, sticky=W)
    Label(canvasFrame,text='Examiner', width=18, anchor='w').grid(row=1,column=11, sticky=W)
    Label(canvasFrame,text='Case Password', width=20, anchor='w').grid(row=1,column=12, sticky=W)
    Label(canvasFrame,text='Date Commenced', width=15, anchor='w').grid(row=1,column=13, sticky=W)
    
    column+=1
        
    row+=1
    column=1
    for each in cases:
        Label(canvasFrame,text=(each[0])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[1])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[2])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[3])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[4])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[5])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[6])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[7])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[8])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[9])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[10])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[12])).grid(row=row,column=column, sticky=W)
        column+=1
        Label(canvasFrame,text=(each[13])).grid(row=row,column=column, sticky=W)
        column=1
        row+=1
    
    photoScrollv = Scrollbar(photoFrame, orient=VERTICAL)
    photoScrollv.config(command=photoCanvas.yview)
    photoScrollh = Scrollbar(photoFrame, orient=HORIZONTAL)
    photoScrollh.config(command=photoCanvas.xview)
    photoCanvas.config(yscrollcommand=photoScrollv.set)
    photoCanvas.config(xscrollcommand=photoScrollh.set)
    photoScrollv.grid(row=0, column=1, sticky="ns")
    photoScrollh.grid(row=row, column=0, sticky="ew")
    canvasFrame.bind("<Configure>", update_scrollregion)

    case.mainloop()
############################################################################################
def updatepinentry():
    d = (edita.get())
    x = d.split(' - ')
    print(x)
    
    z=a1a.get()
    y=a2a.get()
    w=a3a.get()
    q=a4a.get()
    t=a6a.get()

    slashes = ('/', ':', '|', '\\', '?')
    for item in slashes:
        q = str.replace(q, item,'-')

    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE pindecryptlog SET crime_ref = ?,
              exhibit = ?,
              oic = ?,
              datestamp = ?,
              pin_notes = ? WHERE full_dft_ref like '%'||?||'%' and exhibit = ?""", (z, y, w, q, t, x[0], x[1],))
    '''pin_notes = ? WHERE full_dft_ref like '%'||?||'%' and exhibit like '%'||?||'%'""", (z, y, w, q, t, x[0], x[1],))
    WHERE id = ?'''
    conn.commit()
    conn.close()
    enable_edit1()
    clearall()
############################################################################################
def updateopencaseentry():
    d = (edit2a.get())

    z=a1a.get()
    y=a2a.get()
    w=a3a.get()
    q=a4a.get()
    p=a5a.get()
    
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE opencases SET crime_ref = ?,
              no_of_exhib = ?,
              oic = ?,
              datestamp = ?,
              note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))
    conn.commit()
    conn.close()
    enable_edit2()
    clearall2()
############################################################################################
def updateassignedcaseentry():
    d = (edit3a.get())

    z=a1a.get()
    y=a2a.get()
    w=a3a.get()
    q=a4a.get()
    p=a5a.get()
    
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE assignedcases SET crime_ref = ?,
              no_of_exhib = ?,
              oic = ?,
              datestamp = ?,
              notes = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))
    conn.commit()
    conn.close()
    enable_edit3()
    clearall3()
############################################################################################
def disable_edit1():
    edit1b.config(state='disabled')
    startup.update
    add_new_PD_entry()
############################################################################################
def enable_edit1():
    edit1b.config(state='normal')
    startup.update
############################################################################################
def disable_edit2():
    edit1b.config(state='disabled')
    startup.update
    edit_opencase_entry()
############################################################################################
def enable_edit2():
    edit1.destroy()
    edit1b.config(state='normal')
    startup.update
############################################################################################
def disable_edit3():
    edit1b.config(state='disabled')
    startup.update
    edit_assignedcase_entry()
############################################################################################
def enable_edit3():
    edit2.destroy()
    edit1b.config(state='normal')
    startup.update
############################################################################################
def callback_Indiv_pd_case2(Indiv_pd_case2a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    get=Indiv_pd_case2.get()
    print(get)
    try:
        for sql in sqlitedbs:
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET crime_ref = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_pd_case,))
            '''c.execute("""UPDATE pindecryptlog SET crime_ref = ?,
                      exhibit = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
#        Indiv_pd_case2.config(state='disabled')
#        Indiv_Case3.config(state='disabled')
#        Indiv_Case4.config(state='disabled')
#        Indiv_Case5.config(state='disabled')
#        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_pd_case3(Indiv_Case3a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    get=Indiv_pd_case3.get()
    print(get)
    try:
        for sql in sqlitedbs:
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET exhibit = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_pd_case,))
            '''c.execute("""UPDATE pindecryptlog SET crime_ref = ?,
                      exhibit = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
#        Indiv_pd_case2.config(state='disabled')
#        Indiv_Case3.config(state='disabled')
#        Indiv_Case4.config(state='disabled')
#        Indiv_Case5.config(state='disabled')
#        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_pd_case4(Indiv_Case4a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    get=Indiv_pd_case4.get()
    try:
        for sql in sqlitedbs:
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET oic = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_pd_case,))
            '''c.execute("""UPDATE pindecryptlog SET crime_ref = ?,
                      exhibit = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
#        Indiv_pd_case2.config(state='disabled')
#        Indiv_Case3.config(state='disabled')
#        Indiv_Case4.config(state='disabled')
#        Indiv_Case5.config(state='disabled')
#        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_pd_case5(Indiv_Case5a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    get=Indiv_pd_case5.get()
    try:
        for sql in sqlitedbs:
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET datestamp = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_pd_case,))
            '''c.execute("""UPDATE pindecryptlog SET crime_ref = ?,
                      exhibit = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
#        Indiv_pd_case2.config(state='disabled')
#        Indiv_Case3.config(state='disabled')
#        Indiv_Case4.config(state='disabled')
#        Indiv_Case5.config(state='disabled')
#        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def Indiv_pd_notes(*args):
    get=Indiv_pd_case6.get(1.0, END)
    print(get)
    print(Indiv_pd_case)
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET pin_notes = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_pd_case,))
            conn.commit()
            conn.close()
    except:
        pass
#        Indiv_Case2.config(state='disabled')
#        Indiv_Case3.config(state='disabled')
#        Indiv_Case4.config(state='disabled')
#        Indiv_Case5.config(state='disabled')
#        Indiv_Case6.config(state='disabled')
        pd_case.update
############################################################################################
def callback_Indiv_Case2(Indiv_Case2a):
    sqlitedbs = (sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_Case2.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE opencases SET crime_ref = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_Case3(Indiv_Case3a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_Case3.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE opencases SET no_of_exhib = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_Case4(Indiv_Case4a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_Case4.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE opencases SET oic = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_Case5(Indiv_Case5a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_Case5.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE opencases SET datestamp = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def Indiv_Case_notes(*args):
    get=Indiv_Case6.get(1.0, END)
    print(get)
    print(Indiv_Case)
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE opencases SET note = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_Case_status(status_dda):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=status_dda.get()
            print(get)
            print(Indiv_Case)
            print(the_exhibit)
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE cases SET status = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case))#, the_exhibit ))

            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        overview.update
############################################################################################
def callback_Loop_Indiv_Case_property_ref(Loop_Indiv_Case3a):
    get=Loop_Indiv_Case3a.get()
    
    print(get)
    print(Indiv_Case)
    tab_id = tabs.index("current")
    tab_data =(tabs.tab(tab_id))
    print(tab_data)
    for data in tab_data:
        print('tab data: ', data)
        #print(tabs.select(tab_id))
        the_exhibit = tab_data['text']
    print(the_exhibit)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE cases SET g83 = ? WHERE full_dft_ref like '%'||?||'%' and exhibit = ?""", (get, Indiv_Case, the_exhibit ))

    conn.commit()
    conn.close()
############################################################################################
def callback_Loop_Indiv_Case_exhibit(Loop_Indiv_Case2a):
    get=Loop_Indiv_Case2a.get()
    
    print(get)
    print(Indiv_Case)
    tab_id = tabs.index("current")
    tab_data =(tabs.tab(tab_id))
    print(tab_data)
    for data in tab_data:
        print('tab data: ', data)
        #print(tabs.select(tab_id))
        the_exhibit = tab_data['text']
    print(the_exhibit)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE cases SET bag_seal = ? WHERE full_dft_ref like '%'||?||'%' and exhibit = ?""", (get, Indiv_Case, the_exhibit ))

    conn.commit()
    conn.close()
###########################################################################################
def Loop_Indiv_Action_exhibit_notes(*args):
    print(Indiv_Case)
    tab_id = tabs.index("current")
    print('tab id:', tab_id)
    tab_data =(tabs.tab(tab_id))
    #print(tab_data)
    for data in tab_data:
        print('tab data: ', data)
        print(tabs.select(tab_id))
        the_exhibit = tab_data['text']
    print(the_exhibit)
    get=ryan[the_exhibit].get(1.0, END)#('t'(str(tab_id))).get(1.0, END)
    print('get',get)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE cases SET exhib_notes = ? WHERE full_dft_ref like '%'||?||'%' and exhibit = ?""", (get, Indiv_Case, the_exhibit ))

    conn.commit()
    conn.close()
###########################################################################################
def Loop_Indiv_Case_exhibit_notes(*args):
    print(Indiv_Case)
    tab_id = tabs.index("current")
    print('tab id:', tab_id)
    tab_data =(tabs.tab(tab_id))
    #print(tab_data)
    for data in tab_data:
        print('tab data: ', data)
        print(tabs.select(tab_id))
        the_exhibit = tab_data['text']
    print(the_exhibit)
    get=ryan[the_exhibit].get(1.0, END)#('t'(str(tab_id))).get(1.0, END)
    print('get',get)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE cases SET exhib_notes = ? WHERE full_dft_ref like '%'||?||'%' and exhibit = ?""", (get, Indiv_Case, the_exhibit ))

    conn.commit()
    conn.close()
###########################################################################################
def Loop_Indiv_Case_exhibit_location(*args):
    print(Indiv_Case)
    tab_id = tabs.index("current")
    print('tab id:', tab_id)
    tab_data =(tabs.tab(tab_id))
    #print(tab_data)
    for data in tab_data:
        print('tab data: ', data)
        print(tabs.select(tab_id))
        the_exhibit = tab_data['text']
    print(the_exhibit)
    get=ryan[the_exhibit].get(1.0, END)#('t'(str(tab_id))).get(1.0, END)
    print('get',get)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE cases SET exhib_notes = ? WHERE full_dft_ref like '%'||?||'%' and exhibit = ?""", (get, Indiv_Case, the_exhibit ))

    conn.commit()
    conn.close()
###########################################################################################
def Loop_Indiv_Case_exhibit_property_of(*args):
    print(Indiv_Case)
    tab_id = tabs.index("current")
    print('tab id:', tab_id)
    tab_data =(tabs.tab(tab_id))
    #print(tab_data)
    for data in tab_data:
        print('tab data: ', data)
        print(tabs.select(tab_id))
        the_exhibit = tab_data['text']
    print(the_exhibit)
    get=ryan[the_exhibit].get(1.0, END)#('t'(str(tab_id))).get(1.0, END)
    print('get',get)
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute("""UPDATE cases SET exhib_notes = ? WHERE full_dft_ref like '%'||?||'%' and exhibit = ?""", (get, Indiv_Case, the_exhibit ))

    conn.commit()
    conn.close()
###########################################################################################
def close_individual_case():
    case.destroy()
    #return_to_casework_from_individual_case()
    #print('case_work_vars', case_work_vars)
    #populateopencases()
############################################################################################    
def close_individual_pd_case():
    pd_case.destroy()
    populatepinlog()
##############################################################################################
def get_individual_case_data():
    global d, case, tabs, key, value, ryan, Indiv_Case, Indiv_Case1, Indiv_Case2, Indiv_Case3, Indiv_Case4, Indiv_Case5, Indiv_Case6, Indiv_Case7, Indiv_Case8, Indiv_Case8a, the_exhibit, Loop_Indiv_Case1, Loop_Indiv_Case2, Loop_Indiv_Case3, Loop_Indiv_Case1a, Loop_Indiv_Case2a, Loop_Indiv_Case3a, Loop_Indiv_Case6
    the_case = []
    ryan = {}
#    launch.grid_forget()
    d = (TEXT[5:12])
    print('ddddddddddddd', d)
    if d == "No entries":
        pass
    else:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute("SELECT * FROM cases WHERE full_dft_ref like '%'||?||'%' and status != ?", (d, 'closed',))
        
        for each in c:
            print(each)
            the_case.append(each)
        case_status = (each[15])
        the_exhibit = (each[4])
        case_pw = (each[12])
        exhibit_notes = (each[14])
        print('The Case no of entries: ', (len(the_case)))
        try:
            if 'normal' == case.state():
                case.destroy()
            elif 'normal' == pd_case.state():
                pd_case.destroy()
            else:
                pass
        except:
            pass
        case = Toplevel()
        width  = case.winfo_screenwidth()
        height = case.winfo_screenheight()
        case.geometry('740x600') #widthxheight
        print(width)
        print(height)
        case.title("CASE - "+TEXT)
        case.resizable(width=False, height=False)
        case.protocol("WM_DELETE_WINDOW", close_individual_case)

        tabs = ttk.Notebook(case)
        overview=ttk.Frame(tabs)
        tabs.add(overview,text='Overview')

        print('The case: ', the_case)
        count_no2 = 1 
        for exhib in the_case:
            #tabs = ttk.Notebook(case)
            print('the case: ',exhib)
            name=str(exhib[4])
            print(name)
            print(type(name))
            name=ttk.Frame(tabs)
            tabs.add(name,text=(exhib[4]))

            rows=31
            columns=9
            conn.close()
            # Set row height here
            height=5
            for frame in the_case:
                count=1
                count2=1

                if frame == the_case[0]:
                    while count < rows:
                        Label(name, text="", width=1, height=1).grid(row=count, column=0, sticky=EW)
                        count+=1
                    while count2 < columns:
                        Label(name, text="", width=1, height=1).grid(row=0, column=count2, sticky=EW)
                        count2+=1
                        
            row = 0
            oic_label=Label(name, text="Exhibit Ref")
            oic_label.grid(row=row, column=1, sticky=W)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Exhibit Seal")
            oic_label.grid(row=row, column=3, sticky=W)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Property Ref")
            oic_label.grid(row=row, column=5, sticky=W)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Property of: ")
            oic_label.grid(row=row, column=7, sticky=EW)
            individualcasevariables.append(oic_label)

            row+=1
            
            Loop_Indiv_Case1a = StringVar()
            Loop_Indiv_Case1a.set(exhib[4])
            Loop_Indiv_Case1a.trace("w", lambda name, index, mode, Loop_Indiv_Case1a=Loop_Indiv_Case1a: callback(Loop_Indiv_Case1a))
            Loop_Indiv_Case1 = Entry(name, relief=SUNKEN,width=15, state='normal', textvariable=Loop_Indiv_Case1a)
            Loop_Indiv_Case1.grid(row=row, column=1, columnspan=2, sticky=EW)
            
            individualcasevariables.append(Loop_Indiv_Case1)
            count += 1
            Loop_Indiv_Case2 = (str("e"+(str(count))))
            Loop_Indiv_Case2a = StringVar()
            Loop_Indiv_Case2a.set(exhib[5])
            Loop_Indiv_Case2a.trace("w", lambda name, index, mode, Loop_Indiv_Case2a=Loop_Indiv_Case2a: callback_Loop_Indiv_Case_exhibit(Loop_Indiv_Case2a))
            Loop_Indiv_Case2 = Entry(name, relief=SUNKEN,width=15, state='normal', textvariable=Loop_Indiv_Case2a)
            Loop_Indiv_Case2.grid(row=row, column=3, columnspan=2, sticky=EW)
            individualcasevariables.append(Loop_Indiv_Case2)
            
            count += 1
            Loop_Indiv_Case3 = (str("e"+(str(count))))
            Loop_Indiv_Case3a = StringVar()
            Loop_Indiv_Case3a.set(exhib[9])
            Loop_Indiv_Case3a.trace("w", lambda name, index, mode, Loop_Indiv_Case3a=Loop_Indiv_Case3a: callback_Loop_Indiv_Case_property_ref(Loop_Indiv_Case3a))
            Loop_Indiv_Case3 = Entry(name, relief=SUNKEN,width=15, state='normal', textvariable=Loop_Indiv_Case3a)
            Loop_Indiv_Case3.grid(row=row, column=5, columnspan=2, sticky=EW)
            individualcasevariables.append(Loop_Indiv_Case3)

            vic_or_susa = StringVar()
            vic_or_susa.set('--Select--')
            vic_or_susa.trace("w", lambda name, index, mode, vic_or_susa=vic_or_susa: callback(vic_or_susa))
            vic_or_susab = OptionMenu(name, vic_or_susa, *property_of, command=qwerty)
            vic_or_susab.grid(row=row, column=7, columnspan=3, sticky=W, pady=5, padx=5)
            individualcasevariables.append(vic_or_susab)
            vic_or_susab.config(width=18)
            
            row+=1

            oic_label=Label(name, text="PRQC:")
            oic_label.grid(row=row, column=1, sticky=EW)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Strategy:")
            oic_label.grid(row=row, column=2, sticky=EW)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Pre-imaging:")
            oic_label.grid(row=row, column=3, sticky=EW)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Imaging:")
            oic_label.grid(row=row, column=4, sticky=EW)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Analysis:")
            oic_label.grid(row=row, column=5, sticky=EW)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Reports:")
            oic_label.grid(row=row, column=6, sticky=EW)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="QC:")
            oic_label.grid(row=row, column=7, sticky=EW)
            individualcasevariables.append(oic_label)

            oic_label=Label(name, text="Docs:")
            oic_label.grid(row=row, column=8, sticky=EW)
            individualcasevariables.append(oic_label)
            
            
            row+=1
            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=1, sticky=EW)
            overviewvariables.append(e30b)
            
            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=2, padx=1, sticky=EW)
            overviewvariables.append(e30b)

            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=3, sticky=EW)
            overviewvariables.append(e30b)
            
            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=4, padx=1, sticky=EW)
            overviewvariables.append(e30b)

            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=5, sticky=EW)
            overviewvariables.append(e30b)
            
            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=6, padx=1, sticky=EW)
            overviewvariables.append(e30b)

            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=7, padx=1, sticky=EW)
            overviewvariables.append(e30b)

            e30a = IntVar()
            e30a.set(0)
            e30a.trace("w", lambda name, index, mode, e30a=e30a: callbacke30(e30a))
            e30b = Checkbutton(name, text=" ", width=5, onvalue=1,offvalue=0,variable=e30a)
            e30b.grid(row=row, column=8, padx=1, sticky=EW)
            overviewvariables.append(e30b)
            

##
##            row=4
##
##            locationa = StringVar()
##            locationa.set('--Select--')
##            locationa.trace("w", lambda name, index, mode, locationa=locationa: callback(locationa))
##            locationab = OptionMenu(name, locationa, *location, command=qwerty)
##            locationab.grid(row=row, column=8, columnspan=1, sticky=W)
##            individualcasevariables.append(locationab)
##            locationab.config(width=18)
##            count_no2+=1
##
##            actiona = StringVar()
##            actiona.set('--Select--')
##            actiona.trace("w", lambda name, index, mode, actiona=actiona: callback(actiona))
##            actionab = OptionMenu(name, actiona, *action, command=qwerty)
##            actionab.grid(row=row, column=1, columnspan=1, sticky=W)
##            individualcasevariables.append(actionab)
##            actionab.config(width=18)
##            count_no2+=1

            row=2

            
##            default_font = font.Font(family="Arial", size=8)
##            submit2a= StringVar()
##            submit2a.trace("w", lambda name, index, mode, submi2ta=submit2a: callback(submit2a))
##            submit2b = Button(name, text="Photography", font=default_font, command=copytoclip_individual_case)
##            submit2b.config(state='disabled')
##            submit2b.grid(row=7, column=1, columnspan=2, sticky=EW)
##            
##
##            default_font = font.Font(family="Arial", size=8)
##            submit2a= StringVar()
##            submit2a.trace("w", lambda name, index, mode, submi2ta=submit2a: callback(submit2a))
##            submit2b = Button(name, text="Pre-imaging", font=default_font, command=copytoclip_individual_case)
##            submit2b.grid(row=7, column=3, columnspan=2, sticky=EW)
##
##            default_font = font.Font(family="Arial", size=8)
##            submit2a= StringVar()
##            submit2a.trace("w", lambda name, index, mode, submi2ta=submit2a: callback(submit2a))
##            submit2b = Button(name, text="Imaging", font=default_font, command=copytoclip_individual_case)
##            submit2b.grid(row=7, column=5, columnspan=2, sticky=EW)
##
##            default_font = font.Font(family="Arial", size=8)
##            submit2a= StringVar()
##            submit2a.trace("w", lambda name, index, mode, submi2ta=submit2a: callback(submit2a))
##            submit2b = Button(name, text="Post-imaging", font=default_font, command=copytoclip_individual_case)
##            submit2b.grid(row=7, column=7, columnspan=2, sticky=EW)
##
##            default_font = font.Font(family="Arial", size=8)
##            submit2a= StringVar()
##            submit2a.trace("w", lambda name, index, mode, submi2ta=submit2a: callback(submit2a))
##            submit2b = Button(name, text="Exhibit seal", font=default_font, command=copytoclip_individual_case)
##            submit2b.grid(row=7, column=9, columnspan=2, sticky=EW)

            Indiv_Case_note_label=Label(name, text="Notes")
            Indiv_Case_note_label.grid(row=6, column=1, sticky=W)
            individualcasevariables.append(Indiv_Case_note_label)
            Loop_Indiv_Action6 = (str("t"+(str(count_no2))))
            print(Loop_Indiv_Action6)
            key=(exhib[4])
            Loop_Indiv_Action6scrollbar = Scrollbar(name)
            Loop_Indiv_Action6=Text(name, wrap=WORD, width=87, yscrollcommand=Loop_Indiv_Action6scrollbar.set)
            print(Loop_Indiv_Action6)
            Loop_Indiv_Action6.grid(row=7, column=1, columnspan=8, sticky=W, rowspan=1)
            print(Loop_Indiv_Action6)
            value=Loop_Indiv_Action6
            
            Loop_Indiv_Action6_notes = (str.replace(exhib[14],"\\n","\n"))
            Loop_Indiv_Action6.insert(INSERT, Loop_Indiv_Action6_notes)

            Loop_Indiv_Action6.grid_propagate(False)
            Loop_Indiv_Action6.bind('<KeyRelease>', Loop_Indiv_Action_exhibit_notes)
            Loop_Indiv_Action6scrollbar.config(command = Loop_Indiv_Action6.yview)
            Loop_Indiv_Action6scrollbar.grid(row=7, column=9, columnspan=1, rowspan=1,  padx= 5, sticky='NS')

            ryan[key] = value

        ## overview ##
            
        tabs.pack(fill='both', expand=Y)
        for key in ryan:
            print('key: ',key)
        print(ryan.values())
        
        # Set row height here
        for row in range(1):
            for col in range(40):
                tk.Label(overview,text='', width=1, height=1).grid(row=row, column=col)
        for col in range(1):
            for row in range(50):
                tk.Label(overview,text='', width=1, height=1).grid(row=row, column=col)
                
        Indiv_Case_note_label=Label(overview, text="Notes")
        Indiv_Case_note_label.grid(row=7, column=1, columnspan=4, sticky=W)
        individualcasevariables.append(Indiv_Case_note_label)

        oic_label=Label(overview, text="DFT Ref")
        oic_label.grid(row=0, column=1, sticky=W)
        individualcasevariables.append(oic_label)

        oic_label=Label(overview, text="Niche/Crime Ref")
        oic_label.grid(row=0, column=2, sticky=W)
        individualcasevariables.append(oic_label)

        oic_label=Label(overview, text="Suspect")
        oic_label.grid(row=0, column=3, sticky=W)
        individualcasevariables.append(oic_label)
        
        oic_label=Label(overview, text="# of exhibits")
        oic_label.grid(row=0, column=5, sticky=W)
        individualcasevariables.append(oic_label)
        
        oic_label=Label(overview, text="OIC")
        oic_label.grid(row=0, column=4, sticky=W)
        individualcasevariables.append(oic_label)

        oic_label=Label(overview, text="Date allocated")
        oic_label.grid(row=2, column=1, columnspan=4, sticky=W)
        individualcasevariables.append(oic_label)

        oic_label=Label(overview, text="Date started")
        oic_label.grid(row=2, column=2, sticky=W)
        individualcasevariables.append(oic_label)

        oic_label=Label(overview, text="Days open")
        oic_label.grid(row=2, column=4, sticky=W)
        individualcasevariables.append(oic_label)

        oic_label=Label(overview, text="Date finished")
        oic_label.grid(row=2, column=3, sticky=W)
        individualcasevariables.append(oic_label)

        case_password=Label(overview, text="Case Password: ")
        case_password.grid(row=5, column=1, sticky=W)


        conn.close()
        


        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute("SELECT * FROM opencases WHERE full_dft_ref like '%'||?||'%'", (d,))
        #c.execute("SELECT * FROM opencases WHERE full_dft_ref like '%'||?||'%'", (d,))

        for each in c:
            print(each)
        count = 1   
        row = 1

        print('each1: ', each[1])

        Indiv_Case = each[1]
        #DFT REF #
        Indiv_Case1a = StringVar()
        Indiv_Case1a.set(each[1])
        Indiv_Case1a.trace("w", lambda name, index, mode, Indiv_Case1a=Indiv_Case1a: callback(Indiv_Case1a))
        Indiv_Case1 = Entry(overview, relief=SUNKEN,width=10, state='normal', textvariable=Indiv_Case1a)
        Indiv_Case1.grid(row=row, column=1, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case1)

        #NICHE REF#
        count += 1
        Indiv_Case2 = (str("e"+(str(count))))
        Indiv_Case2a = StringVar()
        Indiv_Case2a.set(each[2])
        Indiv_Case2a.trace("w", lambda name, index, mode, Indiv_Case2a=Indiv_Case2a: callback_Indiv_Case2(Indiv_Case2a))
        Indiv_Case2 = Entry(overview, relief=SUNKEN,  width=15, state='normal', textvariable=Indiv_Case2a)
        Indiv_Case2.grid(row=row, column=2, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case2)
        count += 1

        #SUSPECT#
        Indiv_Case9 = (str("e"+(str(count))))
        Indiv_Case9a = StringVar()
        Indiv_Case9a.set(each[8])
        Indiv_Case9a.trace("w", lambda name, index, mode, Indiv_Case9a=Indiv_Case9a: callback(Indiv_Case9a))
        Indiv_Case9 = Entry(overview, relief=SUNKEN, state='normal', textvariable=Indiv_Case9a)
        Indiv_Case9.grid(row=row, column=3, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case9)

        # OIC
        Indiv_Case11 = (str("e"+(str(count))))
        Indiv_Case11a = StringVar()
        Indiv_Case11a.set('oic')#each[4])
        Indiv_Case11a.trace("w", lambda name, index, mode, Indiv_Case11a=Indiv_Case11a: callback_Indiv_Case11(Indiv_Case11a))
        Indiv_Case11 = Entry(overview, relief=SUNKEN,width=20, state='normal', textvariable=Indiv_Case11a)
        Indiv_Case11.grid(row=row, column=4, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case11)
        count += 1

        # NO OF EXHIBITS
        Indiv_Case3 = (str("e"+(str(count))))
        Indiv_Case3a = StringVar()
        Indiv_Case3a.set((len(the_case)))
        Indiv_Case3a.trace("w", lambda name, index, mode, Indiv_Case3a=Indiv_Case3a: callback_Indiv_Case3(Indiv_Case3a))
        Indiv_Case3 = Entry(overview, relief=SUNKEN,width=15, state='normal', textvariable=Indiv_Case3a)
        Indiv_Case3.grid(row=row, column=5, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case3)
        count += 1
        row += 2

        #date allocated
        Indiv_Case10 = (str("e"+(str(count))))
        Indiv_Case10a = StringVar()
        Indiv_Case10a.set('1 05/04/1980')#each[13])
        Indiv_Case10a.trace("w", lambda name, index, mode, Indiv_Case10a=Indiv_Case10a: callback_Indiv_Case10(Indiv_Case10a))
        Indiv_Case10 = Entry(overview, relief=SUNKEN, width=15, state='disabled', textvariable=Indiv_Case10a)
        Indiv_Case10.grid(row=row, column=1, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case10)
        count += 1

        #date started
        Indiv_Case5 = (str("e"+(str(count))))
        Indiv_Case5a = StringVar()
        Indiv_Case5a.set('2 05/04/1980')#each[13])
        Indiv_Case5a.trace("w", lambda name, index, mode, Indiv_Case5a=Indiv_Case5a: callback_Indiv_Case5(Indiv_Case5a))
        Indiv_Case5 = Entry(overview, relief=SUNKEN,width=15, state='disabled', textvariable=Indiv_Case5a)
        Indiv_Case5.grid(row=row, column=2, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case5)
        count += 1

        # Date finished
        Indiv_Case4 = (str("e"+(str(count))))
        Indiv_Case4a = StringVar()
        Indiv_Case4a.set('Indiv_Case4')#each[4])
        Indiv_Case4a.trace("w", lambda name, index, mode, Indiv_Case4a=Indiv_Case4a: callback_Indiv_Case4(Indiv_Case4a))
        Indiv_Case4 = Entry(overview, relief=SUNKEN,width=15, state='normal', textvariable=Indiv_Case4a)
        Indiv_Case4.grid(row=row, column=3, sticky=EW)
        individualcasevariables.append(Indiv_Case4)
        count += 1
        selected_month_rec = (timestamp[:-9])
        
        #timestamp = (each[13])
        selected_month_rec = (timestamp[:-9])

        #selected_month_rec = (each[13])
        print(selected_month_rec)
        start = date(int(selected_month_rec.split('-')[0]),int(selected_month_rec.split('-')[1]),int(selected_month_rec.split('-')[2]))
        today = date.today()
        res = today - start
        res.days

        #days open
        Indiv_Case7 = (str("e"+(str(count))))
        Indiv_Case7a = StringVar()
        Indiv_Case7a.set(res.days)
        Indiv_Case7a.trace("w", lambda name, index, mode, Indiv_Case7a=Indiv_Case7a: callback(Indiv_Case7a))
        Indiv_Case7 = Entry(overview, relief=SUNKEN,width=20, state='disabled', textvariable=Indiv_Case7a)
        Indiv_Case7.grid(row=row, column=4, columnspan=1, sticky=EW)
        individualcasevariables.append(Indiv_Case7)
        count += 1
        row += 2
        #case password
        Indiv_Case8 = (str("e"+(str(count))))
        Indiv_Case8a = StringVar()
        Indiv_Case8a.set(case_pw)
        Indiv_Case8a.trace("w", lambda name, index, mode, Indiv_Case8a=Indiv_Case8a: callback(Indiv_Case8a))
        Indiv_Case8 = Entry(overview, relief=SUNKEN,  width=30, state='disabled', textvariable=Indiv_Case8a)
        Indiv_Case8.grid(row=row, column=2, columnspan=2, sticky=EW)
        individualcasevariables.append(Indiv_Case8)

        default_font = font.Font(family="Arial", size=8)
        # copy to clip button
        submit2a= StringVar()
        submit2a.trace("w", lambda name, index, mode, submi2ta=submit2a: callback(submit2a))
        submit2b = Button(overview, text="Copy to clip", font=default_font, command=copytoclip_individual_case)
        submit2b.grid(row=row, column=4, columnspan=1, sticky=EW)
        row += 3

        #notes &scrollbar
        Indiv_Case6scrollbar = Scrollbar(overview)
        Indiv_Case6=Text(overview, wrap=WORD,  width=87, yscrollcommand=Indiv_Case6scrollbar.set)
        Indiv_Case6.grid(row=row, column=1, columnspan=7,  rowspan=3, sticky=W)

        Indiv_Case6_notes = (str.replace(each[6],"\\n","\n"))
        Indiv_Case6.insert(INSERT, Indiv_Case6_notes)

        Indiv_Case6.grid_propagate(False)
        Indiv_Case6.bind('<KeyRelease>', Indiv_Case_notes)
        Indiv_Case6scrollbar.config( command = Indiv_Case6.yview)
        Indiv_Case6scrollbar.grid(row=row, column=8,  columnspan=1, rowspan=3, padx= 5, sticky='NS')
        
        individualcasevariables.append(Indiv_Case6scrollbar)
        individualcasevariables.append(Indiv_Case6)

        count += 1

##        del_close1=Button(overview, text="Delete/Close", width=9, command=close)
##        del_close1.grid(row=13, column=5, columnspan=1, sticky=W)
##        del_close1.config(state='normal')
##        individualcasevariables.append(del_close1)

        count += 1
        Indiv_Case6.focus()
        case.update()
        name.update()
        case.attributes('-topmost',True)
        case.after_idle(case.attributes,'-topmost',False)
        case.mainloop()
##############################################################################################
def get_individual_pd_data():
    global d, x, pd_case, Indiv_pd_case, Indiv_pd_case1, Indiv_pd_case2, Indiv_pd_case3, Indiv_pd_case4, Indiv_pd_case5, Indiv_pd_case6, Indiv_pd_case7
    
    d = TEXT
    if d == "No entries":
        pass
    else:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute("SELECT * FROM pindecryptlog WHERE full_dft_ref like '%'||?||'%'", (d,))
        try:
            if 'normal' == pd_case.state():
                pd_case.destroy()
            elif 'normal' == case.state():
                case.destroy()
            else:
                pass
        except:
            pass

        pd_case = Toplevel()
        width  = root.winfo_screenwidth()
        height = root.winfo_screenheight()
        pd_case.geometry('720x680') #widthxheight
        print(width)
        print(height)
        pd_case.title("PIN decryption: "+d)
        pd_case.resizable(width=False, height=False)
        pd_case.protocol("WM_DELETE_WINDOW", close_individual_pd_case)
        
        Indiv_pd_case_note_label=Label(pd_case, text=" Notes")
        Indiv_pd_case_note_label.grid(row=5, column=1, sticky=W)

        oic_label=Label(pd_case, text=" DFT Ref")
        oic_label.grid(row=1, column=1, sticky=W)
    
        oic_label=Label(pd_case, text=" Niche/Crime Ref")
        oic_label.grid(row=1, column=2, sticky=W)

        oic_label=Label(pd_case, text=" Exhibit Ref")
        oic_label.grid(row=1, column=3, sticky=W)
        
        oic_label=Label(pd_case, text=" OIC")
        oic_label.grid(row=3, column=3, sticky=W)

        oic_label=Label(pd_case, text=" Date started")
        oic_label.grid(row=3, column=1, sticky=W)

        oic_label=Label(pd_case, text=" Days open")
        oic_label.grid(row=3, column=2, sticky=W)
        frames=["pd_case"] 
        rows=31
        columns=7
        # Set row height here
        height=5
        for frame in frames:
            count=0
            count2=0
            frame=(str(frame))
            if frame == frames[0]:
                while count < rows:
                    Label(pd_case, text=" ", width=2).grid(row=count, column=0, sticky=EW, ipady=height)
                    count+=1
                while count2 < columns:
                    Label(pd_case, text=" ", width=2).grid(row=0, column=count2, sticky=EW, ipady=height)
                    count2+=1
        for each in c:
            print(each)
            
        count = 1   
        row = 2

        Indiv_pd_case = each[1]

        Indiv_pd_case1a = StringVar()
        Indiv_pd_case1a.set(each[1])
        Indiv_pd_case1a.trace("w", lambda name, index, mode, Indiv_pd_case1a=Indiv_pd_case1a: callback(Indiv_pd_case1a))
        Indiv_pd_case1 = Entry(pd_case, relief=SUNKEN,width=10, state='disabled', textvariable=Indiv_pd_case1a)
        Indiv_pd_case1.grid(row=row, column=1, sticky=EW)

        count += 1
        Indiv_pd_case2 = (str("e"+(str(count))))
        Indiv_pd_case2a = StringVar()
        Indiv_pd_case2a.set(each[2])
        Indiv_pd_case2a.trace("w", lambda name, index, mode, Indiv_pd_case2a=Indiv_pd_case2a: callback_Indiv_pd_case2(Indiv_pd_case2a))
        Indiv_pd_case2 = Entry(pd_case, relief=SUNKEN,width=15, state='normal', textvariable=Indiv_pd_case2a)
        Indiv_pd_case2.grid(row=row, column=2, sticky=EW)

        count += 1
        Indiv_pd_case3 = (str("e"+(str(count))))
        Indiv_pd_case3a = StringVar()
        Indiv_pd_case3a.set(each[3])
        Indiv_pd_case3a.trace("w", lambda name, index, mode, Indiv_pd_case3a=Indiv_pd_case3a: callback_Indiv_pd_case3(Indiv_pd_case3a))
        Indiv_pd_case3 = Entry(pd_case, relief=SUNKEN,width=35, state='normal', textvariable=Indiv_pd_case3a)
        Indiv_pd_case3.grid(row=row, column=3, columnspan=2, sticky=W)
        x=Indiv_pd_case3a.get()
        
        count += 1
        row += 2

        Indiv_pd_case4 = (str("e"+(str(count))))
        Indiv_pd_case4a = StringVar()
        Indiv_pd_case4a.set(each[4])
        Indiv_pd_case4a.trace("w", lambda name, index, mode, Indiv_pd_case4a=Indiv_pd_case4a: callback_Indiv_pd_case4(Indiv_pd_case4a))
        Indiv_pd_case4 = Entry(pd_case, relief=SUNKEN,width=35, state='normal', textvariable=Indiv_pd_case4a)
        Indiv_pd_case4.grid(row=row, column=3, sticky=W)

        count += 1

        Indiv_pd_case5 = (str("e"+(str(count))))
        Indiv_pd_case5a = StringVar()
        Indiv_pd_case5a.set(each[5])
        Indiv_pd_case5a.trace("w", lambda name, index, mode, Indiv_pd_case5a=Indiv_pd_case5a: callback_Indiv_pd_case5(Indiv_pd_case5a))
        Indiv_pd_case5 = Entry(pd_case, relief=SUNKEN,width=25, state='disabled', textvariable=Indiv_pd_case5a)
        Indiv_pd_case5.grid(row=row, column=1, columnspan=1, sticky=W)

        count += 1

        selected_month_rec = (each[5])
        start = date(int(selected_month_rec.split('-')[0]),int(selected_month_rec.split('-')[1]),int(selected_month_rec.split('-')[2]))
        today = date.today()
        res = today - start
        res.days

        Indiv_pd_case7 = (str("e"+(str(count))))
        Indiv_pd_case7a = StringVar()
        Indiv_pd_case7a.set(res.days)
        Indiv_pd_case7a.trace("w", lambda name, index, mode, Indiv_pd_case7a=Indiv_pd_case7a: callback(Indiv_pd_case7a))
        Indiv_pd_case7 = Entry(pd_case, relief=SUNKEN,width=25, state='disabled', textvariable=Indiv_pd_case7a)
        Indiv_pd_case7.grid(row=row, column=2, columnspan=1, sticky=W)

        count += 1
        row += 2

        Indiv_pd_case6scrollbar = Scrollbar(pd_case)
        Indiv_pd_case6=Text(pd_case, wrap=WORD, width=130, yscrollcommand=Indiv_pd_case6scrollbar.set)
        Indiv_pd_case6.grid(row=6, column=1, columnspan=8, sticky=W, rowspan=13)

        Indiv_pd_case6_notes = (str.replace(each[8],"\\n","\n"))
        Indiv_pd_case6.insert(INSERT, Indiv_pd_case6_notes)

        Indiv_pd_case6.grid_propagate(False)
        Indiv_pd_case6.bind('<KeyRelease>', Indiv_pd_notes)
        Indiv_pd_case6scrollbar.config( command = Indiv_pd_case6.yview)
        Indiv_pd_case6scrollbar.grid(row=6, column=6, rowspan=3,  sticky='NS')


        #        Label(pd_case, text=" Status:").grid(row=row, column=1, sticky=W)
        status_dda = StringVar()
        status_dda.set('--Status--')
        status_dda.trace("w", lambda name, index, mode, status_dda=status_dda: callback(status_dda))
        status_ddb = OptionMenu(pd_case, status_dda, *status, command=qwerty)
        status_ddb.grid(row=19, column=1, columnspan=2, sticky=EW)
        individual_pd_casevariables.append(status_ddb)

        count += 1

        del_close1=Button(pd_case, text="Delete/Close", width=17, command=close_pd)
        del_close1.grid(row=19, column=4, sticky=W)
        del_close1.config(state='normal')
        individual_pd_casevariables.append(del_close1)

        count += 1

        pd_case.update
        c.close()
        pd_case.mainloop()
############################################################################################
def callback_Indiv_pd2(Indiv_Case2a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_pd_case2.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET crime_ref = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_pd3(Indiv_Case3a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_pd_case3.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET exhibit = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_pd4(Indiv_Case4a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_pd_case4.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET oic = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update
############################################################################################
def callback_Indiv_pd5(Indiv_Case5a):
    sqlitedbs = (sqlprolocal, sqlprolocal)
    try:
        for sql in sqlitedbs:
            get=Indiv_pd_case5.get()
            
            conn = sqlite3.connect(sqlprolocal, isolation_level=None)
            conn.execute('pragma journal_mode=wal')
            c=conn.cursor()
            c.execute("""UPDATE pindecryptlog SET datestamp = ? WHERE full_dft_ref like '%'||?||'%'""", (get, Indiv_Case,))
            '''c.execute("""UPDATE opencases SET crime_ref = ?,
                      no_of_exhib = ?,
                      oic = ?,
                      datestamp = ?,
                      note = ? WHERE full_dft_ref like '%'||?||'%'""", (z, y, w, q, p, d,))'''
            conn.commit()
            conn.close()
    except:
        Indiv_Case2.config(state='disabled')
        Indiv_Case3.config(state='disabled')
        Indiv_Case4.config(state='disabled')
        Indiv_Case5.config(state='disabled')
        Indiv_Case6.config(state='disabled')
        startup.update

############################################################################################
def qwerty(sv):
    global d, x
    d = TEXT
    if d == "No entries":
        pass
    else:
        x = d.split(' - ')
        
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute("SELECT * FROM pindecryptlog WHERE full_dft_ref like '%'||?||'%' and exhibit = ?", (x[0], x[1],))
        for each in c:
            print(each)
            
            a1a.set(each[2])
            a2a.set(each[3])
            a3a.set(each[4])
            a4a.set(each[5])
            a5a.set('YYYY-MM-DD')
            a6a.set(each[8])
        edit.update()
        c.close()
############################################################################################
def qwerty1(sv):
    global d, x
    d = (edit2a.get())
    if d == "No entries":
        pass
    else:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute("SELECT * FROM opencases WHERE full_dft_ref like '%'||?||'%'", (d,))
        for each in c:
            print(each)
            a1a.set(each[2])
            a2a.set(each[3])
            a3a.set(each[4])
            a4a.set(each[5])
            a5a.set(each[6])
        edit1.update()
        c.close()
############################################################################################
def qwerty2(sv):
    global d, x
    d = (edit3a.get())
    if d == "No entries":
        pass
    else:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute("SELECT * FROM assignedcases WHERE full_dft_ref like '%'||?||'%'", (d,))
        for each in c:
            print(each)
            a1a.set(each[2])
            a2a.set(each[3])
            a3a.set(each[4])
            a4a.set(each[5])
            a5a.set(each[6])
        edit2.update()
        c.close()
############################################################################################
def edit_entry():
    global edita, a1a, a2a, a3a, a4a, a5a, a6a, editb, edit, options

    options=[]
    for each in pindecryptlogs:
        options.append(str(each[1]+' - '+each[3]))
        
    gui = ("Edit")
    edit = Toplevel()
    edit.title("Edit")
    edit.resizable(width=False, height=False)
    edit.lift(aboveThis=root)

    frames=["edit"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(edit, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    noentries="False"
    if (len(options))==0:
        options=('No entries',)
        noentries = "True"
        
    Label(edit, text=" DFT Ref/Exhibit").grid(row=0, column=1, sticky=W)
    edita = StringVar()
    edita.set('--Select--')
    edita.trace("w", lambda name, index, mode, edita=edita: callback(edita))
    editb = OptionMenu(edit, edita, *options, command=qwerty)
    editb.grid(row=1, column=1, columnspan=1, sticky=EW)

    def cut_from_options():
        options['menu'].delete(0, 'end')
        for each in pindecryptlogs:
            options.add_command(label=each, command=tk._setit(edita, (str(each[1]+' - '+each[3]))))

    Label(edit, text=" Crime Ref").grid(row=0, column=2, sticky=W)
    a1a = StringVar()
    a1a.trace("w", lambda name, index, mode, a1a=a1a: callback1(a1a))
    a1b = Entry(edit,relief=SUNKEN,width=20, textvariable=a1a)
    a1b.grid(row=1, column=2)

    Label(edit, text=" Exhibit Ref").grid(row=0, column=3, sticky=W)
    a2a = StringVar()
    a2a.trace("w", lambda name, index, mode, a2a=a2a: callback1(a2a))
    a2b = Entry(edit,relief=SUNKEN,width=20, textvariable=a2a)
    a2b.grid(row=1, column=3)

    Label(edit, text=" OIC/Contact").grid(row=0, column=4, sticky=W)
    a3a = StringVar()
    a3a.trace("w", lambda name, index, mode, a3a=a3a: callback1(a3a))
    a3b = Entry(edit,relief=SUNKEN,width=20, textvariable=a3a)
    a3b.grid(row=1, column=4)

    Label(edit, text=" Date started").grid(row=0, column=5, sticky=W)
    a4a = StringVar()
    a4a.trace("w", lambda name, index, mode, a4a=a4a: callback1(a4a))
    a4b = Entry(edit,relief=SUNKEN,width=20, textvariable=a4a)
    a4b.grid(row=1, column=5)
    
    Label(edit, text=" Days running").grid(row=0, column=6, sticky=W)
    a5a = StringVar()
    a5a.trace("w", lambda name, index, mode, a5a=a5a: callback1(a5a))
    a5b = Entry(edit,relief=SUNKEN,width=10, textvariable=a5a)
    a5b.grid(row=1, column=6)
    a5b.config(state='disabled')

    Label(edit, text=" Notes").grid(row=0, column=7, sticky=W)
    a6a = StringVar()
    a6a.trace("w", lambda name, index, mode, a6a=a6a: callback1(a6a))
    a6b = Entry(edit,relief=SUNKEN,width=80, textvariable=a6a)
    a6b.grid(row=2, column=2, columnspan=8, sticky=EW)


    Label(edit, text=" ",width=3).grid(row=1, column=10, columnspan=1, sticky=EW)
    Label(edit, text=" ",width=20).grid(row=3, column=1, columnspan=1, sticky=EW)
    
    if noentries == 'True':
        del_close2=Button(edit, text="Delete/Close", width=17, command=pdresult)
        del_close2.grid(row=1, column=8, sticky=W)
        del_close2.config(state='disable')
        save2 = Button(edit, text="Save changes", width=17, command=updatepinentry)
        save2.grid(row=1, column=9, sticky=W)
        save2.config(state='disable')
    else:
        del_close2=Button(edit, text="Delete/Close", width=17, command=pdresult)
        del_close2.grid(row=1, column=8, sticky=W)
        del_close2.config(state='normal')
        save2 = Button(edit, text="Save changes", width=17, command=updatepinentry)
        save2.grid(row=1, column=9, sticky=W)
        save2.config(state='normal')

    edit.lift()
    edit.attributes('-topmost',True)
    edit.after_idle(root.attributes,'-topmost',False)
    edit.protocol("WM_DELETE_WINDOW", enable_edit1)
    edit.mainloop()
    
############################################################################################
def edit_opencase_entry():
    global save1, edit2a, a1a, a2a, a3a, a4a, a5a, editb, edit1, options

    options=[]
    for each in opencases:
        options.append(each[1])
        
    gui = ("Edit")
    edit1 = Toplevel()
    edit1.title("Edit")
    edit1.resizable(width=False, height=False)
    edit1.lift(aboveThis=root)

    frames=["edit"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(edit1, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    '''if (len(options))==0:
        options=('No entries',)'''
        
    Label(edit1, text=" DFT Ref/Exhibit").grid(row=0, column=1, sticky=W)
    edit2a = StringVar()
    edit2a.set('--Select--')
    edit2a.trace("w", lambda name, index, mode, edit2a=edit2a: callback(edit2a))
    edit2b = OptionMenu(edit1, edit2a, *options, command=qwerty1)
    edit2b.grid(row=1, column=1, columnspan=1, sticky=EW)

    def cut_from_options():
        options['menu'].delete(0, 'end')
        for each in opencases:
            options.add_command(label=each, command=tk._setit(edita, (str(each[1]+' - '+each[3]))))

    Label(edit1, text=" Crime Ref").grid(row=0, column=2, sticky=W)
    a1a = StringVar()
    a1a.trace("w", lambda name, index, mode, a1a=a1a: callback1(a1a))
    a1b = Entry(edit1,relief=SUNKEN,width=20, textvariable=a1a)
    a1b.grid(row=1, column=2)

    Label(edit1, text=" No of exhibits").grid(row=0, column=3, sticky=W)
    a2a = StringVar()
    a2a.trace("w", lambda name, index, mode, a2a=a2a: callback1(a2a))
    a2b = Entry(edit1,relief=SUNKEN,width=20, textvariable=a2a)
    a2b.grid(row=1, column=3)

    Label(edit1, text=" OIC/Contact").grid(row=0, column=4, sticky=W)
    a3a = StringVar()
    a3a.trace("w", lambda name, index, mode, a3a=a3a: callback1(a3a))
    a3b = Entry(edit1,relief=SUNKEN,width=20, textvariable=a3a)
    a3b.grid(row=1, column=4)

    Label(edit1, text=" Date started").grid(row=0, column=5, sticky=W)
    a4a = StringVar()
    a4a.trace("w", lambda name, index, mode, a4a=a4a: callback1(a4a))
    a4b = Entry(edit1,relief=SUNKEN,width=20, textvariable=a4a)
    a4b.grid(row=1, column=5)
    
    Label(edit1, text=" Notes").grid(row=2, column=1, sticky=E)
    a5a = StringVar()
    a5a.trace("w", lambda name, index, mode, a5a=a5a: callback1(a5a))
    a5b = Entry(edit1,relief=SUNKEN,width=80, textvariable=a5a)
    a5b.grid(row=2, column=2, columnspan=8, sticky=EW)

    Label(edit1, text=" ",width=3).grid(row=1, column=10, columnspan=1, sticky=EW)
    Label(edit1, text=" ",width=20).grid(row=3, column=1, columnspan=1, sticky=EW)

    del_close1 = Button(edit1, text="Delete/Close", width=17, command=deleteentry2).grid(row=1, column=8, sticky=W)
    
    save1=Button(edit1, text="Save changes", width=17, command=updateopencaseentry)
    save1.grid(row=1, column=9, sticky=W)
    save1.config(command=updateopencaseentry)

    edit1.lift()
    edit1.attributes('-topmost',True)
    edit1.after_idle(root.attributes,'-topmost',False)
    edit1.protocol("WM_DELETE_WINDOW", enable_edit2)
    edit1.mainloop()
############################################################################################
def edit_assignedcase_entry():
    global save1, edit3a, a1a, a2a, a3a, a4a, a5a, editb, edit2, assigned, start

    assigned=[]
    for each in assignedcases:
        assigned.append(each[1])
        print(each)
        
    gui = ("Edit")
    edit2 = Toplevel()
    edit2.title("Edit")
    edit2.resizable(width=False, height=False)
    edit2.lift(aboveThis=root)

    frames=["edit"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(edit2, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    '''if (len(options))==0:
        options=('No entries',)'''
        
    Label(edit2, text=" DFT Ref").grid(row=0, column=1, sticky=W)
    edit3a = StringVar()
    edit3a.set('--Select--')
    edit3a.trace("w", lambda name, index, mode, edit3a=edit3a: callback(edit3a))
    edit3b = OptionMenu(edit2, edit3a, *assigned, command=qwerty2)
    edit3b.grid(row=1, column=1, columnspan=1, sticky=EW)

    def cut_from_options():
        options['menu'].delete(0, 'end')
        for each in opencases:
            options.add_command(label=each, command=tk._setit(edita, (str(each[1]+' - '+each[3]))))

    Label(edit2, text=" Crime Ref").grid(row=0, column=2, sticky=W)
    a1a = StringVar()
    a1a.trace("w", lambda name, index, mode, a1a=a1a: callback1(a1a))
    a1b = Entry(edit2,relief=SUNKEN,width=20, textvariable=a1a)
    a1b.grid(row=1, column=2)

    Label(edit2, text=" No of exhibits").grid(row=0, column=3, sticky=W)
    a2a = StringVar()
    a2a.trace("w", lambda name, index, mode, a2a=a2a: callback1(a2a))
    a2b = Entry(edit2,relief=SUNKEN,width=20, textvariable=a2a)
    a2b.grid(row=1, column=3)

    Label(edit2, text=" OIC/Contact").grid(row=0, column=4, sticky=W)
    a3a = StringVar()
    a3a.trace("w", lambda name, index, mode, a3a=a3a: callback1(a3a))
    a3b = Entry(edit2,relief=SUNKEN,width=20, textvariable=a3a)
    a3b.grid(row=1, column=4)

    Label(edit2, text=" Date Assigned").grid(row=0, column=5, sticky=W)
    a4a = StringVar()
    a4a.trace("w", lambda name, index, mode, a4a=a4a: callback1(a4a))
    a4b = Entry(edit2,relief=SUNKEN,width=20, textvariable=a4a)
    a4b.grid(row=1, column=5)
    
    Label(edit2, text=" Notes").grid(row=2, column=1, sticky=E)
    a5a = StringVar()
    a5a.trace("w", lambda name, index, mode, a5a=a5a: callback1(a5a))
    a5b = Entry(edit2,relief=SUNKEN,width=50, textvariable=a5a)
    a5b.grid(row=2, column=2, columnspan=7, sticky=EW)

    Label(edit2, text=" ",width=3).grid(row=1, column=10, columnspan=1, sticky=EW)
    Label(edit2, text=" ",width=20).grid(row=3, column=1, columnspan=1, sticky=EW)

    del_close2 = Button(edit2, text="Delete/Clear", width=17, command=deleteentry3).grid(row=1, column=8, sticky=W)
    
    save2=Button(edit2, text="Save changes", width=17, command=updateassignedcaseentry)
    save2.grid(row=1, column=9, sticky=W)
    save2.config(command=updateassignedcaseentry)

    start=Button(edit2, text="Start", width=17, command=start_assigned)
    start.grid(row=2, column=9, sticky=W)

    edit2.lift()
    edit2.attributes('-topmost',True)
    edit2.after_idle(root.attributes,'-topmost',False)
    edit2.protocol("WM_DELETE_WINDOW", enable_edit3)
    edit2.mainloop()
    
############################################################################################
def callback1_pd_entry(sv):
    pd_dft=(str(Newentry1a.get().upper()))
    print(sv.get())
    pd_dft=(pd_dft[:7])
    if len(pd_dft) == 3:
        if pd_dft[-1:3] != '-':
            Newentry1a.set(pd_dft[:2])
            pd_dft=(pd_dft[:2])
        else:
            Newentry1a.set(pd_dft[:7])
    elif len(pd_dft) >= 7:
        Newentry1a.set(pd_dft[:7])
        pd_dftrefcheck()
    else:
        Newentry1a.set(pd_dft[:7])
        Newentry1b.config({"background": "white"})
    add.update() 
############################################################################################
def callback2_pd_entry(sv):
    pd_exhibit=(str(Newentry3a.get().upper()))
    print(sv.get())
    if len(pd_exhibit) >= 2:
        pd_exhibitcheck()
    else:
        Newentry3b.config({"background": "white"})
############################################################################################
def callback3_pd_entry(sv):
    pd_crime=(str(Newentry2a.get().upper()))
    print(sv.get())
    if len(pd_crime) >= 11:
        pd_crimecheck()
    else:
        Newentry2b.config({"background": "white"})
############################################################################################
def pd_dftrefcheck():
    Newentry1b.config({"background": "pale green"})
    add.update()
############################################################################################
def pd_exhibitcheck():
    Newentry3b.config({"background": "pale green"})
    add.update()
############################################################################################
def pd_crimecheck():
    Newentry2b.config({"background": "pale green"})
    add.update()
############################################################################################
def add_new_PD_entry():
    global add, save1, edit3a, a1a, a2a, a3a, a4a, a5a, editb, edit2, assigned, start, Newentry0a, Newentry1a, Newentry2a, Newentry3a, Newentry4a, Newentry5a, Newentry6a, Newentry0b, Newentry1b, Newentry2b, Newentry3b, Newentry4b, Newentry5b, Newentry6b

    assigned=[]
    for each in assignedcases:
        assigned.append(each[1])
        print(each)
        
    gui = ("Add")
    add = Toplevel()
    add.title("Add PIN decryption Log entry")
    add.resizable(width=False, height=False)
    add.lift(aboveThis=root)

    frames=["add"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(add, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    addnewrow = StringVar()
    addnewrow.trace("w", lambda name, index, mode, addnewrow=addnewrow: callback(addnewrow))
    addnewrowb = Button(add, text="+Add New", width=10, command=process_add_new_pd_entry)
    addnewrowb.grid(row=4, column=3, sticky=EW)
    addnewrowb.config(state='normal')
    
    Label(add, text=" DFT Ref").grid(row=1, column=1, columnspan=1, sticky=W)
    Newentry1a = StringVar()
    Newentry1a.trace("w", lambda name, index, mode, Newentry1a=Newentry1a: callback1_pd_entry(Newentry1a))
    Newentry1b = Entry(add, relief=SUNKEN,width=35, textvariable=Newentry1a)
    Newentry1b.grid(row=2, column=1, sticky=W)

    Label(add, text=" Exhibit Ref").grid(row=1, column=2, columnspan=1, sticky=W)
    Newentry3a = StringVar()
    Newentry3a.trace("w", lambda name, index, mode, Newentry3a=Newentry3a: callback2_pd_entry(Newentry3a))
    Newentry3b = Entry(add, relief=SUNKEN, width=35, textvariable=Newentry3a)
    Newentry3b.grid(row=2, column=2, sticky=W)

    Label(add, text=" Niche Ref").grid(row=1, column=3, columnspan=1, sticky=W)
    Newentry2a = StringVar()
    Newentry2a.trace("w", lambda name, index, mode, Newentry2a=Newentry2a: callback3_pd_entry(Newentry2a))
    Newentry2b = Entry(add, relief=SUNKEN, width=25, textvariable=Newentry2a)
    Newentry2b.grid(row=2, column=3, sticky=W)
    
    Label(add, text=" OIC/Contact").grid(row=3, column=1, columnspan=1, sticky=W)
    Newentry4a = StringVar()
    Newentry4a.trace("w", lambda name, index, mode, Newentry4a=Newentry4a: callback(Newentry4a))
    Newentry4b = Entry(add, relief=SUNKEN,width=35, textvariable=Newentry4a)
    Newentry4b.grid(row=4, column=1, sticky=W)

    Label(add, text=" Date started").grid(row=3, column=2, columnspan=1, sticky=W)
    Newentry5a = StringVar()
    Newentry5a.set('YYYY-MM-DD')
    Newentry5a.trace("w", lambda name, index, mode, Newentry5a=Newentry5a: callback(Newentry5a))
    Newentry5b = Entry(add, relief=SUNKEN,width=35, state='disabled', textvariable=Newentry5a)
    Newentry5b.grid(row=4, column=2, sticky=W)
    
    Label(add, text=" ").grid(row=5, column=1, columnspan=1, sticky=W)
    
    '''Newentry6a = StringVar()
    Newentry6a.trace("w", lambda name, index, mode, Newentry6a=Newentry6a: callback(Newentry6a))
    Newentry6b = Entry(add, relief=SUNKEN,width=25, state='disabled', textvariable=Newentry6a)
    Newentry6b.grid(row=4, column=3, sticky=W)'''


    add.lift()
    add.attributes('-topmost',True)
    add.after_idle(root.attributes,'-topmost',False)
    add.protocol("WM_DELETE_WINDOW", addclose)
    add.mainloop()
############################################################################################
def search_pw():
    global search_pw, save1, edit3a, a1a, a2a, a3a, a4a, a5a, editb, edit2, assigned, start, Newentry1a, pw_search_textbox

    assigned=[]
    for each in assignedcases:
        assigned.append(each[1])
        print(each)
        
    gui = ("search_pw")
    search_pw = Toplevel()
    search_pw.title("CASE - Search Case Password(s)")
    search_pw.resizable(width=False, height=False)
    search_pw.lift(aboveThis=root)

    frames=["search_pw"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(search_pw, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1

    
    Label(search_pw, text=" DFT Ref:").grid(row=1, column=1, columnspan=1, sticky=W)
    Newentry1a = StringVar()
    Newentry1a.trace("w", lambda name, index, mode, Newentry1a=Newentry1a: callback_pw_search(Newentry1a))
    Newentry1b = Entry(search_pw, relief=SUNKEN,width=35, textvariable=Newentry1a)
    Newentry1b.grid(row=1, column=2, sticky=EW)
    Newentry1b.focus()
    
    Label(search_pw, text=" ").grid(row=2, column=1, columnspan=1, sticky=W)
    
    ### Notepad tab (scroll WORKING)
    notepad_scrollbar = Scrollbar(search_pw)
    #Label(notepad, text=" Results/Findings: ").grid(row=5, column=0, sticky=W, pady=0, columnspan=6)
    pw_search_textbox=Text(search_pw, wrap=WORD, width=40, height=10, yscrollcommand=notepad_scrollbar.set)
    pw_search_textbox.grid(row=3, column=1, columnspan=2, sticky=EW, rowspan=1)
    print('Len notepadentries:',(len(notepadentries)))

    pw_search_textbox.grid_propagate(False)
    pw_search_textbox.bind('<KeyRelease>', notepadnotes1)
    notepad_scrollbar.config( command = pw_search_textbox.yview)
    notepad_scrollbar.bind('<MouseWheel>', pw_search_textbox)
    notepad_scrollbar.grid(row=3, column=3, rowspan=1,  sticky='NS')
    
    Label(search_pw, text=" ").grid(row=4, column=1, columnspan=1, sticky=W)
    
    search_pw.lift()
    search_pw.attributes('-topmost',True)
    search_pw.after_idle(root.attributes,'-topmost',False)
    search_pw.protocol("WM_DELETE_WINDOW", search_pwclose)
    search_pw.mainloop()
############################################################################################    
def process_add_new_pd_entry():
    global year, dftonly, sus_op_crime
    pd_dft = Newentry1a.get()
    pd_exhibit = Newentry3a.get()
    pd_crime = Newentry2a.get()
    if len(pd_dft)!=7: 
        Newentry1b.config({"background": "orange red"})
        startup.update()
        dftok=0
    else:
        dftok=1
    
    if len(pd_exhibit)<=1: 
        Newentry3b.config({"background": "orange red"})
        startup.update()
        crimeok=0
    else:
        crimeok=1
        
    if len(pd_crime)<=10: 
        Newentry2b.config({"background": "orange red"})
        startup.update()
        exhibok=0
    else:
        exhibok=1

    
    if dftok+crimeok+exhibok==3:
        addline()
    else:
        pass
    
                    
    '''if dftok ==1:
        if crimeok ==1:
            if exhibok==1:
                if len(suspect)<=1: '''
                               
###########################################################################################
def start_assigned_from_dropdown(value):
    e1a.set('')
    e2a.set('')
    e3a.set('')
    e4.delete(first=0,last=100)
    e5a.set('')
    e6a.set('')
    e7a.set('')
    e8.delete(first=0,last=100)
    e9a.set('')
    for each in assignedcases:
        if each[1] == value:
            e1a.set(each[1].upper())
            e2a.set(each[2].upper())
            if (len(each[4]))==0:
                    e5a.set("-")
            else:
                e5a.set(each[4].upper())
############################################################################################    
def start_assigned():
    dft_ref = (edit3a.get().upper())
    e1a.set(dft_ref)
    
    crime_ref = (a1a.get().upper())
    e2a.set(crime_ref)
    
    oic = (a3a.get().upper())
    
    if (len(oic))==0:
            e5a.set("-")
    else:
        e5a.set(oic)
    enable_edit3()
############################################################################################
def expandview():
    viewtype = viewTypea.get()
    if viewtype == 'PIN decryption log':
        showpdlogdata()
    elif viewtype == 'Open cases':
        showopencasedata()
    else:
        showassignedcasedata()
############################################################################################
def notepadnotes1(*args):
    get=notepad_n59b.get(1.0, END)
    print('in func.', get)
    timestamp = '{:%Y-%m-%d %H:%M:%S}'.format(datetime.datetime.now())
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" UPDATE notepad SET notes = ? WHERE _rowid_ = 1""", (get,))
    idno=2
    c.execute(""" UPDATE notepad SET last_written = ? WHERE _rowid_ = 1""", (timestamp,))
    conn.commit()
    conn.close()
    notepad_overview.delete(1.0,END)
    notepad_overview.insert(INSERT, get)
############################################################################################    
def notepad_overview_notes1(*args):
    get=notepad_overview.get(1.0, END)
    print('in func.', get)
    timestamp = '{:%Y-%m-%d %H:%M:%S}'.format(datetime.datetime.now())
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" UPDATE notepad SET notes = ? WHERE _rowid_ = 1""", (get,))
    idno=2
    c.execute(""" UPDATE notepad SET last_written = ? WHERE _rowid_ = 1""", (timestamp,))
    conn.commit()
    conn.close()
    notepad_n59b.delete(1.0,END)
    notepad_n59b.insert(INSERT, get)
############################################################################################
def notifications_1(*args):
    get=notepad_n60b.get(1.0, END)
    print('in func.', get)
    timestamp = '{:%Y-%m-%d %H:%M:%S}'.format(datetime.datetime.now())
    conn = sqlite3.connect(sqlprolocal, isolation_level=None)
    conn.execute('pragma journal_mode=wal')
    c=conn.cursor()
    c.execute(""" UPDATE notepad SET notes = ? WHERE _rowid_ = 1""", (get,))
    idno=2
    c.execute(""" UPDATE notepad SET last_written = ? WHERE _rowid_ = 1""", (timestamp,))
    conn.commit()
    conn.close()

############################################################################################
def clear_notification(*args):
    if (len(assigned_not_notified))==0:
        root.update()
        print('LENGTH of new_notification_variables = 0')
        pass
    else:
        mark_notified=(assigned_not_notified[0])
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        c.execute(" UPDATE assignedcases SET notified = ? WHERE full_dft_ref like '%'||?||'%'", (1, mark_notified,))
        c.execute(" UPDATE assignedcases SET timestamp = ? WHERE full_dft_ref like '%'||?||'%'", (timestamp, mark_notified,))
        conn.commit()
        conn.close()
        entry.destroy()

        assigned_not_notified.remove(mark_notified)
        new_notification_variables[0].destroy()
        
        root.update()
        writenotifications()
        readnotifications()
        update_jackie()
    readassignedcases()
############################################################################################
def notification():
    global e8zb, message,entry, notific1
    if (len(assigned_not_notified))>=2:
        pass
    else:
        message=(str('NOTIFICATION: You have a newly assigned case: '+assigned_not_notified[0]))
        entry = 'notific1' #(str(assigned_not_notified[0]))
        entry = StringVar()
        entry.set(message)
        entry = Entry(startup,relief=SUNKEN,width=30, textvariable=entry)
        entry.grid(row=0, column=1, columnspan=6, sticky=EW)
        new_notification_variables.append(entry)
        entry.config(state='disabled')
        entry.bind('<Double-Button-1>', lambda x: clear_notification())
        root.update()
############################################################################################
def jackie():
    notifications_count1=2
    
    for each in notifications_list:
        entry = (str("e"+(str(notifications_count1))))
        item ='  -  Date:  '.join(each)
        entry = Label(notifications, text=item)
        entry.grid(row=notifications_count1, column=1, sticky=W)
        notification_variables.append(entry)
        notifications_count1+=1
############################################################################################
def update_jackie():
    for each in notification_variables:
        each.destroy()
    notifications_count1=1
    for each in notifications_list:
        item ='  -  Date:  '.join(each)
        entry = Label(notifications, text=item).grid(row=notifications_count1, column=1, sticky=W)
        notifications_count1+=1
    root.update()
############################################################################################
def menu():
    menubar = Menu(new_case)

    casesmenu = Menu(new_case)
    casesmenu.add_command(label="New Case", command=new_case_gui)
    casesmenu.add_separator()

    #casesmenu.add_command(label="New PD entry", command=add_new_PD_entry)
    #casesmenu.add_separator()
    
    #casesmenu.add_command(label="New Triage entry", command=new_triage_gui)
    #casesmenu.add_separator()
    
    casesmenu.add_command(label="Exit", command=new_case.quit)
    menubar.add_cascade(label="File", menu=casesmenu)
    
    #exportmenu = Menu(new_case)
    #exportmenu.add_command(label="Overview Export", command=export_overview)
    #menubar.add_cascade(label="Export", menu=exportmenu)
                        
    settingsmenu = Menu(new_case)
    settingsmenu.add_command(label="Admin", command=admin_login_gui)
    settingsmenu.add_separator()
    settingsmenu.add_command(label="Settings", command=settings_gui)
    menubar.add_cascade(label="Settings", menu=settingsmenu)

    passwordmenu = Menu(new_case)
    passwordmenu.add_command(label="Generate Password", command=pw_generator_gui)
    passwordmenu.add_command(label="Search Case Password", command=search_pw)
    menubar.add_cascade(label="Password", menu=passwordmenu)
    
    aboutmenu = Menu(new_case)
    aboutmenu.add_command(label="About", command=about_gui)
    menubar.add_cascade(label="About", menu=aboutmenu)
    
    new_case.config(menu=menubar)
############################################################################################
def menu_new_case():
    menubar = Menu(new_case)

    casesmenu = Menu(new_case)
#    casesmenu.add_command(label="New Case", command=new_case_gui)
#    casesmenu.add_command(label="Load Case", command=0)
#    casesmenu.add_separator()
    casesmenu.add_command(label="Exit", command=new_case.quit)
    menubar.add_cascade(label="File", menu=casesmenu)

    settingsmenu = Menu(new_case)
    settingsmenu.add_command(label="Settings", command=settings_gui)
    menubar.add_cascade(label="Settings", menu=settingsmenu)

    aboutmenu = Menu(new_case)
    aboutmenu.add_command(label="About", command=about_gui)
    menubar.add_cascade(label="About", menu=aboutmenu)
    
    new_case.config(menu=menubar)
############################################################################################
def change_startup(estartb):
    for sql in sqlitedbs:
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor() 
        try:
            c.execute('UPDATE profile SET startup = ? WHERE _rowid_ = 1', (estartb,))
        except:
                pass
        conn.close()
############################################################################################
def about_gui():
    global popup
    gui = ("popup")
    popup = Toplevel()
    popup.title("CASE - About")
    popup.resizable(width=False, height=False)

    frames=["popup"]
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count2 < columns:
                Label(popup, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
    Label(popup, text=" Version: 2.0.0 - sdft").grid(row=0, column=1, sticky=W)
    Label(popup, text=" Build date: 18/12/2021").grid(row=1, column=1, sticky=W)
    Label(popup, text=" Contact: ryan.ward@sussex.pnn.police.uk").grid(row=2, column=1, sticky=W)
    
    Label(popup, text=" ").grid(row=5, column=0, columnspan=3, sticky=W)
    
    popup.mainloop()
############################################################################################
def exportpwtolocation():
    global desktop
    desktop = (str(os.environ['USERPROFILE']))
    current_setting = e10a.get()
    d = argv[1] if len(argv)>1 else filedialog.askdirectory(initialdir=desktop)
    D = os.path.realpath(d)
    if d == '':
        e10a.set(current_setting)
    else:
        e10a.set(D)
############################################################################################

# Create system folder
systempath = 'C:\\Case Creator\\'
if not os.path.exists(systempath):      
    os.makedirs(systempath)
sqlprolocal =("C:\\Case Creator\\casecreator.sqlite3")
if not os.path.exists(sqlprolocal):
    createprofile()
############################################################################################
# Read profile sql
readversion()
############################################################################################
# Check software and DB version
'''if version != updateversion:
    print("Version mis-match")
    if version >= updateversion:
        print("Version: ", version)
        print("Update Version: ", updateversion)
    else:
        print("Version: ", version)
        print("Update Version: ", updateversion)
        updateprofile()
else:
    print("Correct DB version")'''
############################################################################################
def grab_case_data(sv):
    data = d1a.get()
    print('print data: ', data)

############################################################################################
def expandview():
    new.geometry('765x427')
    panel.bind('<Triple-Button-1>', lambda x: unexpandview())
    extend.config(command=unexpandview)
    new.update()
############################################################################################
def unexpandview():
    new.geometry('570x427')
    panel.bind('<Triple-Button-1>', lambda x: expandview())
    extend.config(command=expandview)
    new.update()
############################################################################################
def remove_resize_border():
    # Get window handle
    hwnd = ctypes.windll.user32.GetParent(new.winfo_id())
    
    # Windows constants
    GWL_STYLE = -16
    WS_THICKFRAME = 0x00040000
    
    # Get current window style
    current_style = ctypes.windll.user32.GetWindowLongPtrW(hwnd, GWL_STYLE)
    
    # Remove thick frame (resizing border) but keep other functionality
    new_style = current_style & ~WS_THICKFRAME
    
    # Set new window style
    ctypes.windll.user32.SetWindowLongPtrW(hwnd, GWL_STYLE, new_style)
############################################################################################
def setup_new_case():
    global extend, panel, root, new, new_case_vars, d1a, d1b, d2a, d2b, d3a, d3b, e1a, e1b, e2a, e2b, e3a, e3b, e4, e4a, e4b, e5a, e5b, e6a, e6b, e7a, e7b, e24a, e24b, e2, e3, e4, e5, e6, e7, e8, e9, e10a, e10b, submitb, LoadButtona, LoadButtonb, LoadButtona1, LoadButtonb1, text_zoom, pw_generator, desktop, Refresh, e1a, e1b, e2a, e2b, e3a, e3b, e4, e5a, e5b, e6a, e6b, e7a, e7b, e24a, e24b, LoadButtona, LoadButtonb, notepad_n59b, canvas, text_var, lab1, lab2, viewTypea, v0, v1, v2, v3, v4, v5, v6, delEntry_a, delEntry_b, pw_generator, newcase, pw_generator_e1, e1a, e1b, e2, e3, e4, e5, e6, e7, e8, e9, submitb, pw_generator_e1a, label1, submit, edit1b, e2a, e3a, e4a, e5a, e6a, e7a, e8a, e9a, launch, e21a, e21b, e22a, e22b, startup_e1a, startup_e1b, startup_e2a, startup_e2b, pw_customa, pw_customb, pw_generator, default_font, zoomed_font

    new = Tk()# ThemedTk(theme="aqua")  # Use a macOS-like theme#
    #new_case.iconbitmap(iconfile)
    new.geometry('570x427')#600x410')765x427')
    new.resizable(False, False)#width=False, height=False)
    new.title("CASE - V1.0 - ryan.ward@sussex.police.uk")
    
    new.protocol('WM_DELETE_WINDOW', sys.exit)

    # Remove the resize border using Windows API
##    new.update()  # Ensure window is created
##    remove_resize_border()
    
    width  = new.winfo_screenwidth()
    height = new.winfo_screenheight()
    print(width)
    print(height)
    new_case_vars = []
    rows=31
    columns=7
##    tabs = ttk.Notebook(root)
##    tabs.pack(fill='both', expand=Y)
##    new=ttk.Frame()
##    settings_tab=ttk.Frame()
##    structure=ttk.Frame()
##    pw_generator=ttk.Frame()
##    about=ttk.Frame()
##    tabs.add(new, text='New')
##    tabs.add(settings_tab, text='Settings')
##    tabs.add(structure, text='Folder Structure')
##    tabs.add(pw_generator, text='Password')
##    tabs.add(about, text='About')
    # Set row height here
    height=4
    frames=[new]#, settings_tab, structure]
    for frame in frames:
        print(frame)
        count=0
        count2=0
        if frame == frames[0]:
            while count < rows:
                if count == 0:
                    Label(frame, text=" ", width=2).grid(row=count, column=1, sticky=EW, ipady=1)
                    count+=1
                else:
                    Label(frame, text=" ", width=2).grid(row=count, column=1, sticky=EW, ipady=height)
                    count+=1
            while count2 < columns:
                Label(frame, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1


    img = Image.open("C:\\Case Creator\\_logos\\Surrey-Sussex-Police-logo.jpg")#C:\\Case Creator\\_logos\\Surrey-Sussex-Police-logo.jpg")
    #img = img.resize((300, 100), Image.ANTIALIAS)
    img = img.resize((125, 100), Image.ADAPTIVE)
    img = ImageTk.PhotoImage(img)
    panel = Label(new, image=img, width=125, background="white")
    panel.image = img
    panel.grid(row=1, column=1, columnspan=1, rowspan=4, sticky=EW)
    panel.bind('<Triple-Button-1>', lambda x: expandview())
    
    Label1=Label(new, text=" New Case ", width=15, borderwidth=2, relief="groove")
    Label1.grid(row=1, column=3, columnspan=3, sticky=EW)
    new_case_vars.append(Label1)
    Label2=Label(new, text=" DFT Ref(YY-DFTREF):")
    Label2.grid(row=2, column=3, sticky=W)
    new_case_vars.append(Label2)
    
    LoadButtona = StringVar()
    LoadButtona.trace("w", lambda name, index, mode, LoadButtona=LoadButtona: callback(LoadButtona))
    LoadButtonb = Button(new, text="  Load  ", width= 10, command=loadcases_dft)
    LoadButtonb.grid(row=2, column=5, sticky=W)
    LoadButtonb.config(state='disable')
    new_case_vars.append(LoadButtonb)
    
    LoadButtona1 = StringVar()
    LoadButtona1.trace("w", lambda name, index, mode, LoadButtona=LoadButtona: callback(LoadButtona1))
    LoadButtonb1 = Button(new, text="  Load  ", width= 10, command=loadcases_crime)
    LoadButtonb1.grid(row=3, column=5, sticky=W)
    LoadButtonb1.config(state='disable')
    new_case_vars.append(LoadButtonb1)
    
    Label3=Label(new, text=" Crime Reference: ")
    Label3.grid(row=3, column=3, sticky=W)
    new_case_vars.append(Label3)
    Label4=Label(new, text=" Exhibit Reference: ")
    Label4.grid(row=4, column=3, sticky=W)
    new_case_vars.append(Label4)
    Label5=Label(new, text=" Bag Seal Reference: ")
    Label5.grid(row=5, column=3, sticky=W)
    new_case_vars.append(Label5)
    Label6=Label(new, text=" OIC/Contact: ")
    Label6.grid(row=6, column=3, sticky=W)
    new_case_vars.append(Label6)
    Label7=Label(new, text=" Operation Name: ")
    Label7.grid(row=7, column=3, sticky=W)
    new_case_vars.append(Label7)
    Label8=Label(new, text=" Suspect Name: ")
    Label8.grid(row=8, column=3, sticky=W)
    new_case_vars.append(Label8)
    Label9=Label(new, text=" Property Reference: ")
    Label9.grid(row=9, column=3, sticky=W)
    new_case_vars.append(Label9)
    Label10=Label(new, text=" Date of offence: ")
    Label10.grid(row=10, column=3, sticky=W)
    new_case_vars.append(Label10)
    Label11=Label(new, text=" Case Password: ")
    Label11.grid(row=11, column=3, sticky=W)
    new_case_vars.append(Label11)

    ############################################################################################
    def callbacke1(sv):
        print(sv.get())
        entry = e1a.get()
        #print(str('entry: ' + entry[-1:3]))

        if len(entry) == 3:
            if entry[-1:3] != '-' :
                e1a.set(entry[:2])
            else:
                e1a.set(entry[:7])
        else:
            e1a.set(entry[:7])

    ############################################################################################
    row = 2
    column = 4
    e1a = StringVar()
    e1a.trace("w", lambda name, index, mode, e1a=e1a: callback1(e1a))
    e1b = Entry(new,relief=SUNKEN,width=30, textvariable=e1a)
    e1b.grid(row=2, column=column)
    e1b.focus()
    new_case_vars.append(e1b)
    
    e2a = StringVar()
    e2a.trace("w", lambda name, index, mode, e2a=e2a: callback2(e2a))
    e2b = Entry(new,relief=SUNKEN,width=30, textvariable=e2a)
    e2b.grid(row=3, column=column)
    new_case_vars.append(e2b)

    e3a = StringVar()
    e3a.trace("w", lambda name, index, mode, e3a=e3a: callback3(e3a))
    e3b = Entry(new,relief=SUNKEN,width=30, textvariable=e3a)
    e3b.grid(row=4, column=column)
    new_case_vars.append(e3b)

    e4a = StringVar() # bag seal
    e4a.trace("w", lambda name, index, mode, e4a=e4a: callback(e4a))
    e4b = Entry(new,relief=SUNKEN,width=30, textvariable=e4a)
    e4b.grid(row=5, column=column)
    new_case_vars.append(e4b)

    e5a = StringVar() # OIC
    e5a.trace("w", lambda name, index, mode, e5a=e5a: callback(e5a))
    e5b = Entry(new,relief=SUNKEN,width=30, textvariable=e5a)
    e5b.grid(row=6, column=column)
    new_case_vars.append(e5b)

    e6a = StringVar()
    e6a.trace("w", lambda name, index, mode, e6a=e6a: callback(e6a))
    e6b = Entry(new,relief=SUNKEN,width=30, textvariable=e6a)
    e6b.grid(row=7, column=column)
    new_case_vars.append(e6b)

    e7a = StringVar()
    e7a.trace("w", lambda name, index, mode, e7a=e7a: callback(e7a))
    e7b = Entry(new,relief=SUNKEN,width=30, textvariable=e7a)
    e7b.grid(row=8, column=column)
    new_case_vars.append(e7b)

    e8 = Entry(new,relief=SUNKEN,width=30)
    e8.grid(row=9, column=column)
    new_case_vars.append(e8)
    
    submita = StringVar()
    submita.trace("w", lambda name, index, mode, submita=submita: callback(submita))
    submitb = Button(new, text="  Submit  ", width= 10, height=3,  command=duplicate_entry)
    submitb.grid(row=10, column=5, rowspan=2, sticky=W)
    new_case_vars.append(submitb)
    submitb.config(state='normal')

    e24a = StringVar()
    e24b = Entry(new,relief=SUNKEN,width=30, textvariable=e24a)
    e24b.grid(row=10, column=column)
    new_case_vars.append(e24b)

    gen_pw()
    
    e9a = StringVar()
    e9a.set(res)
    e9a.trace("w", lambda name, index, mode, e9a=e9a: callback(e9a))
    e9b = Entry(new,relief=SUNKEN,width=30, textvariable=e9a)
    e9b.grid(row=11, column=column)
    e9b.config(state='disabled')
    new_case_vars.append(e9b)

    Label(new, text=" ").grid(row=12, column=1, sticky=E)

    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
    Label(new, text=" Case Directory: ").grid(row=13, column=1, sticky=E)
    casedir = Button(new, text="  Browse  ", width= 10, command=exportpwtolocation, state='normal')
    casedir.grid(row=13, column=5, sticky=W)

    e10a = StringVar()
    e10a.set(case_dir)
    e10a.trace("w", lambda name, index, mode, e10a=e10a: callback(e10a))
    e10b = Entry(new,relief=SUNKEN, width=40, textvariable=e10a)
    e10b.grid(row=13, column=3, columnspan=2, sticky=EW)
    e10b.config(state='disabled')

    ##Extra Cases section
    extend = Button(new, text="  Options  ", command=structure_gui)#expandview)
    extend.grid(row=6, column=1, sticky=EW)
                 
    Label_c1=Label(new, text=" ", width=2)
    Label_c1.grid(row=2, column=6, sticky=W)
    
##    c1a = StringVar()
##    c1a.set('')
##    c1a.trace("w", lambda name, index, mode, e10a=e10a: callback(e10a))
##    c1b = Entry(new,relief=SUNKEN, width=20, textvariable=e10a)
##    c1b.grid(row=2, column=7, columnspan=1, sticky=EW)

    Label(new, text=" Open Cases ", borderwidth=2, relief="groove").grid(row=1, column=7, sticky=EW)
    d1a = StringVar()
    d1a.set(opencases_easy_view[0])
    d1a.trace("w", lambda name, index, mode, d1a=d1a: callback(d1a))
    d1b = OptionMenu(new, d1a, *opencases_easy_view, command=0)
    d1b.grid(row=2, column=7, sticky=W)
    d1b.config(width=23)
    
    casedir = Button(new, text="  Open  ", command=populate_individual_case, state='normal')
    casedir.grid(row=3, column=7, sticky=EW)

    
    Label(new, text=" Closed Cases ", borderwidth=2, relief="groove").grid(row=5, column=7, sticky=EW)
    d2a = StringVar()
    d2a.set(opencases_easy_view[0])
    d2a.trace("w", lambda name, index, mode, d2a=d2a: callback(d2a))
    d2b = OptionMenu(new, d2a, *opencases_easy_view, command=0)
    d2b.grid(row=6, column=7, sticky=W)
    d2b.config(width=23)

    def callback_d3a(d3a):
        a=d3a.get()
        global case_password, opencases_easy_view, cases, casecount, case_dir, template_dir, contemp_file, examiner
        conn = sqlite3.connect(sqlprolocal, isolation_level=None)
        conn.execute('pragma journal_mode=wal')
        c=conn.cursor()
        b=[]
        c.execute(""" select * from cases""")
        c.execute("SELECT * FROM cases WHERE full_dft_ref like '%'||?||'%'", (a[:7],))
        for each in c:
            print(each[12])
            b.append(each[12])
        conn.close()
        print('b', b)
        d4a.set(b[0])
        
    Label(new, text=" Case Passwords ", borderwidth=2, relief="groove").grid(row=8, column=7, sticky=EW)
    d3a = StringVar()
    d3a.set(case_password[0])
    d3a.trace("w", lambda name, index, mode, d3a=d3a: callback_d3a(d3a))
    d3b = OptionMenu(new, d3a, *case_password, command=0)
    d3b.grid(row=9, column=7, sticky=W)
    d3b.config(width=23)

    d4a = StringVar()
    d4a.set('')
    d4a.trace("w", lambda name, index, mode, d4a=d4a: callback(d4a))
    d4b = Entry(new,relief=SUNKEN,  textvariable=d4a)
    d4b.grid(row=10, column=7, sticky=EW)
    
    new.mainloop()
#######################################################################################
def structure_gui():
    global structure, row, column, f1variables, r1c1
    structure = Tk()
    structure.title("CASE - Structure")
    structure.resizable(width=False, height=False)
    structure.geometry('875x400')
    
    #Frame1
    row=2
    column = 2

    f1variables=[]
    cell_width=22
    Label(structure, text=" ", width=cell_width).grid(row=1, column=2, sticky=W)
    Label(structure, text=" ", width=cell_width).grid(row=1, column=3, sticky=W)
    Label(structure, text=" ", width=cell_width).grid(row=1, column=4, sticky=W)
    Label(structure, text=" ", width=cell_width).grid(row=1, column=5, sticky=W)
    Label(structure, text=" ", width=cell_width).grid(row=1, column=6, sticky=W)

    
    Label(structure, text="  ").grid(row=1, column=1, sticky=W)
    Label(structure, text="Level 1 ").grid(row=2, column=1, sticky=W)
    Label(structure, text="Level 2 ").grid(row=4, column=1, sticky=W)
    Label(structure, text="Level 3 ").grid(row=6, column=1, sticky=W)
    Label(structure, text="Level 4 ").grid(row=8, column=1, sticky=W)
    Label(structure, text="Level 5 ").grid(row=10, column=1, sticky=W)


    case_structure_options=['Please select', 'Crime reference', 'DFT reference', 'Exhibit reference', 'Generated material', 'Source files']
    
    r1c1a = StringVar()
    r1c1a.set(case_structure_options[0])
    r1c1a.trace("w", lambda name, index, mode, r1c1a=r1c1a: callback(r1c1a))
    r1c1b = OptionMenu(structure, r1c1a, *case_structure_options, command=0)
    r1c1b.grid(row=row, column=column, sticky=EW)
    r1c1b.config(state='normal', fg="black", activebackground="lightgrey", activeforeground="black")
    #r1c1a.bind()
    column+=1

    r1c2a = StringVar()
    r1c2a.set(case_structure_options[0])
    r1c2a.trace("w", lambda name, index, mode, r1c2a=r1c2a: callback(r1c2a))
    r1c2b = OptionMenu(structure, r1c2a, *case_structure_options, command=0)
    r1c2b.grid(row=row, column=column, sticky=EW)
    r1c2b.config(state='normal')

    column+=1

    r1c3 = Entry(structure,relief=SUNKEN,width=cell_width).grid(row=row, column=column, sticky=W)
    r1c3a = StringVar()
    r1c3a.set(case_structure_options[0])
    r1c3a.trace("w", lambda name, index, mode, r1c3a=r1c3a: callback(r1c3a))
    r1c3b = OptionMenu(structure, r1c3a, *case_structure_options, command=0)
    r1c3b.grid(row=row, column=column, sticky=EW)
    r1c3b.config(state='normal')

    column+=1

    r1c4a = StringVar()
    r1c4a.set(case_structure_options[0])
    r1c4a.trace("w", lambda name, index, mode, r1c4a=r1c4a: callback(r1c4a))
    r1c4b = OptionMenu(structure, r1c4a, *case_structure_options, command=0)
    r1c4b.grid(row=row, column=column, sticky=EW)
    r1c4b.config(state='normal')

    column+=1

    r1c5a = StringVar()
    r1c5a.set(case_structure_options[0])
    r1c5a.trace("w", lambda name, index, mode, r1c5a=r1c5a: callback(r1c5a))
    r1c5b = OptionMenu(structure, r1c5a, *case_structure_options, command=0)
    r1c5b.grid(row=row, column=column, sticky=EW)
    r1c5b.config(state='normal')

    row+=2
    column = 2
    Label(structure, text=" ").grid(row=3, column=1, sticky=W)
    ###Row 2###
    r2c1a = StringVar()
    r2c1a.set(case_structure_options[0])
    r2c1a.trace("w", lambda name, index, mode, r2c1a=r2c1a: callback(r2c1a))
    r2c1b = OptionMenu(structure, r2c1a, *case_structure_options, command=0)
    r2c1b.grid(row=row, column=column, sticky=EW)
    r2c1b.config(state='normal')

    column+=1

    r2c2a = StringVar()
    r2c2a.set(case_structure_options[0])
    r2c2a.trace("w", lambda name, index, mode, r2c2a=r2c2a: callback(r2c2a))
    r2c2b = OptionMenu(structure, r2c2a, *case_structure_options, command=0)
    r2c2b.grid(row=row, column=column, sticky=EW)
    r2c2b.config(state='normal')

    column+=1

    r2c3a = StringVar()
    r2c3a.set(case_structure_options[0])
    r2c3a.trace("w", lambda name, index, mode, r2c3a=r2c3a: callback(r2c3a))
    r2c3b = OptionMenu(structure, r2c3a, *case_structure_options, command=0)
    r2c3b.grid(row=row, column=column, sticky=EW)
    r2c3b.config(state='normal')

    column+=1

    r2c4a = StringVar()
    r2c4a.set(case_structure_options[0])
    r2c4a.trace("w", lambda name, index, mode, r2c4a=r2c4a: callback(r2c4a))
    r2c4b = OptionMenu(structure, r2c4a, *case_structure_options, command=0)
    r2c4b.grid(row=row, column=column, sticky=EW)
    r2c4b.config(state='normal')

    column+=1

    r2c5a = StringVar()
    r2c5a.set(case_structure_options[0])
    r2c5a.trace("w", lambda name, index, mode, r2c5a=r2c5a: callback(r2c5a))
    r2c5b = OptionMenu(structure, r2c5a, *case_structure_options, command=0)
    r2c5b.grid(row=row, column=column, sticky=EW)
    r2c5b.config(state='normal')
    row+=2
    column = 2
    Label(structure, text=" ").grid(row=5, column=1, sticky=W)
    ###Row 3###
    r3c1a = StringVar()
    r3c1a.set(case_structure_options[0])
    r3c1a.trace("w", lambda name, index, mode, r3c1a=r3c1a: callback(r3c1a))
    r3c1b = OptionMenu(structure, r3c1a, *case_structure_options, command=0)
    r3c1b.grid(row=row, column=column, sticky=EW)
    r3c1b.config(state='normal')

    column+=1

    r3c2a = StringVar()
    r3c2a.set(case_structure_options[0])
    r3c2a.trace("w", lambda name, index, mode, r3c2a=r3c2a: callback(r3c2a))
    r3c2b = OptionMenu(structure, r3c2a, *case_structure_options, command=0)
    r3c2b.grid(row=row, column=column, sticky=EW)
    r3c2b.config(state='normal')

    column+=1

    r3c3a = StringVar()
    r3c3a.set(case_structure_options[0])
    r3c3a.trace("w", lambda name, index, mode, r3c3a=r3c3a: callback(r3c3a))
    r3c3b = OptionMenu(structure, r3c3a, *case_structure_options, command=0)
    r3c3b.grid(row=row, column=column, sticky=EW)
    r3c3b.config(state='normal')

    column+=1

    r3c4a = StringVar()
    r3c4a.set(case_structure_options[0])
    r3c4a.trace("w", lambda name, index, mode, r3c4a=r3c4a: callback(r3c4a))
    r3c4b = OptionMenu(structure, r3c4a, *case_structure_options, command=0)
    r3c4b.grid(row=row, column=column, sticky=EW)
    r3c4b.config(state='normal')

    column+=1


    r3c5a = StringVar()
    r3c5a.set(case_structure_options[0])
    r3c5a.trace("w", lambda name, index, mode, r3c5a=r3c5a: callback(r3c5a))
    r3c5b = OptionMenu(structure, r3c5a, *case_structure_options, command=0)
    r3c5b.grid(row=row, column=column, sticky=EW)
    r3c5b.config(state='normal')
    row+=2
    column = 2
    Label(structure, text=" ").grid(row=7, column=1, sticky=W)
    ###Row 4###
    r4c1a = StringVar()
    r4c1a.set(case_structure_options[0])
    r4c1a.trace("w", lambda name, index, mode, r4c1a=r4c1a: callback(r4c1a))
    r4c1b = OptionMenu(structure, r4c1a, *case_structure_options, command=0)
    r4c1b.grid(row=row, column=column, sticky=EW)
    r4c1b.config(state='normal')

    column+=1

    r4c2a = StringVar()
    r4c2a.set(case_structure_options[0])
    r4c2a.trace("w", lambda name, index, mode, r4c2a=r4c2a: callback(r4c2a))
    r4c2b = OptionMenu(structure, r4c2a, *case_structure_options, command=0)
    r4c2b.grid(row=row, column=column, sticky=EW)
    r4c2b.config(state='normal')

    column+=1

    r4c3a = StringVar()
    r4c3a.set(case_structure_options[0])
    r4c3a.trace("w", lambda name, index, mode, r4c3a=r4c3a: callback(r4c3a))
    r4c3b = OptionMenu(structure, r4c3a, *case_structure_options, command=0)
    r4c3b.grid(row=row, column=column, sticky=EW)
    r4c3b.config(state='normal')

    column+=1

    r4c4a = StringVar()
    r4c4a.set(case_structure_options[0])
    r4c4a.trace("w", lambda name, index, mode, r4c4a=r4c4a: callback(r4c4a))
    r4c4b = OptionMenu(structure, r4c4a, *case_structure_options, command=0)
    r4c4b.grid(row=row, column=column, sticky=EW)
    r4c4b.config(state='normal')

    column+=1

    r4c5a = StringVar()
    r4c5a.set(case_structure_options[0])
    r4c5a.trace("w", lambda name, index, mode, r4c5a=r4c5a: callback(r4c5a))
    r4c5b = OptionMenu(structure, r4c5a, *case_structure_options, command=0)
    r4c5b.grid(row=row, column=column, sticky=EW)
    r4c5b.config(state='normal')
    row+=2
    column = 2
    Label(structure, text=" ").grid(row=9, column=1, sticky=W)
    ###Row 5###
    r5c1a = StringVar()
    r5c1a.set(case_structure_options[0])
    r5c1a.trace("w", lambda name, index, mode, r5c1a=r5c1a: callback(r5c1a))
    r5c1b = OptionMenu(structure, r5c1a, *case_structure_options, command=0)
    r5c1b.grid(row=row, column=column, sticky=EW)
    r5c1b.config(state='normal')

    column+=1

    r5c2a = StringVar()
    r5c2a.set(case_structure_options[0])
    r5c2a.trace("w", lambda name, index, mode, r5c2a=r5c2a: callback(r5c2a))
    r5c2b = OptionMenu(structure, r5c2a, *case_structure_options, command=0)
    r5c2b.grid(row=row, column=column, sticky=EW)
    r5c2b.config(state='normal')

    column+=1

    r5c3a = StringVar()
    r5c3a.set(case_structure_options[0])
    r5c3a.trace("w", lambda name, index, mode, r5c3a=r5c3a: callback(r5c3a))
    r5c3b = OptionMenu(structure, r5c3a, *case_structure_options, command=0)
    r5c3b.grid(row=row, column=column, sticky=EW)
    r5c3b.config(state='normal')

    column+=1

    r5c4a = StringVar()
    r5c4a.set(case_structure_options[0])
    r5c4a.trace("w", lambda name, index, mode, r5c4a=r5c4a: callback(r5c4a))
    r5c4b = OptionMenu(structure, r5c4a, *case_structure_options, command=0)
    r5c4b.grid(row=row, column=column, sticky=EW)
    r5c4b.config(state='normal')

    column+=1

    r5c5a = StringVar()
    r5c5a.set(case_structure_options[0])
    r5c5a.trace("w", lambda name, index, mode, r5c5a=r5c5a: callback(r5c5a))
    r5c5b = OptionMenu(structure, r5c5a, *case_structure_options, command=0)
    r5c5b.grid(row=row, column=column, sticky=EW)
    r5c5b.config(state='normal')

    row+=1
    Label(structure, text=" ").grid(row=row, column=1, sticky=W)
    row+=1
    commit_button = Button(structure, text="  Submit / Update  ", command=addrow)
    commit_button.grid(row=row, column=column, padx=0, ipadx=0, sticky=EW)

    
    print(locals())
    structure.attributes('-topmost',True)
    structure.mainloop()
############################################################################################
def addrow():
    global row, entry_b
    print(r1c1)
    entry = (str("e"+(str(row))+(str(row))))
    entry_b = (str("e"+(str(row))+(str(row))))+("_b")
    row_str = (str(row-1))
    name = (str('Add column '+row_str))
    print('entry_b', entry_b)
    entry_b=Button(structure, text=name, command=addcolumn)
    entry_b.grid(row=row, column=0, sticky=W)
    print('entry_b', entry_b)
    print('entry', entry)
    f1variables.append(entry_b)
    entry = Entry(structure,relief=SUNKEN,width=40)
    entry.grid(row=row, column=1)
    f1variables.append(entry)
    row+=1
    entry.insert(10,"ryan1")
    structure.update()

    print('locals 1:', locals())
    #print('globals 1:', globals())
############################################################################################
def addcolumn():
    global row, column
    text = entry_b.cget("text")
    # Get grid information for the clicked button
    info = entry_b.grid_info()
    row = info['row']
    col = info['column']
    print('btn_text: ',text)
    #messagebox.showinfo("Button Info", f"Row: {row}, Column: {col}")
    dynamic_column = text[11:]
    dynamic_column = int(dynamic_column)
    dynamic_column+=1
    dynamic_row = text[11:]
    dynamic_row = int(dynamic_row)
    dynamic_row+=1
    print('dynamic_column:', dynamic_column)
    column+=1
    #Label(structure, text=row).grid(row=row, column=0, sticky=W)
    entry = (str("e"+(str(row))+(str(row))))
    print(entry)
    entry = Entry(structure,relief=SUNKEN,width=40)
    f1variables.append(entry)
    entry.grid(row=1, column=dynamic_column)
    entry.insert(10,"ryan1")
    structure.update()
    
    print('locals 2:', locals())
############################################################################################
def printdata():
        count=0
        for each in f1variables:
                value=(f1variables[count].get())
                print(value)
                count+=1
############################################################################################
##    ### Settings_tab ######################################################################
##    Label1=Label(settings_tab, text=" Case & template settings_tab ", width=10, borderwidth=2).grid(row=1, column=1, columnspan=8, sticky=EW)
##    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
##    Label(settings_tab, text=" Case Directory: ").grid(row=2, column=1, sticky=W)
##    casedir = Button(settings_tab, text="  Browse  ", width= 10, command=selectExamDir, state='disabled')
##    casedir.grid(row=2, column=6, sticky=W)
##    tempdir = Button(settings_tab, text="  Browse  ", width= 10, command=selectTempDir, state='disabled')
##    tempdir.grid(row=3, column=6, sticky=W)
##
##    Label1=Label(settings_tab, text=" Document settings_tab ", width=10, borderwidth=2).grid(row=4, column=1, columnspan=8, sticky=EW)
##    Label(settings_tab, text=" Template Directory: ").grid(row=3, column=1, sticky=W)
##    Label(settings_tab, text=" Contemp Notes: ").grid(row=5, column=1, sticky=W)
##    Label(settings_tab, text=" SFR: ").grid(row=7, column=1, sticky=W)
##    Label(settings_tab, text=" Disclosure Cert: ").grid(row=9, column=1, sticky=W)
##    ############################################################################################
##    def callbacke1(sv):
##        entry = e1a.get()
##        if len(entry) == 3:
##            if entry[-1:3] != '-' :
##                e1a.set(entry[:2])
##            else:
##                e1a.set(entry[:7])
##        else:
##            e1a.set(entry[:7])
##
##    ############################################################################################
##    entry_width = 50
##    
##    e10 = StringVar()
##    e10.set(case_dir)
##    e10 = Entry(settings_tab,relief=SUNKEN, width=55, textvariable=e10)
##    e10.grid(row=2, column=2, columnspan=2, sticky=W)
##    e10.config(state='disabled')
##
##    e11 = StringVar()
##    e11.set(template_dir)
##    e11 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e11)
##    e11.grid(row=3, column=2, columnspan=2, sticky=W)
##    e11.config(state='disabled')
##
##    e12 = StringVar()
##    e12.set(contemp_file) #notes_file
##    e12 = Entry(settings_tab,relief=SUNKEN,width=entry_width, textvariable=e12)
##    e12.grid(row=5, column=2, columnspan=2, sticky=W)
##
##    #Copy notes_file to case checkbox
##    e33a = IntVar()
##    e33a.set(contemp_copy)
##    e33a.trace("w", lambda name, index, mode, e33a=e33a: callbacke33(e33a))
##    e33b = Checkbutton(settings_tab, text="Copy", onvalue=1,offvalue=0,variable=e33a, command=e33checkbox)
##    e33b.grid(row=5, column=3, sticky=EW)
##    e33b.config(state='disabled')
##
####    #Populate notes_file checkbox
####    e34a = IntVar()
####    e34a.set(contemp_populate)
####    e34a.trace("w", lambda name, index, mode, e34a=e34a: callbacke34(e34a))
####    e34b = Checkbutton(settings_tab, text="Populate?", onvalue=1,offvalue=0,variable=e34a)
####    e34b.grid(row=6, column=3, sticky=W)
####    e34b.config(state='disabled')
##
##    e25 = StringVar()
##    e25.set(sfr_file) #sfr_file
##    e25 = Entry(settings_tab,relief=SUNKEN,width=entry_width, textvariable=e25)
##    e25.grid(row=7, column=2, columnspan=2, sticky=W)
##
##    #Copy sfr_file to case checkbox
##    e35a = IntVar()
##    e35a.set(sfr_copy)
##    e35a.trace("w", lambda name, index, mode, e35a=e35a: callbacke35(e35a))
##    e35b = Checkbutton(settings_tab, text="Copy", onvalue=1,offvalue=0,variable=e35a, command=e35checkbox)
##    e35b.grid(row=7, column=3, sticky=EW)
##    e35b.config(state='disabled')
##
####    #Populate sfr_file checkbox
####    e36a = IntVar()
####    e36a.set(sfr_populate)
####    e36a.trace("w", lambda name, index, mode, e36a=e36a: callbacke36(e36a))
####    e36b = Checkbutton(settings_tab, text="Populate?", onvalue=1,offvalue=0,variable=e36a)
####    e36b.grid(row=8, column=3, sticky=W)
####    e36b.config(state='disabled')
##
##    e32 = StringVar()
##    e32.set(disclosure_file) #disclosure_file
##    e32 = Entry(settings_tab,relief=SUNKEN,width=entry_width, textvariable=e32)
##    e32.grid(row=9, column=2, columnspan=2, sticky=W)
##
##    #Copy disclosure_file to case checkbox
##    e37a = IntVar()
##    e37a.set(disclosure_copy)
##    e37a.trace("w", lambda name, index, mode, e37a=e37a: callbacke37(e37a))
##    e37b = Checkbutton(settings_tab, text="Copy", onvalue=1,offvalue=0,variable=e37a, command=e37checkbox)
##    e37b.grid(row=9, column=3, sticky=EW)
##    e37b.config(state='disabled')
##
####    #Populate disclosure_file checkbox
####    e38a = IntVar()
####    e38a.set(disclosure_populate)
####    e38a.trace("w", lambda name, index, mode, e38a=e38a: callbacke38(e38a))
####    e38b = Checkbutton(settings_tab, text="?", onvalue=1,offvalue=0,variable=e38a)
####    e38b.grid(row=10, column=3, sticky=W)
####    e38b.config(state='disabled')
##
##    contempdir = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectContemp)
##    contempdir.grid(row=5, column=6, sticky=W)
##    sfrdir = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectSfr)
##    sfrdir.grid(row=7, column=6, sticky=W)
##    disclosuredir = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectDisclosure)
##    disclosuredir.grid(row=9, column=6, sticky=W)
####
####    Label(settings_tab, text=" Database details ", width=10, borderwidth=2).grid(row=11, column=1, columnspan=8, sticky=EW)
####
####    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
####    Label(settings_tab, text=" Network database: ").grid(row=12, column=1, sticky=W)
####    Label(settings_tab, text=" Local database: ").grid(row=13, column=1, sticky=W)
####
####    db10 = StringVar()
####    db10.set(sqlprolocal)
####    db10 = Entry(settings_tab,relief=SUNKEN, width=55, textvariable=db10)
####    db10.grid(row=12, column=2, columnspan=2, sticky=EW)
####    db10.config(state='disabled')
####
####    db11 = StringVar()
####    db11.set(sqlprolocal)
####    db11 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=db11)
####    db11.grid(row=13, column=2, columnspan=2, sticky=EW)
####    db11.config(state='disabled')
####
####    networkdb = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectNetworkDB)
####    networkdb.grid(row=12, column=6, sticky=W)
####    localdb = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectLocalDB)
####    localdb.grid(row=13, column=6, sticky=W)
##    
##    e10.config(state='disabled')
##    e11.config(state='disabled')
##    e12.config(state='disabled')
##    e25.config(state='disabled')
##    e32.config(state='disabled')
##
##    ## Folder - settings_tab
##    Label1=Label(settings_tab, text=" Automatic Folder creation ", width=10, borderwidth=2).grid(row=14, column=1, columnspan=8, sticky=EW)
##    Label(settings_tab, text=" Exam Folder 1: ").grid(row=15, column=1, sticky=W)
##    Label(settings_tab, text=" Exam Folder 2: ").grid(row=16, column=1, sticky=W)
##    Label(settings_tab, text=" Exam Folder 3: ").grid(row=17, column=1, sticky=W)
##    Label(settings_tab, text=" Exam Folder 4: ").grid(row=18, column=1, sticky=W)
##
##    e13 = StringVar()
##    e13.set(fld1)
##    e13 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e13)
##    e13.grid(row=15, column=2, columnspan=2, sticky=EW)
##
##    e26a = IntVar()
##    e26a.set(pp1)
##    e26a.trace("w", lambda name, index, mode, e26a=e26a: callbacke26(e26a))
##    e26b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e26a)
##    e26b.grid(row=15, column=5, sticky=E)
##    e26b.config(state='disabled')
##
##    e14 = StringVar()
##    e14.set(fld2)
##    e14 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e14)
##    e14.grid(row=16, column=2, columnspan=2, sticky=EW)
##
##    e27a = IntVar()
##    e27a.set(pp2)
##    e27a.trace("w", lambda name, index, mode, e27a=e27a: callbacke27(e27a))
##    e27b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e27a)
##    e27b.grid(row=16, column=5, sticky=E)
##    e27b.config(state='disabled')
##
##    e15 = StringVar()
##    e15.set(fld3)
##    e15 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e15)
##    e15.grid(row=17, column=2, columnspan=2, sticky=EW)
##
##    e28a = IntVar()
##    e28a.set(pp3)
##    e28a.trace("w", lambda name, index, mode, e28a=e28a: callbacke28(e28a))
##    e28b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e28a)
##    e28b.grid(row=17, column=5, sticky=E)
##    e28b.config(state='disabled')
##
##    e16 = StringVar()
##    e16.set(fld4)
##    e16 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e16)
##    e16.grid(row=18, column=2, columnspan=2, sticky=EW)
##
##    e29a = IntVar()
##    e29a.set(pp4)
##    e29a.trace("w", lambda name, index, mode, e29a=e29a: callbacke29(e29a))
##    e29b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e29a)
##    e29b.grid(row=18, column=5, sticky=E)
##    e29b.config(state='disabled')
##
##
##    e13.config(state='disabled')
##    e14.config(state='disabled')
##    e15.config(state='disabled')
##    e16.config(state='disabled')
## 
##
##    Label1=Label(settings_tab, text=" Analyst settings_tab ", width=10, borderwidth=2).grid(row=19, column=1, columnspan=8, sticky=EW)
#### Folder - settings_tab
##    Label(settings_tab, text=" Analyst: ").grid(row=20, column=1, sticky=W)
##    e9a = StringVar()
##    e9a.set(examiner)
##    e9a.trace("w", lambda name, index, mode, e9a=e9a: callback(e9a))
##    e9b = OptionMenu(settings_tab, e9a, *analysts, command=updateexaminer)
##    e9b.config(width=30)
##    e9b.config(state='disabled')
##    e9b.grid(row=20, column=2, sticky=W)
##
#### Structure ##############################################################
#############################################################################
##    
##    Label1=Label(structure, text=" Folder structure  ", width=10, borderwidth=2).grid(row=1, column=1, columnspan=8, sticky=EW)
##    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
##    Label(structure, text=" Level 1: ").grid(row=2, column=1, sticky=W)
##    casedir = Button(structure, text="  Browse  ", width= 10, command=selectExamDir, state='disabled')
##    casedir.grid(row=2, column=6, sticky=W)
##    tempdir = Button(structure, text="  Browse  ", width= 10, command=selectTempDir, state='disabled')
##    tempdir.grid(row=3, column=6, sticky=W)
##
##    #Label1=Label(structure, text=" Document structure  ", width=10, borderwidth=2).grid(row=4, column=1, columnspan=8, sticky=EW)
##    Label(structure, text=" Level 2: ").grid(row=3, column=1, sticky=W)
##    Label(structure, text=" Level 3: ").grid(row=5, column=1, sticky=W)
##    Label(structure, text=" Level 4: ").grid(row=7, column=1, sticky=W)
##    Label(structure, text=" Level 5: ").grid(row=9, column=1, sticky=W)
##    ############################################################################################
##    def callbacke1(sv):
##        entry = e1a.get()
##        if len(entry) == 3:
##            if entry[-1:3] != '-' :
##                e1a.set(entry[:2])
##            else:
##                e1a.set(entry[:7])
##        else:
##            e1a.set(entry[:7])
##
##    ############################################################################################
##    entry_width = 50
##    
##    e10 = StringVar()
##    e10.set(case_dir)
##    e10 = Entry(structure ,relief=SUNKEN, width=55, textvariable=e10)
##    e10.grid(row=2, column=2, columnspan=2, sticky=W)
##    e10.config(state='disabled')
##
##    e11 = StringVar()
##    e11.set(template_dir)
##    e11 = Entry(structure ,relief=SUNKEN,width=55, textvariable=e11)
##    e11.grid(row=3, column=2, columnspan=2, sticky=W)
##    e11.config(state='disabled')
##
##    e12 = StringVar()
##    e12.set(contemp_file) #notes_file
##    e12 = Entry(structure ,relief=SUNKEN,width=entry_width, textvariable=e12)
##    e12.grid(row=5, column=2, columnspan=2, sticky=W)
##
##    #Copy notes_file to case checkbox
##    e33a = IntVar()
##    e33a.set(contemp_copy)
##    e33a.trace("w", lambda name, index, mode, e33a=e33a: callbacke33(e33a))
##    e33b = Checkbutton(structure, text="Copy", onvalue=1,offvalue=0,variable=e33a, command=e33checkbox)
##    e33b.grid(row=5, column=3, sticky=EW)
##    e33b.config(state='disabled')
##
####    #Populate notes_file checkbox
####    e34a = IntVar()
####    e34a.set(contemp_populate)
####    e34a.trace("w", lambda name, index, mode, e34a=e34a: callbacke34(e34a))
####    e34b = Checkbutton(structure, text="Populate?", onvalue=1,offvalue=0,variable=e34a)
####    e34b.grid(row=6, column=3, sticky=W)
####    e34b.config(state='disabled')
##
##    e25 = StringVar()
##    e25.set(sfr_file) #sfr_file
##    e25 = Entry(structure ,relief=SUNKEN,width=entry_width, textvariable=e25)
##    e25.grid(row=7, column=2, columnspan=2, sticky=W)
##
##    #Copy sfr_file to case checkbox
##    e35a = IntVar()
##    e35a.set(sfr_copy)
##    e35a.trace("w", lambda name, index, mode, e35a=e35a: callbacke35(e35a))
##    e35b = Checkbutton(structure, text="Copy", onvalue=1,offvalue=0,variable=e35a, command=e35checkbox)
##    e35b.grid(row=7, column=3, sticky=EW)
##    e35b.config(state='disabled')
##
####    #Populate sfr_file checkbox
####    e36a = IntVar()
####    e36a.set(sfr_populate)
####    e36a.trace("w", lambda name, index, mode, e36a=e36a: callbacke36(e36a))
####    e36b = Checkbutton(structure, text="Populate?", onvalue=1,offvalue=0,variable=e36a)
####    e36b.grid(row=8, column=3, sticky=W)
####    e36b.config(state='disabled')
##
##    e32 = StringVar()
##    e32.set(disclosure_file) #disclosure_file
##    e32 = Entry(structure ,relief=SUNKEN,width=entry_width, textvariable=e32)
##    e32.grid(row=9, column=2, columnspan=2, sticky=W)
##
##    #Copy disclosure_file to case checkbox
##    e37a = IntVar()
##    e37a.set(disclosure_copy)
##    e37a.trace("w", lambda name, index, mode, e37a=e37a: callbacke37(e37a))
##    e37b = Checkbutton(structure, text="Copy", onvalue=1,offvalue=0,variable=e37a, command=e37checkbox)
##    e37b.grid(row=9, column=3, sticky=EW)
##    e37b.config(state='disabled')
##
####    #Populate disclosure_file checkbox
####    e38a = IntVar()
####    e38a.set(disclosure_populate)
####    e38a.trace("w", lambda name, index, mode, e38a=e38a: callbacke38(e38a))
####    e38b = Checkbutton(structure, text="?", onvalue=1,offvalue=0,variable=e38a)
####    e38b.grid(row=10, column=3, sticky=W)
####    e38b.config(state='disabled')
##
##    contempdir = Button(structure, text="  Browse  ",  width= 10, state='disabled', command=selectContemp)
##    contempdir.grid(row=5, column=6, sticky=W)
##    sfrdir = Button(structure, text="  Browse  ",  width= 10, state='disabled', command=selectSfr)
##    sfrdir.grid(row=7, column=6, sticky=W)
##    disclosuredir = Button(structure, text="  Browse  ",  width= 10, state='disabled', command=selectDisclosure)
##    disclosuredir.grid(row=9, column=6, sticky=W)
####
####    Label(structure, text=" Database details ", width=10, borderwidth=2).grid(row=11, column=1, columnspan=8, sticky=EW)
####
####    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
####    Label(structure, text=" Network database: ").grid(row=12, column=1, sticky=W)
####    Label(structure, text=" Local database: ").grid(row=13, column=1, sticky=W)
####
####    db10 = StringVar()
####    db10.set(sqlprolocal)
####    db10 = Entry(structure ,relief=SUNKEN, width=55, textvariable=db10)
####    db10.grid(row=12, column=2, columnspan=2, sticky=EW)
####    db10.config(state='disabled')
####
####    db11 = StringVar()
####    db11.set(sqlprolocal)
####    db11 = Entry(structure ,relief=SUNKEN,width=55, textvariable=db11)
####    db11.grid(row=13, column=2, columnspan=2, sticky=EW)
####    db11.config(state='disabled')
####
####    networkdb = Button(structure, text="  Browse  ",  width= 10, state='disabled', command=selectNetworkDB)
####    networkdb.grid(row=12, column=6, sticky=W)
####    localdb = Button(structure, text="  Browse  ",  width= 10, state='disabled', command=selectLocalDB)
####    localdb.grid(row=13, column=6, sticky=W)
##    
##    e10.config(state='disabled')
##    e11.config(state='disabled')
##    e12.config(state='disabled')
##    e25.config(state='disabled')
##    e32.config(state='disabled')
##
##    ## Folder - structure 
##    Label1=Label(structure, text=" Automatic Folder creation ", width=10, borderwidth=2).grid(row=14, column=1, columnspan=8, sticky=EW)
##    Label(structure, text=" Exam Folder 1: ").grid(row=15, column=1, sticky=W)
##    Label(structure, text=" Exam Folder 2: ").grid(row=16, column=1, sticky=W)
##    Label(structure, text=" Exam Folder 3: ").grid(row=17, column=1, sticky=W)
##    Label(structure, text=" Exam Folder 4: ").grid(row=18, column=1, sticky=W)
##
##    e13 = StringVar()
##    e13.set(fld1)
##    e13 = Entry(structure ,relief=SUNKEN,width=55, textvariable=e13)
##    e13.grid(row=15, column=2, columnspan=2, sticky=EW)
##
##    e26a = IntVar()
##    e26a.set(pp1)
##    e26a.trace("w", lambda name, index, mode, e26a=e26a: callbacke26(e26a))
##    e26b = Checkbutton(structure, text=" ", onvalue=1,offvalue=0,variable=e26a)
##    e26b.grid(row=15, column=5, sticky=E)
##    e26b.config(state='disabled')
##
##    e14 = StringVar()
##    e14.set(fld2)
##    e14 = Entry(structure ,relief=SUNKEN,width=55, textvariable=e14)
##    e14.grid(row=16, column=2, columnspan=2, sticky=EW)
##
##    e27a = IntVar()
##    e27a.set(pp2)
##    e27a.trace("w", lambda name, index, mode, e27a=e27a: callbacke27(e27a))
##    e27b = Checkbutton(structure, text=" ", onvalue=1,offvalue=0,variable=e27a)
##    e27b.grid(row=16, column=5, sticky=E)
##    e27b.config(state='disabled')
##
##    e15 = StringVar()
##    e15.set(fld3)
##    e15 = Entry(structure ,relief=SUNKEN,width=55, textvariable=e15)
##    e15.grid(row=17, column=2, columnspan=2, sticky=EW)
##
##    e28a = IntVar()
##    e28a.set(pp3)
##    e28a.trace("w", lambda name, index, mode, e28a=e28a: callbacke28(e28a))
##    e28b = Checkbutton(structure, text=" ", onvalue=1,offvalue=0,variable=e28a)
##    e28b.grid(row=17, column=5, sticky=E)
##    e28b.config(state='disabled')
##
##    e16 = StringVar()
##    e16.set(fld4)
##    e16 = Entry(structure ,relief=SUNKEN,width=55, textvariable=e16)
##    e16.grid(row=18, column=2, columnspan=2, sticky=EW)
##
##    e29a = IntVar()
##    e29a.set(pp4)
##    e29a.trace("w", lambda name, index, mode, e29a=e29a: callbacke29(e29a))
##    e29b = Checkbutton(structure, text=" ", onvalue=1,offvalue=0,variable=e29a)
##    e29b.grid(row=18, column=5, sticky=E)
##    e29b.config(state='disabled')
##
##
##    e13.config(state='disabled')
##    e14.config(state='disabled')
##    e15.config(state='disabled')
##    e16.config(state='disabled')
## 
## 
##### Password#################################################################
##    
##    Label(pw_generator, text=" ").grid(row=1, column=2, sticky=W, columnspan=1)
##    
##    #Label(pw_generator, text=" Password Generator", width=10, borderwidth=2, relief="groove").grid(row=1, column=1, columnspan=8, sticky=EW)
##    Button(pw_generator, text=" Browse ", command=exportpwtolocation).grid(row=4, column=4, sticky=EW)
##    Button(pw_generator, text=" Save to file ", command=exportpwtodesktop).grid(row=4, column=5, sticky=EW)
##
##    ############################################################################################
##    Label(pw_generator, text="Password: ").grid(row=1, column=3, sticky=W, columnspan=1)
##
##    # Define the initial font
##    default_font = font.Font(family="Courier", size=10)
##    zoomed_font = font.Font(family="Helvetica", size=16)
##    
##    select_random_entry()
##    pw_customa = StringVar()
##    pw_customa.set(res)
##    pw_customa.trace("w", lambda name, index, mode, pw_customa=pw_customa: callback_pw(pw_customa))
##    pw_customb = Entry(pw_generator,relief=SUNKEN, width=45, font = default_font, textvariable=pw_customa)
##    pw_customb.grid(row=1, column=4, columnspan=2, sticky=W)
##    
##    # Bind mouse hover events
##    #pw_customb.bind("<Enter>", zoom_in)
##    #pw_customb.bind("<Leave>", zoom_out)
##    #pw_customb.config(state='normal')
##    #pw_customb.font.Font(family="Courier", size=14)#, weight="bold")
##
##    Refresh = Button(pw_generator, text="  Refresh  ", width= 10, state='normal', command=new_pw_guiupdatePw)
##    Refresh.grid(row=2, column=4, sticky=EW)
##
##    submit2a= StringVar()
##    submit2a.trace("w", lambda name, index, mode, submi2ta=submit2a: callback(submit2a))
##    submit2b = Button(pw_generator, text="Copy to clip",  width= 10, command=copytoclip)
##    submit2b.grid(row=2, column=5, sticky=EW)
##    #submitb.config(state='disabled')
##
##    Label(pw_generator, text="Save location:  ").grid(row=3, column=3, sticky=W, columnspan=1)#, ipady=20)
##
##    desktop = (str(os.environ['USERPROFILE']))
##    startup_e2a = StringVar()
##    startup_e2a.set(desktop+'\\Desktop')
##    startup_e2b = Entry(pw_generator, relief=SUNKEN,width=45, textvariable=startup_e2a)
##    startup_e2b.grid(row=3, column=4, sticky=EW, columnspan=2)

############################################################################################
#mantatory()
#dbcheck()
# Read analysts sql
readanalysts()
# Read profile sql
readprofile()
# Read folders sql
readfolders()
# Read cases sql
readcases()
# Read assigned cases sql
#readassignedcases()
# Read opencases sql
#readopencases()
# Read PIN Decryption log sql
#readpindecrypt()
#readpindecryptlog()
# Read notepad
#readnotepad()
# Read notifications
#readnotifications()
# Read Triage
#readtriage()

database()
os_scandir()

#setup_new_case()
#select_random_entry()
structure_gui()
############################################################################################

############################################################################################
def return_to_casework():
    global case_work_vars, casetype, photoCanvas1, left_frame, photoScrollv, photoScrollh
    casetype = ''

    try:
        for each in new_case_vars:
            print('test ',each)
            each.destroy()
        root.update()
    except:
        pass

    for each in case_work_vars:
        print('test 2',each)
        each.destroy()
        
    width  = root.winfo_screenwidth()
    height = root.winfo_screenheight()
    
    new_case_buttonb.config(text='  New Case  ')
    new_case_buttonb.config(command=setup_new_case)
    
    left_frame = Frame(case_work, borderwidth=1,width=100, height=height)
    left_frame.grid(row=1, column=2, rowspan=12, columnspan=3, sticky=NS)
    left_frame.rowconfigure(0, weight=1) 
    left_frame.columnconfigure(0, weight=1) 

    photoCanvas1 = Canvas(left_frame, width=250, height=500)#(height)-45) #width=width, height=(height)-45)
    photoCanvas1.grid(sticky=NSEW)

    dynamic_height = (len(opencases)*24)
    print('dynamic_height', dynamic_height)
    left1 = Frame(photoCanvas1, width=250, height=dynamic_height)
    photoCanvas1.create_window(0, 0, window=left1, anchor='nw')
    

    count=0
    row=1
    column=3
    case_work_dict={}
    case_work_vars=[]
    for each in opencases:
        print(each)
        entry1 = (str("e"+(str(count))))
        Entry1a = StringVar()
        button_text=(str('DFT: '+(each[1])+' - Niche: '+(each[2])+' - OIC: '+(each[4])))
        Entry1a.set(button_text.upper())
        dft=each[1]
        niche=each[2]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(photoCanvas1, text=dft, width= 57, textvariable=Entry1a, anchor="w")
        entry1.bind('<Button-1>', populate_individual_case)
        entry1.grid(row=row, column=column, sticky=EW)
        case_work_vars.append(entry1)
        case_work_dict[dft] = niche
        row+=1

    photoScrollv = Scrollbar(left_frame, orient=VERTICAL)
    photoScrollv.config(command=photoCanvas1.yview)
    photoScrollh = Scrollbar(left_frame, orient=HORIZONTAL)
    photoScrollh.config(command=photoCanvas1.xview)
    photoCanvas1.config(yscrollcommand=photoScrollv.set)
    photoCanvas1.config(xscrollcommand=photoScrollh.set)
    photoScrollv.grid(row=0, column=2, sticky="ns")
    photoScrollh.grid(row=row, column=0, sticky="ew")

    left1.bind("<Configure>", update_scrollregion)
    photoScrollv.bind("<MouseWheel>", update_scrollregion)
    root.update()
    #new_pw_guigenPw()
############################################################################################
def new_case_gui():
    global root, case_work, case_work_vars, new_case_buttona, new_case_buttonb, Refresh, e1a, e1b, e2a, e2b, e3a, e3b, e4, e4a, e4b, e5a, e5b, e6a, e6b, e7a, e7b, e24a, e24b, LoadButtona, LoadButtonb, LoadButtona1, LoadButtonb1, notepad_n59b, canvas, text_var, lab1, lab2, viewTypea, v0, v1, v2, v3, v4, v5, v6, delEntry_a, delEntry_b, new_case, newcase, new_case_e1, e1a, e1b, e2, e3, e4, e5, e6, e7, e8, e9, submitb, new_case_e1a, label1, submit, edit1b, e2a, e3a, e4a, e5a, e6a, e7a, e8a, e9a, launch, e21a, e21b, e22a, e22b, startup_e1a, startup_e1b, photoCanvas1, photoScrollv, photoScrollh, left_frame
    try:
        if 'normal' == pw_generator.state():
            pw_generator.destroy()
        elif 'normal' == new_triage.state():
            new_triage.destroy()
        else:
            pass
    except:
        pass
        
    root = Tk()
    #new_case.iconbitmap(iconfile)
    root.geometry('610x410')
    root.title("CASE - New")
    root.resizable(width=False, height=False)
    #menu_new_case()
    root.protocol('WM_DELETE_WINDOW', sys.exit)
    width  = root.winfo_screenwidth()
    height = root.winfo_screenheight()
    print(width)
    print(height)
    #menu()
    frames=["case_work", "new_case"] #"databases", "notepad", "settings", "about"]   "graykey",
    rows=31
    columns=7
    tabs = ttk.Notebook(root)

    tabs.pack(fill='both', expand=Y)
    case_work=ttk.Frame()
    #new_case=ttk.Frame()
    tabs.add(case_work,text='Casework')
    #tabs.add(new_case,text='New Case')
    # Set row height here
    height=4
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count < rows:
                if count == 0:
                    Label(case_work, text=" ", width=2).grid(row=count, column=0, sticky=EW, ipady=1)
                    count+=1
                else:
                    Label(case_work, text=" ", width=2).grid(row=count, column=0, sticky=EW, ipady=height)
                    count+=1
            while count2 < columns:
                Label(case_work, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
       
    img = Image.open("C:\\Case Creator\\_logos\\Surrey-Sussex-Police-logo.jpg")
    #img = img.resize((300, 100), Image.ANTIALIAS)
    img = img.resize((150, 125), Image.ADAPTIVE)
    img = ImageTk.PhotoImage(img)
    panel = Label(case_work, image=img, width=150, background="white")
    panel.image = img
    panel.grid(row=1, column=1, columnspan=1, rowspan=4, sticky=EW)
    
    new_case_buttona = StringVar()
    new_case_buttona.trace("w", lambda name, index, mode, new_case_buttona=new_case_buttona: callback(new_case_buttona))
    new_case_buttonb = Button(case_work, text="  New Case  ", width= 20, command=setup_new_case)
    new_case_buttonb.grid(row=6, column=1, sticky=W)
    new_case_buttonb.config(state='normal')

    left_frame = Frame(case_work, borderwidth=1,width=100, height=200)
    left_frame.grid(row=1, column=2, rowspan=11, columnspan=3, sticky=NS)
##    left_frame.rowconfigure(0, weight=1) 
##    left_frame.columnconfigure(0, weight=1) 

    photoCanvas1 = Canvas(left_frame)#, width=250, height=200)#(height)-45) #width=width, height=(height)-45)
    photoCanvas1.grid(sticky=NSEW)
    
    row=1
    column=3
    case_work_dict={}
    case_work_vars=[]
    for each in opencases:
        print(each)
        entry1 = (str("e"+(str(count))))
        Entry1a = StringVar()
        button_text=(str('DFT: '+(each[1])+' - Niche: '+(each[2])+' - OIC: '+(each[4])))
        Entry1a.set(button_text.upper())
        dft=each[1]
        niche=each[2]
        Entry1a.trace("w", lambda name, index, mode, Entry1a=Entry1a: callback(Entry1a))
        entry1 = Button(photoCanvas1, text=dft, width= 57, textvariable=Entry1a, anchor="w")
        entry1.bind('<Button-1>', populate_individual_case)
        entry1.grid(row=row, column=column, sticky=EW)
        case_work_vars.append(entry1)
        case_work_dict[dft] = niche
        row+=1
        
    dynamic_height = (len(opencases)*24)
    print('dynamic_height', dynamic_height)
    left1 = Frame(photoCanvas1, width=0, height=750)
    photoCanvas1.create_window(0, 0, window=left1, anchor='nw')
    
    #case_work_vars.append(photoCanvas1)
    '''img = Image.open("C:\\Case Creator\\_logos\\Surrey-Sussex-Police-logo.jpg")
    #img = img.resize((300, 100), Image.ANTIALIAS)
    img = img.resize((150, 125), Image.ADAPTIVE)
    img = ImageTk.PhotoImage(img)
    panel = Label(new_case, image=img, width=150, background="white")
    panel.image = img
    panel.grid(row=1, column=1, columnspan=1, rowspan=4, sticky=EW)
        
    #dbconnect = Label(new_case, text='',width=30)
    #dbconnect.grid(row=1, column=1, columnspan=6, sticky=EW)
    #dbconnect.config(state='disabled')

    Label1=Label(new_case, text=" New Case ", width=15, borderwidth=2, relief="groove").grid(row=1, column=3, columnspan=3, sticky=EW)
    #Label(new_case, text=" Password Generator", width=10, borderwidth=2, relief="groove").grid(row=14, column=1, columnspan=6, sticky=EW)
    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
    Label(new_case, text=" DFT Ref(YY-DFTREF):").grid(row=2, column=3, sticky=W)

    LoadButtona = StringVar()
    LoadButtona.trace("w", lambda name, index, mode, LoadButtona=LoadButtona: callback(LoadButtona))
    LoadButtonb = Button(new_case, text="  Load  ", width= 10, command=loadcases_dft)
    LoadButtonb.grid(row=2, column=5, sticky=W)
    LoadButtonb.config(state='disable')

    LoadButtona1 = StringVar()
    LoadButtona1.trace("w", lambda name, index, mode, LoadButtona=LoadButtona: callback(LoadButtona1))
    LoadButtonb1 = Button(new_case, text="  Load  ", width= 10, command=loadcases_crime)
    LoadButtonb1.grid(row=3, column=5, sticky=W)
    LoadButtonb1.config(state='disable')

    #Button(new_case, text="  Lock  ", width= 10, command=lockCase).grid(row=4, column=4, sticky=W)

    Label(new_case, text=" Crime Reference: ").grid(row=3, column=3, sticky=W)
    Label(new_case, text=" Exhibit Reference: ").grid(row=4, column=3, sticky=W)
    Label(new_case, text=" Bag Seal Reference: ").grid(row=5, column=3, sticky=W)
    Label(new_case, text=" OIC/Contact: ").grid(row=6, column=3, sticky=W)
    Label(new_case, text=" Operation Name: ").grid(row=7, column=3, sticky=W)
    Label(new_case, text=" Suspect Name: ").grid(row=8, column=3, sticky=W)
    Label(new_case, text=" Property Reference: ").grid(row=9, column=3, sticky=W)
    Label(new_case, text=" Date of offence: ").grid(row=10, column=3, sticky=W)


    ############################################################################################
    def callbacke1(sv):
        print(sv.get())
        entry = e1a.get()
        #print(str('entry: ' + entry[-1:3]))

        if len(entry) == 3:
            if entry[-1:3] != '-' :
                e1a.set(entry[:2])
            else:
                e1a.set(entry[:7])
        else:
            e1a.set(entry[:7])

    ############################################################################################
    row = 2
    column = 4
    e1a = StringVar()
    e1a.trace("w", lambda name, index, mode, e1a=e1a: callback1(e1a))
    e1b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e1a)
    e1b.grid(row=2, column=column)
    e1b.focus()

    e2a = StringVar()
    e2a.trace("w", lambda name, index, mode, e2a=e2a: callback2(e2a))
    e2b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e2a)
    e2b.grid(row=3, column=column)

    e3a = StringVar()
    e3a.trace("w", lambda name, index, mode, e3a=e3a: callback3(e3a))
    e3b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e3a)
    e3b.grid(row=4, column=column)

    e4a = StringVar() # bag seal
    e4a.trace("w", lambda name, index, mode, e4a=e4a: callback(e4a))
    e4b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e4a)
    e4b.grid(row=5, column=column)

    e5a = StringVar() # OIC
    e5a.trace("w", lambda name, index, mode, e5a=e5a: callback(e5a))
    e5b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e5a)
    e5b.grid(row=6, column=column)

    e6a = StringVar()
    e6a.trace("w", lambda name, index, mode, e6a=e6a: callback(e6a))
    e6b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e6a)
    e6b.grid(row=7, column=column)

    e7a = StringVar()
    e7a.trace("w", lambda name, index, mode, e7a=e7a: callback(e7a))
    e7b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e7a)
    e7b.grid(row=8, column=column)

    e8 = Entry(new_case,relief=SUNKEN,width=30)
    e8.grid(row=9, column=column)

    submita = StringVar()
    submita.trace("w", lambda name, index, mode, submita=submita: callback(submita))
    submitb = Button(new_case, text="  Submit  ", width= 10, command=duplicate_entry)
    submitb.grid(row=10, column=5, sticky=W)
    submitb.config(state='normal')

    e24a = StringVar()
    e24b = Entry(new_case,relief=SUNKEN,width=30, textvariable=e24a)
    e24b.grid(row=10, column=column)'''

    photoScrollv = Scrollbar(left_frame, orient=VERTICAL, command=photoCanvas1.yview)
    photoScrollh = Scrollbar(left_frame, orient=HORIZONTAL, command=photoCanvas1.xview)
    photoCanvas1.config(yscrollcommand=photoScrollv.set)
    photoCanvas1.config(xscrollcommand=photoScrollh.set)
    photoScrollv.grid(row=0, column=2, sticky="ns")
    photoScrollh.grid(row=row, column=0, sticky="ew")

    left_frame.bind("<Configure>", update_scrollregion)
    photoScrollv.bind("<MouseWheel>", update_scrollregion)

        
    analyst = e9a.get()

    Label(settings_tab, text=" Start-up: ").grid(row=21, column=1, sticky=W)
    estartupa = StringVar()
    estartupa.set(startup_method)
    estartupa.trace("w", lambda name, index, mode, estartupa=estartupa: callback(estartupa))
    estartupb = OptionMenu(settings_tab, estartupa, *startupmode, command=change_startup)
    estartupb.bind('<Triple-Button-1>', lambda x: print('Hi'))
    estartupb.grid(row=21, column=2, sticky=W)
    estartupb.config(width=30)
    estartupb.config(state='disabled')
    
    root.lift()
    root.attributes('-topmost',True)
    #root.after_idle(new_case.attributes,'-topmost',False)
    root.mainloop()
############################################################################################
def menu():
    menubar = Menu(root)

    casesmenu = Menu(root)
    casesmenu.add_command(label="New Case", command=setup_new_case)
    casesmenu.add_separator()

    #casesmenu.add_command(label="New PD entry", command=add_new_PD_entry)
    #casesmenu.add_separator()
    
    #casesmenu.add_command(label="New Triage entry", command=new_triage_gui)
    #casesmenu.add_separator()
    
    casesmenu.add_command(label="Exit", command=root.quit)
    menubar.add_cascade(label="New", menu=casesmenu)
    
    #exportmenu = Menu(new_case)
    #exportmenu.add_command(label="Overview Export", command=export_overview)
    #menubar.add_cascade(label="Export", menu=exportmenu)
                        
    settingsmenu = Menu(root)
    settingsmenu.add_command(label="Admin", command=admin_login_gui)
    settingsmenu.add_separator()
    settingsmenu.add_command(label="Settings", command=settings_gui)
    menubar.add_cascade(label="Settings", menu=settingsmenu)

    passwordmenu = Menu(root)
    passwordmenu.add_command(label="Generate Password", command=pw_generator_gui)
    passwordmenu.add_command(label="Search Case Password", command=search_pw)
    menubar.add_cascade(label="Password", menu=passwordmenu)
    
    aboutmenu = Menu(root)
    aboutmenu.add_command(label="About", command=about_gui)
    menubar.add_cascade(label="About", menu=aboutmenu)
    
    root.config(menu=menubar)
############################################################################################
## Settings
def settings_gui():
    global networkdb, localdb, db10, db11, settings_tab, change_analyst_button, settings_lock, casedir, tempdir, sfrdir, disclosuredir, contempdir, e33b, e34b, e35b, e36b, e37b, e38b, estartupa, estartupb, e12, e13, e14, e15, e16, e17, e18, Refresh, e25, e32, e1a, e1b, e2a, e2b, e3a, e3b, e4, e5a, e5b, e6a, e6b, e7a, e7b, e9a, e9b, e10, e11, e24a, e24b, e26b, e27b, e28b, e29b, e30b, e31b, LoadButtona, LoadButtonb, notepad_n59b, canvas, text_var, lab1, lab2, viewTypea, v0, v1, v2, v3, v4, v5, v6, delEntry_a, delEntry_b, new_case, newcase, new_case_e1, e1a, e1b, e2, e3, e4, e5, e6, e7, e8, e9, submitb, new_case_e1a, label1, submit, edit1b, e2a, e3a, e4a, e5a, e6a, e7a, e8a, e9a, launch, e21a, e21b, e22a, e22b, startup_e1a, startup_e1b

    gui = ("new_case")
    settings_tab = Tk()
    #new_case.iconbitmap(iconfile)
    settings_tab.geometry('610x720')
    settings_tab.title("CASE - Settings")
    settings_tab.resizable(width=False, height=False)
    settings_tab.protocol('WM_DELETE_WINDOW', closesettings)

    
    frames=["new_case"] #"databases", "notepad", "settings_tab", "about"]   "graykey",
    rows=31
    columns=7
    # Set row height here
    height=5
    for frame in frames:
        count=0
        count2=0
        frame=(str(frame))
        if frame == frames[0]:
            while count < rows:
                Label(settings_tab, text=" ", width=2).grid(row=count, column=0, sticky=EW, ipady=height)
                count+=1
            while count2 < columns:
                Label(settings_tab, text=" ", width=2).grid(row=1, column=count2, sticky=EW, ipady=height)
                count2+=1
        
    dbconnect = Label(settings_tab, text='',width=30)
    dbconnect.grid(row=1, column=1, columnspan=6, sticky=EW)
    dbconnect.config(state='disabled')
    settings_lock = Button(settings_tab, text=" Unlock ", width= 10, command=0)#settings_unlock)
    settings_lock.grid(row=21, column=6, sticky=E)
    settings_lock.focus_set()

    Label1=Label(settings_tab, text=" Case & template settings_tab ", width=10, borderwidth=2, relief="groove").grid(row=1, column=1, columnspan=8, sticky=EW)
    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
    Label(settings_tab, text=" Case Directory: ").grid(row=2, column=1, sticky=W)
    casedir = Button(settings_tab, text="  Browse  ", width= 10, command=selectExamDir, state='disabled')
    casedir.grid(row=2, column=6, sticky=W)
    tempdir = Button(settings_tab, text="  Browse  ", width= 10, command=selectTempDir, state='disabled')
    tempdir.grid(row=3, column=6, sticky=W)

    Label1=Label(settings_tab, text=" Document settings_tab ", width=10, borderwidth=2, relief="groove").grid(row=4, column=1, columnspan=8, sticky=EW)
    Label(settings_tab, text=" Template Directory: ").grid(row=3, column=1, sticky=W)
    Label(settings_tab, text=" Contemp Notes: ").grid(row=5, column=1, sticky=W)
    Label(settings_tab, text=" SFR: ").grid(row=7, column=1, sticky=W)
    Label(settings_tab, text=" Disclosure Cert: ").grid(row=9, column=1, sticky=W)
    ############################################################################################
    def callbacke1(sv):
        entry = e1a.get()
        if len(entry) == 3:
            if entry[-1:3] != '-' :
                e1a.set(entry[:2])
            else:
                e1a.set(entry[:7])
        else:
            e1a.set(entry[:7])

    ############################################################################################
    e10 = StringVar()
    e10.set(case_dir)
    e10 = Entry(settings_tab,relief=SUNKEN, width=55, textvariable=e10)
    e10.grid(row=2, column=2, columnspan=2, sticky=EW)
    e10.config(state='disabled')

    e11 = StringVar()
    e11.set(template_dir)
    e11 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e11)
    e11.grid(row=3, column=2, columnspan=2, sticky=EW)
    e11.config(state='disabled')

    e12 = StringVar()
    e12.set(contemp_file) #notes_file
    e12 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e12)
    e12.grid(row=5, column=2, columnspan=2, sticky=EW)

    #Copy notes_file to case checkbox
    e33a = IntVar()
    e33a.set(contemp_copy)
    e33a.trace("w", lambda name, index, mode, e33a=e33a: callbacke33(e33a))
    e33b = Checkbutton(settings_tab, text="Copy to case", onvalue=1,offvalue=0,variable=e33a, command=e33checkbox)
    e33b.grid(row=6, column=2, sticky=EW)
    e33b.config(state='disabled')

    #Populate notes_file checkbox
    e34a = IntVar()
    e34a.set(contemp_populate)
    e34a.trace("w", lambda name, index, mode, e34a=e34a: callbacke34(e34a))
    e34b = Checkbutton(settings_tab, text="Populate?", onvalue=1,offvalue=0,variable=e34a)
    e34b.grid(row=6, column=3, sticky=W)
    e34b.config(state='disabled')

    e25 = StringVar()
    e25.set(sfr_file) #sfr_file
    e25 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e25)
    e25.grid(row=7, column=2, columnspan=2, sticky=EW)

    #Copy sfr_file to case checkbox
    e35a = IntVar()
    e35a.set(sfr_copy)
    e35a.trace("w", lambda name, index, mode, e35a=e35a: callbacke35(e35a))
    e35b = Checkbutton(settings_tab, text="Copy to case", onvalue=1,offvalue=0,variable=e35a, command=e35checkbox)
    e35b.grid(row=8, column=2, sticky=EW)
    e35b.config(state='disabled')

    #Populate sfr_file checkbox
    e36a = IntVar()
    e36a.set(sfr_populate)
    e36a.trace("w", lambda name, index, mode, e36a=e36a: callbacke36(e36a))
    e36b = Checkbutton(settings_tab, text="Populate?", onvalue=1,offvalue=0,variable=e36a)
    e36b.grid(row=8, column=3, sticky=W)
    e36b.config(state='disabled')

    e32 = StringVar()
    e32.set(disclosure_file) #disclosure_file
    e32 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e32)
    e32.grid(row=9, column=2, columnspan=2, sticky=EW)

    #Copy disclosure_file to case checkbox
    e37a = IntVar()
    e37a.set(disclosure_copy)
    e37a.trace("w", lambda name, index, mode, e37a=e37a: callbacke37(e37a))
    e37b = Checkbutton(settings_tab, text="Copy to case", onvalue=1,offvalue=0,variable=e37a, command=e37checkbox)
    e37b.grid(row=10, column=2, sticky=EW)
    e37b.config(state='disabled')

    #Populate disclosure_file checkbox
    e38a = IntVar()
    e38a.set(disclosure_populate)
    e38a.trace("w", lambda name, index, mode, e38a=e38a: callbacke38(e38a))
    e38b = Checkbutton(settings_tab, text="?", onvalue=1,offvalue=0,variable=e38a)
    e38b.grid(row=10, column=3, sticky=W)
    e38b.config(state='disabled')

    contempdir = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectContemp)
    contempdir.grid(row=5, column=6, sticky=W)
    sfrdir = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectSfr)
    sfrdir.grid(row=7, column=6, sticky=W)
    disclosuredir = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectDisclosure)
    disclosuredir.grid(row=9, column=6, sticky=W)

    Label(settings_tab, text=" Database details ", width=10, borderwidth=2, relief="groove").grid(row=11, column=1, columnspan=8, sticky=EW)

    ## Relief options for labels - raised, sunken, flat, ridge, solid, groove
    Label(settings_tab, text=" Network database: ").grid(row=12, column=1, sticky=W)
    Label(settings_tab, text=" Local database: ").grid(row=13, column=1, sticky=W)

    db10 = StringVar()
    db10.set(sqlprolocal)
    db10 = Entry(settings_tab,relief=SUNKEN, width=55, textvariable=db10)
    db10.grid(row=12, column=2, columnspan=2, sticky=EW)
    db10.config(state='disabled')

    db11 = StringVar()
    db11.set(sqlprolocal)
    db11 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=db11)
    db11.grid(row=13, column=2, columnspan=2, sticky=EW)
    db11.config(state='disabled')

    networkdb = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectNetworkDB)
    networkdb.grid(row=12, column=6, sticky=W)
    localdb = Button(settings_tab, text="  Browse  ",  width= 10, state='disabled', command=selectLocalDB)
    localdb.grid(row=13, column=6, sticky=W)
    
    e10.config(state='disabled')
    e11.config(state='disabled')
    e12.config(state='disabled')
    e25.config(state='disabled')
    e32.config(state='disabled')

    ## Folder - settings_tab
    Label1=Label(settings_tab, text=" Automatic Folder creation ", width=10, borderwidth=2, relief="groove").grid(row=14, column=1, columnspan=8, sticky=EW)
    Label(settings_tab, text=" Exam Folder 1: ").grid(row=15, column=1, sticky=W)
    Label(settings_tab, text=" Exam Folder 2: ").grid(row=16, column=1, sticky=W)
    Label(settings_tab, text=" Exam Folder 3: ").grid(row=17, column=1, sticky=W)
    Label(settings_tab, text=" Exam Folder 4: ").grid(row=18, column=1, sticky=W)

    e13 = StringVar()
    e13.set(fld1)
    e13 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e13)
    e13.grid(row=15, column=2, columnspan=2, sticky=EW)

    e26a = IntVar()
    e26a.set(pp1)
    e26a.trace("w", lambda name, index, mode, e26a=e26a: callbacke26(e26a))
    e26b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e26a)
    e26b.grid(row=15, column=5, sticky=E)
    e26b.config(state='disabled')

    e14 = StringVar()
    e14.set(fld2)
    e14 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e14)
    e14.grid(row=16, column=2, columnspan=2, sticky=EW)

    e27a = IntVar()
    e27a.set(pp2)
    e27a.trace("w", lambda name, index, mode, e27a=e27a: callbacke27(e27a))
    e27b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e27a)
    e27b.grid(row=16, column=5, sticky=E)
    e27b.config(state='disabled')

    e15 = StringVar()
    e15.set(fld3)
    e15 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e15)
    e15.grid(row=17, column=2, columnspan=2, sticky=EW)

    e28a = IntVar()
    e28a.set(pp3)
    e28a.trace("w", lambda name, index, mode, e28a=e28a: callbacke28(e28a))
    e28b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e28a)
    e28b.grid(row=17, column=5, sticky=E)
    e28b.config(state='disabled')

    e16 = StringVar()
    e16.set(fld4)
    e16 = Entry(settings_tab,relief=SUNKEN,width=55, textvariable=e16)
    e16.grid(row=18, column=2, columnspan=2, sticky=EW)

    e29a = IntVar()
    e29a.set(pp4)
    e29a.trace("w", lambda name, index, mode, e29a=e29a: callbacke29(e29a))
    e29b = Checkbutton(settings_tab, text=" ", onvalue=1,offvalue=0,variable=e29a)
    e29b.grid(row=18, column=5, sticky=E)
    e29b.config(state='disabled')


    e13.config(state='disabled')
    e14.config(state='disabled')
    e15.config(state='disabled')
    e16.config(state='disabled')
 
## Folder - settings_tab
    Label1=Label(settings_tab, text=" Analyst settings_tab ", width=10, borderwidth=2, relief="groove").grid(row=19, column=1, columnspan=8, sticky=EW)

    Label(settings_tab, text=" Analyst: ").grid(row=20, column=1, sticky=W)
    e9a = StringVar()
    e9a.set(examiner)
    e9a.trace("w", lambda name, index, mode, e9a=e9a: callback(e9a))
    e9b = OptionMenu(settings_tab, e9a, *analysts, command=updateexaminer)
    e9b.config(width=30)
    e9b.config(state='disabled')
    e9b.grid(row=20, column=2, sticky=W)
    
    analyst = e9a.get()

    Label(settings_tab, text=" Start-up: ").grid(row=21, column=1, sticky=W)
    estartupa = StringVar()
    estartupa.set(startup_method)
    estartupa.trace("w", lambda name, index, mode, estartupa=estartupa: callback(estartupa))
    estartupb = OptionMenu(settings_tab, estartupa, *startupmode, command=change_startup)
    estartupb.bind('<Triple-Button-1>', lambda x: print('Hi'))
    estartupb.grid(row=21, column=2, sticky=W)
    estartupb.config(width=30)
    estartupb.config(state='disabled')

    settings_tab.lift()
    settings_tab.attributes('-topmost',True)
    #settings.after_idle(new_case.attributes,'-topmost',False)
    settings_tab.mainloop()
############################################################################################

##
##if startup_method == 'Case':
##    lockstate=0
##
##    ##login_gui()
##    # BIND EXAMPLE
##    #startup.bind('<Return>', lambda x: check())
##
##    global triage, tabs, notepad_n59b, canvas, text_var, lab1, lab2, viewTypea, v0, v1, v2, v3, v4, v5, v6, delEntry_a, delEntry_b, startup, newcase, startup_e1, e1a, e1b, e2, e3, e4, e5, e6, e7, e8, e9, submitb, startup_e1a, label1, submit, edit1b, e2a, e3a, e4a, e5a, e6a, e7a, e8a, e9a, launch
##
##    gui = ("startup")
##    root = Tk()
##
##    root.iconbitmap(iconfile)
##    root.title("CASE - "+analyst)
##    root.resizable(width=True, height=True)
##    root.geometry('1800x720')
##    width  = root.winfo_screenwidth()
##    height = root.winfo_screenheight()
##    print(width)
##    print(height)
##
##    root.attributes("-fullscreen", False)
##    menu()
##
##    #FRAME SETUP / SPACING
##    tabs = ttk.Notebook(root)
##
##    tabs.pack(fill='both', expand=Y)
##
##    open_cases=ttk.Frame()
##    #graykey=ttk.Frame()
##    notepad=ttk.Frame()
##    #notifications=ttk.Frame()
##    #settings=ttk.Frame()
##    #about=ttk.Frame()
##    #triage=ttk.Frame()
##
##    tabs.add(open_cases,text='Casework')
##    #tabs.add(graykey,text='Graykey')
##    tabs.add(notepad,text='Notepad')
##    #tabs.add(notifications,text='Notifications')
##    #tabs.add(settings,text='Settings')
##    #tabs.add(about,text='About')
##    #tabs.add(triage,text='Triage')
##
##    ### Notepad tab (scroll WORKING)
##    notepad_scrollbar = Scrollbar(notepad)
##    #Label(notepad, text=" Results/Findings: ").grid(row=5, column=0, sticky=W, pady=0, columnspan=6)
##    notepad_n59b=Text(notepad, wrap=WORD, width=width, height=50, yscrollcommand=notepad_scrollbar.set)
##    notepad_n59b.grid(row=1, column=0, columnspan=1, sticky=EW, rowspan=1)
##    print('Len notepadentries:',(len(notepadentries)))
##    if (len(notepadentries))==2:
##            print('in func')
##            notepadnotes1_text = (str(notepadentries[1]))
##            notepadnotes1_text = (str.replace(notepadnotes1_text,"\\n","\n"))
##            notepad_n59b.insert(INSERT, notepadnotes1_text)
##    else:
##        pass
##    notepad_n59b.grid_propagate(False)
##    notepad_n59b.bind('<KeyRelease>', notepadnotes1)
##    notepad_scrollbar.config( command = notepad_n59b.yview)
##    notepad_scrollbar.bind('<MouseWheel>', notepad_n59b)
##    notepad_scrollbar.grid(row=1, column=1, rowspan=1,  sticky='NS')
##
##    ############################################################################################
##
##    left_frame = Frame(open_cases, borderwidth=1,relief=RIDGE)
##    left_frame.grid(row=0, column=1, columnspan=10, sticky=E) 
##
##    photoCanvas1 = Canvas(left_frame, width=1000, height=(height)-45) #width=width, height=(height)-45)
##    photoCanvas1.grid()#sticky=NSEW)
##
##    left1 = Frame(photoCanvas1, width=width, height=10000)
##    photoCanvas1.create_window(0, 0, window=left1, anchor='nw')
##    ############################################################################################
##
##    right_frame = Frame(open_cases, borderwidth=1,relief=RIDGE)
##    right_frame.grid(row=0, column=11, columnspan=10, sticky=W)
##
##    photoCanvas2 = Canvas(right_frame, width=800, height=(height)-45) #width=width, height=(height)-45)
##    photoCanvas2.grid()#sticky=NSEW)
##
##    right1 = Frame(photoCanvas2, width=width, height=10000)
##    photoCanvas2.create_window(0, 0, window=right1, anchor='nw')
##    ############################################################################################
##
## 
##
##    ############################################################################################
##    # Monitoring tab changes
##    tabs.bind('<<NotebookTabChanged>>', lambda x: ontabchangelock())
##    #everyfive()
##
##    root.mainloop()
##else:
##    new_case_gui()
##
##
##
##
